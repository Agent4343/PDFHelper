"""
Procedure Builder routes — AI-guided procedure writing with business templates.

Covers: /procedures, /procedures/{id}, /procedures/{id}/chat,
        /procedures/{id}/generate, /procedures/{id}/download
"""

import json
import uuid
from datetime import datetime, timezone

from anthropic import Anthropic
from fastapi import APIRouter, Depends, HTTPException, Query, Request, UploadFile, File
from fastapi.responses import Response, StreamingResponse

from auth import verify_auth, get_db
from config import CHAT_MODEL, CHAT_MAX_TOKENS
from database import (
    DBDocument,
    DBProcedureSession,
    DBProcedureMessage,
    DBAppSetting,
    SessionLocal,
)
from helpers import _encrypt_text, _decrypt_text, _safe_decrypt, _load_stored_text, _stored_text_to_structured

FACILITY_PRESETS = {
    "Hebron": {
        "crafts": ["Operations", "Instrumentation", "Mechanical", "Electrical", "Process"],
        "doc_prefix": "CAHE-EC-OOPRO",
    },
    "Hibernia": {
        "crafts": ["Operations", "Instrumentation", "Mechanical", "Electrical", "Process"],
        "doc_prefix": "HS-O-O",
    },
}

router = APIRouter(dependencies=[Depends(verify_auth)])


def _extract_docx_structured(file_bytes: bytes) -> str:
    """Extract structured text from a .docx file preserving headings, tables, notes."""
    from docx import Document as DocxDocument
    import io

    doc = DocxDocument(io.BytesIO(file_bytes))
    parts = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style_name = (para.style.name or "").lower()
        if "heading 1" in style_name or "proc - heading 1" in style_name:
            parts.append(f"# {text}")
        elif "heading 2" in style_name or "proc - heading 2" in style_name:
            parts.append(f"## {text}")
        elif "heading 3" in style_name or "proc - heading 3" in style_name:
            parts.append(f"### {text}")
        elif "caution" in style_name:
            parts.append(f"CAUTION: {text}")
        elif "warning" in style_name:
            parts.append(f"WARNING: {text}")
        elif "bullet" in style_name or "list" in style_name:
            parts.append(f"- {text}")
        else:
            parts.append(text)

    for table in doc.tables:
        header_cells = [cell.text.strip() for cell in table.rows[0].cells]
        parts.append("")
        parts.append("| " + " | ".join(header_cells) + " |")
        parts.append("| " + " | ".join(["---"] * len(header_cells)) + " |")
        for row in table.rows[1:]:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            note_text = " ".join(cells).upper()
            if note_text.startswith("NOTE:") or note_text.startswith("CAUTION:"):
                parts.append(" ".join(c for c in cells if c))
            else:
                parts.append("| " + " | ".join(cells) + " |")
        parts.append("")

    return "\n".join(parts)

PROCEDURE_SYSTEM_PROMPT = """You are an expert technical procedure writer for upstream oil and gas operations, trained on:
- PPA AP-907-005 Procedure Writer's Manual (Rev. 3)
- Upstream Best Practices: Procedure Writing Rules
- Operations Integrity Protocol 6.1 (OIMS Element 6)
- Upstream Procedure Tools: Task Analysis
- Safety Critical Task Analysis (SCTA) programme requirements
- Critical Task Execution Playbook (Revision 9) — Precise Execution of Critical Tasks

Your job is to help create clear, precise, and safe work procedures that meet OIMS 6.1 requirements and human performance principles.

=== GATHERING INFORMATION ===
Ask focused questions ONE AT A TIME in this order:
1. Facility/site name and craft/discipline (e.g., Facility: Hebron, Craft: Operations)
2. Procedure title and designation number (use logical numbering: SITE-DISC-TYPE-SYS-SEQ-000, e.g., CAHE-EC-OOPRO-01-006-1003-000)
3. Revision number (e.g., Rev. 0, Rev. D18)
4. Revalidation date
5. Purpose — what, when, and why (do NOT simply repeat the title)
6. Scope — activities covered, boundaries, applicable personnel and equipment
7. References and commitments — regulatory docs, operating experience, P&IDs, vendor manuals
8. Definitions — terms unique to this procedure (alphabetical, do not define self-explanatory terms)
9. Responsibilities — who does what (high-level summary, not a repeat of steps)
10. Precautions — equipment/personnel/public protection measures (state effect AND cause)
11. Limitations — specific regulatory or administrative limits with values
12. Prerequisites — PPE, materials, special tools, other prerequisites (LMRA, JSA, valve line-up, etc.)
13. Step-by-step instructions — walk through each action
14. Acceptance criteria — quantitative/qualitative pass/fail criteria
15. Attachments needed — data sheets, checklists, figures, P&ID excerpts, record of results
16. Is this a Safeguard Critical Task per CTE Playbook? If so:
    - Identify all safeguard critical steps
    - Determine hold points (pre-step or post-step)
    - Specify independent verification method (1=in-person not present, 2=in-person present, 3=remote radio, 4=remote electronic)
    - State consequence of failure for each safeguard critical step
17. Who performs each step? Use standard job roles (e.g., CCR Operator, Operations Technician, Instr Tech, FGS Operator)

=== PROCEDURE STRUCTURE ===
The document generator automatically creates the header table, TOC, and page numbering.
Output the procedure content using these EXACT section headings and numbers:

# 1. PURPOSE AND SCOPE
## 1.1. Purpose
## 1.2. Scope

# 2. PRECAUTIONS AND LIMITATIONS
## 2.1. Precautions
## 2.2. Limitations

# 3. PREREQUISITES
## 3.1. Personal Protective Equipment, PPE
## 3.2. Materials
## 3.3. Special Tools and Equipment
## 3.4. Other Prerequisites

| Step | Action / Remarks | Who | Check |
(use this table for prerequisite verification steps like LMRA, JSA, valve line-up)

# 4. INSTRUCTIONS
## 4.1. [First procedure section]
## 4.2. [Second procedure section]
(each subsection gets its own action step table)

# 5. REFERENCES AND COMMITMENT
## 5.1. Performance References
## 5.2. Commitments References
## 5.3. Developmental References

## Attachment 1 - [Title]
## Attachment 2 - [Title]
(attachments start on a new page automatically)

=== WRITING RULES ===

ACTION STEPS:
- Every action step starts with an ACTION VERB in UPPERCASE BOLD (e.g., OPEN, CLOSE, VERIFY, CHECK, RECORD, PERFORM)
- One action per step — never combine two actions
- Active voice only — the step directs the user to act
- Include WHO performs the action and a checkbox

Step format — output as a markdown table with these 4 columns:
| Step | Action / Remarks | Who | Check |
| 1 | OPEN inlet valve XX-XXX-001 to Amine Circulating Pump XX-XXXX | Ops Tech | ☐ |
(Step numbers auto-increment; leave Step column empty or use "-" for auto-numbering)

EMPHASIS TECHNIQUES:
- Action verbs: UPPERCASE BOLD (e.g., OPEN, CLOSE, VERIFY)
- Conditional/logic terms: UPPERCASE, UNDERLINED, BOLD (IF, THEN, WHEN, AND, OR, NOT, WHILE)
- Component positions: UPPERCASE (OPEN, CLOSED, ON, OFF, AUTO)
- Component noun names: Title Case (Heater Pump, Amine Discharge Valve)
- Locations: inside parentheses — (inside the MMC)
- Condition and action on SEPARATE lines

CONDITIONAL STEPS:
- IF introduces a condition that may or may not be true
- WHEN introduces a condition expected to occur
- THEN goes between condition and action (never between actions)
- IF AT ANY TIME introduces a condition that may occur during procedure execution
- WHILE introduces a continuous action
- Do NOT use AND/OR construction
- State conditions positively (IF valve is OPEN, not IF valve is NOT CLOSED)
- For 3+ conditions, use a decision table

IF/THEN table format:
| Step | Action / Remarks | Who | Check |
| - | IF Temperature exceeds 100C | | |
| - | THEN OPEN XX-XXX-0011 inlet to XXX exchanger | Ops Tech | ☐ |

NOTES, CAUTIONS, AND WARNINGS:
- Place BEFORE the step they apply to (never after)
- Sequence: Note first, then Caution, then Warning (most important closest to step)
- Must appear on the same page as the impacted step
- Written in passive voice, short and concise
- Must NOT contain action steps or implied instructions
- If removed, procedure performance would not be affected
- WARNING: personnel injury, loss of life, health hazards — format: ! WARNING: [text]
- CAUTION: equipment damage, process risk — format: CAUTION: [text]
- NOTE: supplemental/explanatory information — format: Note: [text]

HOLD POINTS:
- A pre-selected step beyond which work may NOT proceed until required action is performed
- Identified by SCTA for HC scenarios or by SME
- Format: "Hold Point" label before the action step
- Require explicit authorization to proceed past

INDEPENDENT VERIFICATION:
- Used for safeguard critical steps
- Two individuals working independently to confirm component condition
- Format includes Name/Sign fields after the step

SAFEGUARD CRITICAL STEPS:
- Only used when identified by SCTA
- Preceded by a WARNING box: "WARNING - SAFEGUARD CRITICAL STEP"
- State: independent verification required, what happens if step is done wrong (clear hazard/consequence)
- Easy to understand and execute (limit potential for error)
- Consider hold points for steps requiring additional review

STEP NUMBERING:
- Up to 4 levels: 1. / a. / (1) / (a)
- Alphanumeric steps performed in written order unless stated otherwise
- Bulleted steps within a single alphanumeric step may be performed in any order
- Sections: 1.0 TITLE / 1.1 Subtitle / 1.1.1 Subtitle

VOCABULARY:
- No ambiguous words: replace "ensure," "appropriate," "proper" with specific measurable language
- SHALL = requirement, SHOULD = recommendation, MAY = permission
- Use action verbs from PPA Attachment 1 (ADJUST, ALIGN, CHECK, CLOSE, CONNECT, DE-ENERGIZE, DRAIN, ENERGIZE, FILL, FLUSH, INSTALL, ISOLATE, MARK, MEASURE, MONITOR, NOTIFY, OBSERVE, OPEN, OPERATE, PLACE, POSITION, PRESS, PULL, PUSH, RAISE, RECORD, RELEASE, REMOVE, REPLACE, RESET, RESTORE, ROTATE, SELECT, SET, START, STOP, TAG, TEST, TORQUE, TURN, VERIFY, etc.)

ABBREVIATIONS:
- Spell out on first reference in each major section, followed by acronym in parentheses
- Example: Competency Assurance Standard (CAS)
- Plurals: add lowercase s (BUs not BU's)
- Use standard industry abbreviations: JSA, PSV, FPSO, P&ID, PFD, HAZOP

FORMAT:
- Font: Arial 11 or 12 point
- Paper: Portrait, 8.5 x 11
- Margins: 0.8 inch left/right, 0.5 inch top/bottom
- Left justify all text
- Single line spacing with white space between steps
- Page header on every page: procedure title, number, revision, page X of Y
- Keep steps unbroken on same page
- Use continuation headings: "4.2 (cont.)" when content spans pages

TASK ANALYSIS (when building from scratch):
- Gather P&IDs, PFDs, vendor info, system descriptions
- Identify operating system using engineering system numbering
- List major equipment from P&IDs
- Consolidate common equipment (no duplicates)
- Populate tasks per equipment, then device actions within each task

=== CRITICAL TASK EXECUTION (CTE Playbook Rev. 9) ===

SAFEGUARD CRITICAL TASKS are human actions that:
- Directly impact the process and are themselves a Critical Safeguard, OR
- Directly impact a critical safeguard, and if performed incorrectly, can trigger a Highest Consequence scenario

For all Safeguard Critical Tasks:
1. Pre-requisites must include checking the health of other safeguards
2. Safeguard Critical Steps must be clearly marked in the written documentation
3. Consequence of failure must be included in the preamble AND in a box before the critical step
4. A hold point and verifier signature line must be included (determined by SCTA)

INDEPENDENT VERIFICATION METHODS:
1 = In-person by direct visual inspection (not present during step execution)
2 = In-person by direct visual inspection (present during step execution)
3 = Remote via radio/phone with positive indication (console alarm, flow/pressure readback)
4 = Remote via electronic means (camera, barcode scan, RFID/QR code)

Key attributes:
- Location and/or conditions to be verified must be clearly defined, observable, and written
- Performed by a qualified individual
- Pre-step: both individuals must independently conclude it is safe to proceed
- Post-step: both individuals must independently confirm outcome conditions are met
- Verification method must be clearly described in the written step
- Shall rely on active participation (not passive)

AFTER ACTION REVIEW (AAR):
- For infrequently performed SGC tasks: AAR each time
- For frequently performed SGC tasks: AAR at set frequency
- AAR occurs as soon as practical after task completion, no later than end of work shift
- Questions: Did it go as expected? Did you do something differently? What could be improved?

=== OIMS 6.1 COMPLIANCE ===
Procedures must address:
- Level of Operations Integrity risk (determines detail and verification needed)
- Operating envelopes with consequence of deviation and response
- Transient operations as applicable
- Regulatory requirements
- Human Factors including capabilities and limitations
- Simplifying processes or tasks to reduce potential for error

{style_config}

{template_config}

If the user provides an existing procedure to update, review it against these standards, identify deficiencies, and ask targeted questions about the changes needed.

Always be thorough — a missing step or unclear instruction in a procedure can lead to safety incidents. Every safeguard critical step must be clearly identified with proper warnings and verification requirements."""


@router.get("/procedures/presets")
async def get_facility_presets():
    """Return facility presets with craft options."""
    return {"facilities": FACILITY_PRESETS}


@router.get("/procedures")
async def list_procedures(
    request: Request,
    skip: int = Query(0, ge=0),
    limit: int = Query(20, ge=1, le=100),
    facility: str = Query(None),
    category: str = Query(None),
    status: str = Query(None),
    db=Depends(get_db),
):
    user_id = getattr(request.state, "user_id", None)
    query = db.query(DBProcedureSession)
    if user_id:
        query = query.filter(DBProcedureSession.user_id == user_id)
    if facility:
        query = query.filter(DBProcedureSession.facility == facility)
    if category:
        query = query.filter(DBProcedureSession.category == category)
    if status:
        query = query.filter(DBProcedureSession.status == status)
    total = query.count()
    sessions = query.order_by(DBProcedureSession.updated_at.desc()).offset(skip).limit(limit).all()
    return {
        "sessions": [
            {
                "id": s.id,
                "title": _safe_decrypt(s.title) or "Untitled Procedure",
                "status": s.status,
                "facility": s.facility,
                "category": s.category,
                "created_at": s.created_at.isoformat(),
                "updated_at": s.updated_at.isoformat(),
            }
            for s in sessions
        ],
        "total": total,
    }


@router.post("/procedures")
async def create_procedure(request: Request, db=Depends(get_db)):
    body = await request.json()
    title = body.get("title", "").strip() or "New Procedure"
    source_doc_id = body.get("source_doc_id")
    mode = body.get("mode", "new")
    facility = body.get("facility", "").strip() or None
    craft = body.get("craft", "").strip() or None
    category = body.get("category", "").strip() or None

    if source_doc_id:
        doc = db.query(DBDocument).filter(DBDocument.id == source_doc_id).first()
        if not doc:
            raise HTTPException(status_code=404, detail="Source document not found")

    now = datetime.now(timezone.utc)
    session = DBProcedureSession(
        id=str(uuid.uuid4()),
        user_id=getattr(request.state, "user_id", None),
        title=_encrypt_text(title),
        source_doc_id=source_doc_id,
        facility=facility,
        category=category,
        status="gathering",
        created_at=now,
        updated_at=now,
    )
    db.add(session)
    db.commit()

    if mode == "update" and source_doc_id:
        initial_msg = (
            "I've loaded the existing procedure for review. I'll analyze it against "
            "PPA AP-907-005, OIMS 6.1, and CTE Playbook standards, then guide you through "
            "the updates needed. Let me start by reviewing what's there.\n\n"
            "What changes or updates are you looking to make to this procedure?"
        )
    elif source_doc_id:
        initial_msg = (
            "I can see you've attached a reference document. I'll use it as a starting "
            "point. Let me begin gathering the information needed for the procedure.\n\n"
            "What is the procedure title and designation number?"
        )
    else:
        if facility and craft:
            initial_msg = (
                f"I'm ready to help you write a new procedure for **{facility}** ({craft}). "
                "I'll guide you through each section to build a complete document "
                "that meets PPA AP-907-005 and OIMS 6.1 standards.\n\n"
                "What is the procedure title and designation number?"
            )
        elif facility:
            initial_msg = (
                f"I'm ready to help you write a new procedure for **{facility}**. "
                "I'll guide you through each section.\n\n"
                "What craft/discipline is this for and what is the procedure title?"
            )
        else:
            initial_msg = (
                "I'm ready to help you write a new procedure. I'll guide you through "
                "each section, asking the right questions to build a complete document "
                "that meets PPA AP-907-005 and OIMS 6.1 standards.\n\n"
                "Let's start — what facility and craft/discipline is this for? "
                "(e.g., Facility: Hebron, Craft: Operations)"
            )

    assistant_msg = DBProcedureMessage(
        id=str(uuid.uuid4()),
        session_id=session.id,
        role="assistant",
        content=_encrypt_text(initial_msg),
        created_at=now,
    )
    db.add(assistant_msg)
    db.commit()

    return {
        "id": session.id,
        "title": title,
        "status": session.status,
        "mode": mode,
        "messages": [{"role": "assistant", "content": initial_msg}],
    }


@router.post("/procedures/upload")
async def upload_procedure_doc(
    request: Request,
    file: UploadFile = File(...),
    db=Depends(get_db),
):
    """Upload a .docx procedure file and create a session for updating it."""
    if not file.filename or not file.filename.lower().endswith(".docx"):
        raise HTTPException(status_code=400, detail="Only .docx files are supported")

    file_bytes = await file.read()
    if len(file_bytes) > 20 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="File too large (max 20MB)")

    try:
        extracted = _extract_docx_structured(file_bytes)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to parse .docx: {str(e)}")

    title_match = None
    for line in extracted.split("\n"):
        if line.startswith("# "):
            title_match = line[2:].strip()
            break
    proc_title = title_match or file.filename.rsplit(".", 1)[0]

    now = datetime.now(timezone.utc)
    session = DBProcedureSession(
        id=str(uuid.uuid4()),
        user_id=getattr(request.state, "user_id", None),
        title=_encrypt_text(f"Update: {proc_title}"),
        status="gathering",
        gathered_data=_encrypt_text(extracted),
        created_at=now,
        updated_at=now,
    )
    db.add(session)
    db.commit()

    initial_msg = (
        f"I've loaded **{proc_title}** for review. I'll analyze it against "
        "PPA AP-907-005, OIMS 6.1, and CTE Playbook standards.\n\n"
        "What changes or updates are you looking to make to this procedure?"
    )
    assistant_msg = DBProcedureMessage(
        id=str(uuid.uuid4()),
        session_id=session.id,
        role="assistant",
        content=_encrypt_text(initial_msg),
        created_at=now,
    )
    db.add(assistant_msg)
    db.commit()

    return {
        "id": session.id,
        "title": f"Update: {proc_title}",
        "status": session.status,
        "mode": "update",
        "messages": [{"role": "assistant", "content": initial_msg}],
    }


@router.post("/procedures/logo")
async def upload_logo(
    request: Request,
    file: UploadFile = File(...),
    db=Depends(get_db),
):
    """Upload a company logo for procedure headers."""
    import base64

    if not file.filename:
        raise HTTPException(status_code=400, detail="No file provided")
    lower = file.filename.lower()
    if not any(lower.endswith(ext) for ext in (".png", ".jpg", ".jpeg")):
        raise HTTPException(status_code=400, detail="Only PNG/JPG images are supported")

    file_bytes = await file.read()
    if len(file_bytes) > 2 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="Logo too large (max 2MB)")

    encoded = base64.b64encode(file_bytes).decode("ascii")
    ext = "png" if lower.endswith(".png") else "jpeg"
    logo_data = f"data:image/{ext};base64,{encoded}"

    now = datetime.now(timezone.utc)
    existing = db.query(DBAppSetting).filter(DBAppSetting.key == "company_logo").first()
    if existing:
        existing.value = _encrypt_text(logo_data)
        existing.updated_at = now
    else:
        setting = DBAppSetting(key="company_logo", value=_encrypt_text(logo_data), updated_at=now)
        db.add(setting)
    db.commit()

    return {"detail": "Logo uploaded successfully"}


@router.get("/procedures/logo")
async def get_logo(db=Depends(get_db)):
    """Check if a company logo is uploaded."""
    setting = db.query(DBAppSetting).filter(DBAppSetting.key == "company_logo").first()
    return {"has_logo": setting is not None}


@router.get("/procedures/{session_id}")
async def get_procedure(session_id: str, db=Depends(get_db)):
    session = db.query(DBProcedureSession).filter(DBProcedureSession.id == session_id).first()
    if not session:
        raise HTTPException(status_code=404, detail="Procedure session not found")

    messages = [
        {"role": m.role, "content": _safe_decrypt(m.content) or ""}
        for m in session.messages
    ]

    return {
        "id": session.id,
        "title": _safe_decrypt(session.title) or "Untitled Procedure",
        "status": session.status,
        "facility": session.facility,
        "category": session.category,
        "source_doc_id": session.source_doc_id,
        "messages": messages,
        "created_at": session.created_at.isoformat(),
        "updated_at": session.updated_at.isoformat(),
    }


@router.delete("/procedures/{session_id}")
async def delete_procedure(session_id: str, db=Depends(get_db)):
    session = db.query(DBProcedureSession).filter(DBProcedureSession.id == session_id).first()
    if not session:
        raise HTTPException(status_code=404, detail="Procedure session not found")
    db.delete(session)
    db.commit()
    return {"detail": "Procedure session deleted"}


# --- 7. Status workflow ---
VALID_TRANSITIONS = {
    "gathering": ["drafting", "complete"],
    "drafting": ["review", "gathering", "complete"],
    "review": ["approved", "drafting"],
    "approved": ["drafting", "complete"],
    "complete": ["review", "drafting"],
}


@router.post("/procedures/{session_id}/status")
async def update_procedure_status(session_id: str, request: Request, db=Depends(get_db)):
    session = db.query(DBProcedureSession).filter(DBProcedureSession.id == session_id).first()
    if not session:
        raise HTTPException(status_code=404, detail="Procedure session not found")
    body = await request.json()
    new_status = body.get("status", "").strip()
    allowed = VALID_TRANSITIONS.get(session.status, [])
    if new_status not in allowed:
        raise HTTPException(
            status_code=400,
            detail=f"Cannot transition from '{session.status}' to '{new_status}'. Allowed: {allowed}",
        )
    session.status = new_status
    session.updated_at = datetime.now(timezone.utc)
    db.commit()
    return {"id": session.id, "status": session.status}


# --- 3. Chat file attachments ---
@router.post("/procedures/{session_id}/attach")
async def attach_file_to_procedure(
    session_id: str,
    request: Request,
    file: UploadFile = File(...),
    db=Depends(get_db),
):
    """Attach a reference file (PDF, DOCX, TXT) to the procedure chat."""
    session = db.query(DBProcedureSession).filter(DBProcedureSession.id == session_id).first()
    if not session:
        raise HTTPException(status_code=404, detail="Procedure session not found")

    file_bytes = await file.read()
    if len(file_bytes) > 20 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="File too large (max 20MB)")

    filename = file.filename or "attachment"
    lower = filename.lower()
    extracted = ""

    if lower.endswith(".docx"):
        try:
            extracted = _extract_docx_structured(file_bytes)
        except Exception:
            raise HTTPException(status_code=400, detail="Failed to parse .docx file")
    elif lower.endswith(".pdf"):
        try:
            import pymupdf
            import io as _io
            pdf = pymupdf.open(stream=file_bytes, filetype="pdf")
            pages = []
            for i in range(len(pdf)):
                pages.append(f"--- Page {i+1} ---\n{pdf[i].get_text()}")
            extracted = "\n".join(pages)
        except Exception:
            raise HTTPException(status_code=400, detail="Failed to parse PDF file")
    elif lower.endswith(".txt") or lower.endswith(".csv"):
        try:
            extracted = file_bytes.decode("utf-8", errors="replace")
        except Exception:
            raise HTTPException(status_code=400, detail="Failed to read text file")
    else:
        raise HTTPException(status_code=400, detail="Supported formats: .docx, .pdf, .txt, .csv")

    truncated = extracted[:10000]
    now = datetime.now(timezone.utc)
    attach_msg = DBProcedureMessage(
        id=str(uuid.uuid4()),
        session_id=session.id,
        role="user",
        content=_encrypt_text(f"[Attached file: {filename}]\n\n{truncated}"),
        created_at=now,
    )
    db.add(attach_msg)
    session.updated_at = now
    db.commit()

    return {
        "detail": f"File '{filename}' attached ({len(extracted)} chars extracted)",
        "filename": filename,
        "chars": len(extracted),
    }


# --- 4. Procedure compliance report ---
@router.post("/procedures/{session_id}/compliance")
async def compliance_report(session_id: str, db=Depends(get_db)):
    """Analyze procedure against PPA AP-907-005, OIMS 6.1, CTE Playbook standards."""
    session = db.query(DBProcedureSession).filter(DBProcedureSession.id == session_id).first()
    if not session:
        raise HTTPException(status_code=404, detail="Procedure session not found")

    content = ""
    if session.output_content:
        content = _safe_decrypt(session.output_content) or ""
    elif session.gathered_data:
        content = _safe_decrypt(session.gathered_data) or ""
    if not content:
        raise HTTPException(status_code=400, detail="No procedure content to analyze")

    client = Anthropic()
    analysis_prompt = f"""Analyze this procedure against the following standards and provide a structured compliance report:

1. PPA AP-907-005 Procedure Writer's Manual
2. OIMS Element 6.1 (Operations Integrity)
3. Critical Task Execution Playbook (CTE Rev. 9)
4. Upstream Procedure Writing Rules

PROCEDURE CONTENT:
{content[:10000]}

Provide your analysis in this exact format:

## Overall Score: X/10

## Findings

For each finding, use this format:
- [PASS] or [FAIL] or [WARN] Category: Description

Check these categories:
1. STRUCTURE: Required sections present (Purpose, Scope, Precautions, Prerequisites, Instructions, References)
2. ACTION STEPS: Action verbs uppercase bold, one action per step, active voice
3. EMPHASIS: IF/WHEN/THEN uppercase underlined bold, component positions uppercase
4. NOTES/CAUTIONS/WARNINGS: Placed before steps, correct sequence, passive voice
5. PREREQUISITES: PPE, materials, tools, other prerequisites defined
6. SAFEGUARD CRITICAL: SGC steps identified, hold points defined, IV method specified
7. VOCABULARY: No ambiguous words (ensure, appropriate, proper), SHALL/SHOULD/MAY used correctly
8. STEP NUMBERING: Up to 4 levels, consistent format
9. ABBREVIATIONS: Spelled out on first use
10. HUMAN FACTORS: Steps designed for human capability, error reduction considered

## Recommendations
Numbered list of specific improvements."""

    response = client.messages.create(
        model=CHAT_MODEL,
        max_tokens=4096,
        messages=[{"role": "user", "content": analysis_prompt}],
    )

    report_text = response.content[0].text if response.content else "Analysis failed"
    return {"report": report_text}


# --- 5. Browser preview ---
@router.get("/procedures/{session_id}/preview")
async def preview_procedure(session_id: str, db=Depends(get_db)):
    """Return HTML preview of the procedure content."""
    from fastapi.responses import HTMLResponse
    import re

    session = db.query(DBProcedureSession).filter(DBProcedureSession.id == session_id).first()
    if not session:
        raise HTTPException(status_code=404, detail="Procedure session not found")
    if not session.output_content:
        raise HTTPException(status_code=400, detail="Procedure not yet generated")

    content = _safe_decrypt(session.output_content) or ""
    title = _safe_decrypt(session.title) or "Procedure"

    def esc(t):
        return t.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    html_parts = [
        "<!DOCTYPE html><html><head>",
        f"<title>{esc(title)}</title>",
        "<style>",
        "body{font-family:Arial,sans-serif;font-size:11pt;max-width:8.5in;margin:auto;padding:1in 0.7in;color:#000}",
        "h1{font-size:14pt;border-bottom:2px solid #000;padding-bottom:4px}",
        "h2{font-size:12pt}h3{font-size:11pt}",
        "table{border-collapse:collapse;width:100%;margin:8px 0}",
        "th,td{border:1px solid #000;padding:4px 8px;text-align:left;font-size:11pt}",
        "th{background:#000;color:#fff;font-weight:bold}",
        ".note{background:#DEEAF6;padding:6px 10px;margin:4px 0}",
        ".caution{background:#FFF2CC;padding:6px 10px;margin:4px 0}",
        ".caution-label{color:#BF8F00;font-weight:bold}",
        ".warning{background:#FFE0E0;padding:6px 10px;margin:4px 0;border:2px solid #CC0000}",
        ".warning-label{color:#CC0000;font-weight:bold}",
        ".hold-point{background:#000;color:#fff;padding:8px 10px;margin:4px 0;font-weight:bold}",
        ".sgc-warning{background:#FFE0E0;border:2px solid #CC0000;padding:8px 10px;margin:4px 0}",
        ".sgc-label{color:#CC0000;font-weight:bold}",
        ".iv-table{background:#DEEAF6}",
        ".end-section{text-align:center;font-weight:bold;font-size:12pt;padding:8px}",
        "ul{list-style-type:disc;margin:4px 0 4px 20px}",
        "@media print{body{padding:0.5in}}",
        "</style></head><body>",
    ]

    in_table = False
    for line in content.split("\n"):
        line = line.rstrip()
        if not line:
            if in_table:
                html_parts.append("</table>")
                in_table = False
            continue

        stripped = line.strip()
        upper = stripped.upper()

        if "SAFEGUARD CRITICAL STEP" in upper:
            text = re.sub(r'^.*SAFEGUARD CRITICAL STEP[:\s]*', '', stripped, flags=re.IGNORECASE)
            html_parts.append(f'<div class="sgc-warning"><span class="sgc-label">WARNING - SAFEGUARD CRITICAL STEP</span> {esc(text)}</div>')
        elif upper.startswith("HOLD POINT") or upper.startswith("**HOLD POINT"):
            text = re.sub(r'^\*{0,2}\s*HOLD POINT\s*\*{0,2}[:\s]*', '', stripped, flags=re.IGNORECASE)
            html_parts.append(f'<div class="hold-point">HOLD POINT {esc(text)}</div>')
        elif "INDEPENDENT VERIFICATION" in upper and ("REQUIRED" in upper or "NEEDED" in upper):
            html_parts.append('<table class="iv-table"><tr><th colspan="3" style="background:#DEEAF6;color:#000">Independent Verification Required</th></tr><tr><td><b>Name:</b> ___________</td><td><b>Signature:</b> ___________</td><td><b>Date:</b> ___________</td></tr></table>')
        elif upper.startswith("WARNING:") or upper.startswith("! WARNING:"):
            text = re.sub(r'^!?\s*WARNING:\s*', '', stripped, flags=re.IGNORECASE)
            html_parts.append(f'<div class="warning"><span class="warning-label">WARNING:</span> {esc(text)}</div>')
        elif upper.startswith("CAUTION:"):
            text = re.sub(r'^CAUTION:\s*', '', stripped, flags=re.IGNORECASE)
            html_parts.append(f'<div class="caution"><span class="caution-label">&#9650; CAUTION:</span> {esc(text)}</div>')
        elif upper.startswith("NOTE:"):
            text = re.sub(r'^NOTE:\s*', '', stripped, flags=re.IGNORECASE)
            html_parts.append(f'<div class="note"><b>NOTE:</b> {esc(text)}</div>')
        elif line.startswith("# "):
            html_parts.append(f"<h1>{esc(line[2:])}</h1>")
        elif line.startswith("## "):
            html_parts.append(f"<h2>{esc(line[3:])}</h2>")
        elif line.startswith("### "):
            html_parts.append(f"<h3>{esc(line[4:])}</h3>")
        elif stripped.startswith("**") and stripped.endswith("**"):
            html_parts.append(f"<p><b>{esc(stripped.strip('*'))}</b></p>")
        elif stripped.startswith("- ") or stripped.startswith("• "):
            html_parts.append(f"<ul><li>{esc(stripped[2:])}</li></ul>")
        elif stripped.startswith("|") and stripped.endswith("|"):
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            if all(set(c) <= set("- :") for c in cells):
                continue
            if not in_table:
                in_table = True
                html_parts.append("<table><tr>" + "".join(f"<th>{esc(c)}</th>" for c in cells) + "</tr>")
            else:
                html_parts.append("<tr>" + "".join(f"<td>{esc(c)}</td>" for c in cells) + "</tr>")
        else:
            html_parts.append(f"<p>{esc(stripped)}</p>")

    if in_table:
        html_parts.append("</table>")
    html_parts.append("</body></html>")

    return HTMLResponse("\n".join(html_parts))


@router.post("/procedures/{session_id}/chat")
async def procedure_chat(session_id: str, request: Request, db=Depends(get_db)):
    session = db.query(DBProcedureSession).filter(DBProcedureSession.id == session_id).first()
    if not session:
        raise HTTPException(status_code=404, detail="Procedure session not found")

    body = await request.json()
    user_message = body.get("message", "").strip()
    if not user_message:
        raise HTTPException(status_code=400, detail="Message is required")

    now = datetime.now(timezone.utc)
    user_msg = DBProcedureMessage(
        id=str(uuid.uuid4()),
        session_id=session.id,
        role="user",
        content=_encrypt_text(user_message),
        created_at=now,
    )
    db.add(user_msg)
    session.updated_at = now
    db.commit()

    style_config = ""
    if session.style_config:
        style_config = f"Additional style rules:\n{_safe_decrypt(session.style_config) or ''}"

    template_config = ""
    if session.template_config:
        template_config = f"Template structure:\n{_safe_decrypt(session.template_config) or ''}"

    source_context = ""
    if session.gathered_data:
        gathered = _safe_decrypt(session.gathered_data) or ""
        if gathered:
            source_context = f"\n\n--- EXISTING PROCEDURE (to update/reference) ---\n{gathered[:12000]}\n--- END ---"
    elif session.source_doc_id:
        doc = db.query(DBDocument).filter(DBDocument.id == session.source_doc_id).first()
        if doc:
            pages = _load_stored_text(doc)
            if pages:
                text = _stored_text_to_structured(pages)
                source_context = f"\n\n--- EXISTING PROCEDURE (to update) ---\n{text[:12000]}\n--- END ---"

    system_prompt = PROCEDURE_SYSTEM_PROMPT.format(
        style_config=style_config,
        template_config=template_config,
    )
    if source_context:
        system_prompt += source_context

    history = []
    for m in session.messages:
        content = _safe_decrypt(m.content) or ""
        if not content:
            continue
        history.append({"role": m.role, "content": content})

    client = Anthropic()

    chat_max = min(CHAT_MAX_TOKENS, 4096)

    def generate():
        full_response = []
        with client.messages.stream(
            model=CHAT_MODEL,
            max_tokens=chat_max,
            system=[{"type": "text", "text": system_prompt, "cache_control": {"type": "ephemeral"}}],
            messages=history,
        ) as stream:
            for text in stream.text_stream:
                full_response.append(text)
                yield f"data: {json.dumps({'type': 'text', 'content': text})}\n\n"

        assistant_content = "".join(full_response)
        save_db = SessionLocal()
        try:
            assistant_msg = DBProcedureMessage(
                id=str(uuid.uuid4()),
                session_id=session.id,
                role="assistant",
                content=_encrypt_text(assistant_content),
                created_at=datetime.now(timezone.utc),
            )
            save_db.add(assistant_msg)
            save_db.commit()
        finally:
            save_db.close()

        yield f"data: {json.dumps({'type': 'done'})}\n\n"

    return StreamingResponse(generate(), media_type="text/event-stream")


@router.post("/procedures/{session_id}/generate")
async def generate_procedure(session_id: str, request: Request, db=Depends(get_db)):
    """Generate the final Word document from gathered information."""
    session = db.query(DBProcedureSession).filter(DBProcedureSession.id == session_id).first()
    if not session:
        raise HTTPException(status_code=404, detail="Procedure session not found")

    style_config = ""
    if session.style_config:
        style_config = f"Additional style rules:\n{_safe_decrypt(session.style_config) or ''}"

    template_config = ""
    if session.template_config:
        template_config = f"Template structure:\n{_safe_decrypt(session.template_config) or ''}"

    history = []
    for m in session.messages:
        content = _safe_decrypt(m.content) or ""
        if not content:
            continue
        history.append({"role": m.role, "content": content})

    is_update = bool(session.gathered_data or session.source_doc_id)
    alteration_section = ""
    if is_update:
        alteration_section = """
If this is an update to an existing procedure, include a "Summary of Alterations" section at the end listing what changed:
## Summary of Alterations
| Rev | Date | Description of Change | Author |
| --- | --- | --- | --- |
(List each change made in this revision)
"""

    generation_prompt = f"""Based on all the information gathered in this conversation, generate the complete procedure document now.

OUTPUT FORMAT RULES:
- Use # for main sections (e.g., # 1. PURPOSE AND SCOPE)
- Use ## for subsections (e.g., ## 4.1. Equipment Lineup)
- Use ### for sub-subsections
- Action verbs in UPPERCASE: OPEN, CLOSE, VERIFY, CHECK, RECORD, PERFORM, etc.
- Conditional terms in UPPERCASE: IF, THEN, WHEN, AND, OR, NOT, WHILE
- Component positions in UPPERCASE: OPEN, CLOSED, ON, OFF, AUTO
- Component names in Title Case: Amine Discharge Valve
- Warnings as: WARNING: [text]
- Cautions as: CAUTION: [text]
- Notes as: NOTE: [text]
- Place warnings/cautions/notes BEFORE the step they apply to
- Tables as markdown tables with | pipes |
- Action step tables with these exact 4 columns:
| Step | Action / Remarks | Who | Check |
| --- | --- | --- | --- |
| 1 | OPEN inlet valve XX-XXX-001 | Ops Tech | ☐ |

Use these EXACT section headings:
# 1. PURPOSE AND SCOPE
## 1.1. Purpose
## 1.2. Scope
# 2. PRECAUTIONS AND LIMITATIONS
## 2.1. Precautions
## 2.2. Limitations
# 3. PREREQUISITES
## 3.1. Personal Protective Equipment, PPE
## 3.2. Materials
## 3.3. Special Tools and Equipment
## 3.4. Other Prerequisites
# 4. INSTRUCTIONS
## 4.1. [First instruction section]
(each subsection gets its own step table)
# 5. REFERENCES AND COMMITMENT
## 5.1. Performance References
## 5.2. Commitments References
## 5.3. Developmental References
{alteration_section}
Every action step must start with an uppercase bold action verb and contain only one action.
Attachments start with ## Attachment N - [Title] and get a new page automatically."""

    history.append({"role": "user", "content": generation_prompt})

    client = Anthropic()

    gen_system = PROCEDURE_SYSTEM_PROMPT.format(
        style_config=style_config,
        template_config=template_config,
    )

    def generate():
        full_response = []
        with client.messages.stream(
            model=CHAT_MODEL,
            max_tokens=CHAT_MAX_TOKENS,
            system=[{"type": "text", "text": gen_system, "cache_control": {"type": "ephemeral"}}],
            messages=history,
        ) as stream:
            for text in stream.text_stream:
                full_response.append(text)
                yield f"data: {json.dumps({'type': 'text', 'content': text})}\n\n"

        content = "".join(full_response)
        save_db = SessionLocal()
        try:
            proc = save_db.query(DBProcedureSession).filter(DBProcedureSession.id == session_id).first()
            if proc:
                proc.output_content = _encrypt_text(content)
                proc.status = "complete"
                proc.updated_at = datetime.now(timezone.utc)
                save_db.commit()
        finally:
            save_db.close()

        yield f"data: {json.dumps({'type': 'done'})}\n\n"

    return StreamingResponse(generate(), media_type="text/event-stream")


@router.get("/procedures/{session_id}/download")
async def download_procedure(session_id: str, db=Depends(get_db)):
    """Download the generated procedure as a Word document matching Hebron template."""
    session = db.query(DBProcedureSession).filter(DBProcedureSession.id == session_id).first()
    if not session:
        raise HTTPException(status_code=404, detail="Procedure session not found")
    if not session.output_content:
        raise HTTPException(status_code=400, detail="Procedure not yet generated")

    from docx import Document
    from docx.shared import Pt, Emu, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from lxml import etree
    import io
    import re
    from datetime import date

    BLACK = RGBColor(0x00, 0x00, 0x00)
    WHITE = RGBColor(0xFF, 0xFF, 0xFF)
    AMBER = RGBColor(0xBF, 0x8F, 0x00)
    RED = RGBColor(0xCC, 0x00, 0x00)
    FONT = "Arial"

    content = _safe_decrypt(session.output_content) or ""
    title = _safe_decrypt(session.title) or "Procedure"

    proc_number = ""
    revision = ""
    facility = ""
    craft = ""
    reval_date = ""
    for m in session.messages:
        msg_text = _safe_decrypt(m.content) or ""
        if not proc_number:
            num_match = re.search(
                r'(?:procedure\s+(?:number|#|no\.?|designation)\s*[:\-]?\s*)([A-Z0-9][\w\-\.]+)',
                msg_text, re.IGNORECASE,
            )
            if num_match:
                proc_number = num_match.group(1)
        if not revision:
            rev_match = re.search(r'(?:revision|rev\.?)\s*[:\-]?\s*([A-Z]?\d+)', msg_text, re.IGNORECASE)
            if rev_match:
                revision = rev_match.group(1)
        if not facility:
            fac_match = re.search(r'(?:facility|site|platform)\s*[:\-]?\s*([A-Za-z][\w\s]+)', msg_text, re.IGNORECASE)
            if fac_match:
                facility = fac_match.group(1).strip()
        if not craft:
            craft_match = re.search(r'(?:craft|discipline)\s*[:\-]?\s*([A-Za-z][\w\s]+)', msg_text, re.IGNORECASE)
            if craft_match:
                craft = craft_match.group(1).strip()

    if session.facility and not facility:
        facility = session.facility

    doc_number_display = f"{proc_number} Rev. {revision}" if proc_number and revision else proc_number or ""

    logo_bytes = None
    logo_setting = db.query(DBAppSetting).filter(DBAppSetting.key == "company_logo").first()
    if logo_setting and logo_setting.value:
        import base64
        logo_data = _safe_decrypt(logo_setting.value) or ""
        if logo_data.startswith("data:image"):
            b64_part = logo_data.split(",", 1)[1] if "," in logo_data else ""
            if b64_part:
                logo_bytes = base64.b64decode(b64_part)

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)

    for hs, sz in [("Heading 1", 14), ("Heading 2", 12), ("Heading 3", 11)]:
        if hs in doc.styles:
            h_style = doc.styles[hs]
            h_style.font.name = FONT
            h_style.font.color.rgb = BLACK
            h_style.font.bold = True
            h_style.font.size = Pt(sz)

    section = doc.sections[0]
    section.page_width = Emu(7772400)
    section.page_height = Emu(10058400)
    section.top_margin = Emu(457200)
    section.bottom_margin = Emu(457200)
    section.left_margin = Emu(635000)
    section.right_margin = Emu(635000)
    section.different_first_page_header_footer = True

    def _make_field(parent, instr):
        fld = etree.SubElement(parent, qn("w:fldSimple"))
        fld.set(qn("w:instr"), instr)
        r = etree.SubElement(fld, qn("w:r"))
        t = etree.SubElement(r, qn("w:t"))
        t.text = "1"
        return fld

    def _set_cell_shading(cell, color):
        tc_pr = cell._element.get_or_add_tcPr()
        shd = tc_pr.makeelement(qn("w:shd"), {
            qn("w:fill"): color, qn("w:val"): "clear",
        })
        tc_pr.append(shd)

    def _styled_run(paragraph, text, font_size=11, bold=False, color=None, underline=False):
        r = paragraph.add_run(text)
        r.font.name = FONT
        r.font.size = Pt(font_size)
        r.bold = bold
        if color:
            r.font.color.rgb = color
        if underline:
            r.underline = True
        return r

    def _set_col_widths(tbl, widths_twips):
        for row in tbl.rows:
            for i, w in enumerate(widths_twips):
                if i < len(row.cells):
                    tc_pr = row.cells[i]._element.get_or_add_tcPr()
                    tcw = tc_pr.makeelement(qn("w:tcW"), {
                        qn("w:w"): str(w), qn("w:type"): "dxa",
                    })
                    tc_pr.append(tcw)

    def _build_header_table_xml(element_parent, cells_data):
        h_tbl = element_parent.makeelement(qn("w:tbl"), {})
        tbl_pr = h_tbl.makeelement(qn("w:tblPr"), {})
        tbl_w = tbl_pr.makeelement(qn("w:tblW"), {qn("w:w"): "5000", qn("w:type"): "pct"})
        tbl_pr.append(tbl_w)
        tbl_borders = tbl_pr.makeelement(qn("w:tblBorders"), {})
        for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
            b = tbl_borders.makeelement(qn(f"w:{edge}"), {
                qn("w:val"): "single", qn("w:sz"): "4",
                qn("w:color"): "000000", qn("w:space"): "0",
            })
            tbl_borders.append(b)
        tbl_pr.append(tbl_borders)
        h_tbl.append(tbl_pr)
        for row_data in cells_data:
            tr = h_tbl.makeelement(qn("w:tr"), {})
            for text, bold, size in row_data:
                tc = tr.makeelement(qn("w:tc"), {})
                p = tc.makeelement(qn("w:p"), {})
                if text:
                    r = p.makeelement(qn("w:r"), {})
                    rpr = r.makeelement(qn("w:rPr"), {})
                    rfont = rpr.makeelement(qn("w:rFonts"), {qn("w:ascii"): FONT, qn("w:hAnsi"): FONT})
                    rpr.append(rfont)
                    rsz = rpr.makeelement(qn("w:sz"), {qn("w:val"): str(size * 2)})
                    rpr.append(rsz)
                    if bold:
                        rpr.append(rpr.makeelement(qn("w:b"), {}))
                    r.append(rpr)
                    t = r.makeelement(qn("w:t"), {})
                    t.text = text
                    t.set(qn("xml:space"), "preserve")
                    r.append(t)
                    p.append(r)
                tc.append(p)
                tr.append(tc)
            h_tbl.append(tr)
        element_parent.append(h_tbl)

    # --- 1. PAGE 1 HEADER: 6-row title block table ---
    first_header = section.first_page_header
    first_header.is_linked_to_previous = False
    for p in first_header.paragraphs:
        p.clear()

    logo_label = "LOGO" if not logo_bytes else ""
    _build_header_table_xml(first_header._element, [
        [("Facility:", False, 9), ("", False, 9), (logo_label, False, 9)],
        [(facility or "—", True, 11), (title, True, 14), ("", False, 9)],
        [("Craft:", False, 9), (title, True, 14), ("", False, 9)],
        [(craft or "—", True, 11), ("", False, 9), ("", False, 9)],
        [("Revalidation Date:", False, 9), (doc_number_display, False, 10), ("", False, 9)],
        [(reval_date or "—", True, 11), (doc_number_display, False, 10), ("", False, 9)],
    ])

    if logo_bytes:
        try:
            logo_stream = io.BytesIO(logo_bytes)
            logo_para = first_header.paragraphs[-1] if first_header.paragraphs else first_header.add_paragraph()
            logo_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            logo_para.add_run().add_picture(logo_stream, width=Inches(1.0))
        except Exception:
            pass

    # --- 1. CONTINUATION HEADER (page 2+): compact 1-row table ---
    cont_header = section.header
    cont_header.is_linked_to_previous = False
    for p in cont_header.paragraphs:
        p.clear()

    _build_header_table_xml(cont_header._element, [
        [(doc_number_display, False, 10), (title, True, 11), (craft or "—", False, 10)],
    ])

    # --- FOOTER (all pages): centered "X of Y" page numbering ---
    for ftr in [section.footer, section.first_page_footer]:
        ftr.is_linked_to_previous = False
        fp = ftr.paragraphs[0] if ftr.paragraphs else ftr.add_paragraph()
        fp.text = ""
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        frun = fp.add_run()
        frun.font.name = FONT
        frun.font.size = Pt(8)
        _make_field(fp._element, " PAGE ")
        frun2 = fp.add_run(" of ")
        frun2.font.name = FONT
        frun2.font.size = Pt(8)
        _make_field(fp._element, " NUMPAGES ")

    # --- 2. TABLE OF CONTENTS ---
    toc_heading = doc.add_paragraph()
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    toc_run = toc_heading.add_run("TABLE OF CONTENTS")
    toc_run.bold = True
    toc_run.font.name = FONT
    toc_run.font.size = Pt(14)

    toc_para = doc.add_paragraph()
    fld_begin = etree.SubElement(toc_para._element, qn("w:r"))
    fld_char_begin = etree.SubElement(fld_begin, qn("w:fldChar"))
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_run = etree.SubElement(toc_para._element, qn("w:r"))
    instr_text = etree.SubElement(instr_run, qn("w:instrText"))
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = ' TOC \\o "1-2" \\h \\z \\u '
    fld_sep = etree.SubElement(toc_para._element, qn("w:r"))
    fld_char_sep = etree.SubElement(fld_sep, qn("w:fldChar"))
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    placeholder_run = etree.SubElement(toc_para._element, qn("w:r"))
    placeholder_text = etree.SubElement(placeholder_run, qn("w:t"))
    placeholder_text.text = "Right-click and select 'Update Field' to generate table of contents"
    fld_end = etree.SubElement(toc_para._element, qn("w:r"))
    fld_char_end = etree.SubElement(fld_end, qn("w:fldChar"))
    fld_char_end.set(qn("w:fldCharType"), "end")

    doc.add_paragraph("")

    # --- 8. Step auto-numbering counter ---
    step_counter = [0]

    # --- Helper: NOTE row inside a step table (merged, light blue) ---
    def _add_note_row(tbl, text, num_cols):
        row = tbl.add_row()
        first_cell = row.cells[0]
        if num_cols > 1:
            first_cell.merge(row.cells[num_cols - 1])
        _set_cell_shading(first_cell, "DEEAF6")
        p = first_cell.paragraphs[0]
        p.text = ""
        _styled_run(p, "NOTE: ", font_size=11, bold=True)
        _styled_run(p, text, font_size=11)

    def _add_wingdings_triangle(paragraph):
        r = paragraph.add_run("p")
        r.font.name = "Wingdings 3"
        r.font.size = Pt(11)
        r.font.color.rgb = AMBER
        paragraph.add_run(" ")

    # --- Helper: CAUTION row inside a step table (merged, light yellow) ---
    def _add_caution_row(tbl, text, num_cols):
        row = tbl.add_row()
        first_cell = row.cells[0]
        if num_cols > 1:
            first_cell.merge(row.cells[num_cols - 1])
        _set_cell_shading(first_cell, "FFF2CC")
        p = first_cell.paragraphs[0]
        p.text = ""
        _add_wingdings_triangle(p)
        _styled_run(p, "CAUTION: ", font_size=11, bold=True, color=AMBER)
        _styled_run(p, text, font_size=11)

    # --- Helper: "End of Section" row ---
    def _add_end_of_section(tbl, num_cols):
        row = tbl.add_row()
        first_cell = row.cells[0]
        if num_cols > 1:
            first_cell.merge(row.cells[num_cols - 1])
        p = first_cell.paragraphs[0]
        p.text = ""
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _styled_run(p, "End of Section", font_size=12, bold=True)

    # --- Standalone NOTE/CAUTION/WARNING boxes (outside tables) ---
    def _add_note_box(doc, text):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.style = "Table Grid"
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        cell = tbl.cell(0, 0)
        _set_cell_shading(cell, "DEEAF6")
        p = cell.paragraphs[0]
        _styled_run(p, "NOTE: ", bold=True)
        _styled_run(p, text)

    def _add_caution_box(doc, text):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.style = "Table Grid"
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        cell = tbl.cell(0, 0)
        _set_cell_shading(cell, "FFF2CC")
        p = cell.paragraphs[0]
        _add_wingdings_triangle(p)
        _styled_run(p, "CAUTION: ", bold=True, color=AMBER)
        _styled_run(p, text)

    def _add_warning_box(doc, text):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.style = "Table Grid"
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        cell = tbl.cell(0, 0)
        _set_cell_shading(cell, "FFE0E0")
        p = cell.paragraphs[0]
        _styled_run(p, "WARNING: ", bold=True, color=RED)
        _styled_run(p, text)

    # --- HOLD POINT box ---
    def _add_hold_point(doc, text):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.style = "Table Grid"
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        cell = tbl.cell(0, 0)
        _set_cell_shading(cell, "000000")
        p = cell.paragraphs[0]
        _styled_run(p, "HOLD POINT", font_size=11, bold=True, color=WHITE)
        if text:
            p2 = cell.add_paragraph()
            _styled_run(p2, text, color=WHITE)

    # --- SAFEGUARD CRITICAL STEP warning ---
    def _add_safeguard_warning(doc, text):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.style = "Table Grid"
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        cell = tbl.cell(0, 0)
        _set_cell_shading(cell, "FFE0E0")
        tc_pr = cell._element.get_or_add_tcPr()
        borders = tc_pr.makeelement(qn("w:tcBorders"), {})
        for edge in ["top", "left", "bottom", "right"]:
            b = borders.makeelement(qn(f"w:{edge}"), {
                qn("w:val"): "single", qn("w:sz"): "12",
                qn("w:color"): "CC0000", qn("w:space"): "0",
            })
            borders.append(b)
        tc_pr.append(borders)
        p = cell.paragraphs[0]
        _styled_run(p, "WARNING - SAFEGUARD CRITICAL STEP", font_size=11, bold=True, color=RED)
        if text:
            p2 = cell.add_paragraph()
            _styled_run(p2, text, color=RED)

    # --- INDEPENDENT VERIFICATION fields ---
    def _add_iv_fields(doc):
        tbl = doc.add_table(rows=3, cols=3)
        tbl.style = "Table Grid"
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        header_cell = tbl.cell(0, 0)
        header_cell.merge(tbl.cell(0, 2))
        _set_cell_shading(header_cell, "DEEAF6")
        hp = header_cell.paragraphs[0]
        _styled_run(hp, "Independent Verification Required", font_size=11, bold=True)
        for col_idx, label in enumerate(["Name:", "Signature:", "Date:"]):
            lc = tbl.cell(1, col_idx)
            p = lc.paragraphs[0]
            _styled_run(p, label, font_size=10, bold=True)
        for col_idx in range(3):
            vc = tbl.cell(2, col_idx)
            p = vc.paragraphs[0]
            _styled_run(p, "________________________", font_size=10)

    # --- ACTION STEP TABLE detection and building ---
    table_active = False
    table_ref = None
    is_action_table = False
    table_col_count = 0

    def _detect_action_table(header_cells):
        lower = [c.lower().strip() for c in header_cells]
        return ("action" in lower or "action/remarks" in lower or "action / remarks" in lower) and (
            "who" in lower or "check" in lower or "no." in lower or "no" in lower or "step" in lower
        )

    def _finish_table():
        nonlocal table_active, table_ref, is_action_table, table_col_count
        if table_active and table_ref and is_action_table:
            _add_end_of_section(table_ref, table_col_count)
        table_active = False
        table_ref = None
        is_action_table = False
        table_col_count = 0

    # Column widths in twips (1 inch = 1440 twips, page ~6.5" usable)
    # Step: 0.6", Action/Remarks: 4.1", Who: 1.1", Check: 0.7"
    ACTION_COL_WIDTHS = [864, 5904, 1584, 1008]

    def _build_action_table_header(doc, cells):
        nonlocal table_ref, table_active, is_action_table, table_col_count
        step_counter[0] = 0
        has_check = any("check" in c.lower() for c in cells)
        cols = list(cells)
        if not has_check:
            cols.append("Check")
        tbl = doc.add_table(rows=1, cols=len(cols))
        tbl.style = "Table Grid"
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        for i, val in enumerate(cols):
            cell = tbl.cell(0, i)
            _set_cell_shading(cell, "000000")
            cell.text = ""
            p = cell.paragraphs[0]
            _styled_run(p, val, font_size=11, bold=True, color=WHITE)
        if len(cols) == 4:
            _set_col_widths(tbl, ACTION_COL_WIDTHS)
        table_ref = tbl
        table_active = True
        is_action_table = True
        table_col_count = len(cols)

    def _build_action_table_row(cells):
        nonlocal table_ref
        has_check_col = len(table_ref.columns) > len(cells)
        row_cells = list(cells)
        if has_check_col:
            row_cells.append("☐")
        step_counter[0] += 1
        row = table_ref.add_row()
        for i, val in enumerate(row_cells):
            if i < len(row.cells):
                row.cells[i].text = ""
                p = row.cells[i].paragraphs[0]
                step_col_lower = val.strip().lower()
                if val.strip() in ("☐", "[ ]", "[x]", "[]"):
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    _styled_run(p, "☐", font_size=12)
                elif i == 0 and (not val.strip() or step_col_lower in ("", "-")):
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    _styled_run(p, str(step_counter[0]), font_size=11)
                elif i == 0 and val.strip().isdigit():
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    step_counter[0] = int(val.strip())
                    _styled_run(p, val.strip(), font_size=11)
                else:
                    _add_formatted_text(p, val)

    def _build_generic_table_header(doc, cells):
        nonlocal table_ref, table_active, is_action_table, table_col_count
        tbl = doc.add_table(rows=1, cols=len(cells))
        tbl.style = "Table Grid"
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        for i, val in enumerate(cells):
            cell = tbl.cell(0, i)
            _set_cell_shading(cell, "000000")
            cell.text = ""
            p = cell.paragraphs[0]
            _styled_run(p, val, font_size=11, bold=True, color=WHITE)
        table_ref = tbl
        table_active = True
        is_action_table = False
        table_col_count = len(cells)

    def _build_generic_table_row(cells):
        nonlocal table_ref
        row = table_ref.add_row()
        for i, val in enumerate(cells):
            if i < len(row.cells):
                row.cells[i].text = ""
                p = row.cells[i].paragraphs[0]
                _styled_run(p, val, font_size=11)

    # --- Text formatting: bold+underline IF/WHEN/THEN/NOT, bold action verbs ---
    ACTION_VERBS = {
        "CONFIRM", "NOTIFY", "PERFORM", "RECORD", "CREATE", "REQUEST",
        "LOCATE", "PLACE", "VERIFY", "ENSURE", "OPEN", "CLOSE", "START",
        "STOP", "CHECK", "ADJUST", "MONITOR", "INSTALL", "REMOVE",
        "INSPECT", "ISOLATE", "ALIGN", "ACTIVATE", "DEACTIVATE",
        "POSITION", "TRANSFER", "OBSERVE", "CONNECT", "DISCONNECT",
        "OPERATE", "ENGAGE", "RELEASE", "INITIATE", "COMPLETE",
        "SECURE", "RESET", "RESTORE", "SET", "APPLY", "PRESSURIZE",
        "DEPRESSURIZE", "DRAIN", "FILL", "FLUSH", "PURGE",
    }
    CONDITIONAL_KEYWORDS = {"IF", "WHEN", "THEN", "NOT"}

    def _add_formatted_text(paragraph, text):
        words = text.split()
        i = 0
        buf = []
        while i < len(words):
            w = words[i]
            w_upper = w.rstrip(",:;.").upper()
            if w_upper in CONDITIONAL_KEYWORDS:
                if buf:
                    _styled_run(paragraph, " ".join(buf) + " ", font_size=11)
                    buf = []
                _styled_run(paragraph, w + " ", font_size=11, bold=True, underline=True)
            elif w_upper in ACTION_VERBS:
                if buf:
                    _styled_run(paragraph, " ".join(buf) + " ", font_size=11)
                    buf = []
                _styled_run(paragraph, w + " ", font_size=11, bold=True)
            else:
                buf.append(w)
            i += 1
        if buf:
            _styled_run(paragraph, " ".join(buf), font_size=11)

    # --- 5. Standard section number mapping for heading normalization ---
    STANDARD_SECTIONS = {
        "PURPOSE AND SCOPE": "1",
        "PURPOSE": "1.1",
        "SCOPE": "1.2",
        "PRECAUTIONS AND LIMITATIONS": "2",
        "PRECAUTIONS": "2.1",
        "LIMITATIONS": "2.2",
        "PREREQUISITES": "3",
        "PERSONAL PROTECTIVE EQUIPMENT": "3.1",
        "PPE": "3.1",
        "MATERIALS": "3.2",
        "SPECIAL TOOLS AND EQUIPMENT": "3.3",
        "OTHER PREREQUISITES": "3.4",
        "INSTRUCTIONS": "4",
        "REFERENCES AND COMMITMENT": "5",
        "REFERENCES AND COMMITMENTS": "5",
        "PERFORMANCE REFERENCES": "5.1",
        "COMMITMENTS REFERENCES": "5.2",
        "DEVELOPMENTAL REFERENCES": "5.3",
    }

    def _normalize_heading(text):
        clean = re.sub(r'^[\d.]+\s*', '', text).strip()
        upper = clean.upper()
        if upper in STANDARD_SECTIONS:
            return f"{STANDARD_SECTIONS[upper]}. {clean}"
        return text

    # --- MAIN CONTENT LOOP ---
    for line in content.split("\n"):
        line = line.rstrip()
        if not line:
            if table_active:
                _finish_table()
            doc.add_paragraph("")
            continue

        line_stripped = line.strip()
        line_upper = line_stripped.upper()

        # Safeguard critical step
        if "SAFEGUARD CRITICAL STEP" in line_upper:
            if table_active:
                _finish_table()
            text = re.sub(r'^.*SAFEGUARD CRITICAL STEP[:\s]*', '', line_stripped, flags=re.IGNORECASE)
            _add_safeguard_warning(doc, text)
            continue

        # Hold point
        if line_upper.startswith("HOLD POINT") or line_upper.startswith("**HOLD POINT"):
            if table_active:
                _finish_table()
            text = re.sub(r'^\*{0,2}\s*HOLD POINT\s*\*{0,2}[:\s]*', '', line_stripped, flags=re.IGNORECASE)
            _add_hold_point(doc, text)
            continue

        # Independent verification
        if "INDEPENDENT VERIFICATION" in line_upper and ("REQUIRED" in line_upper or "NEEDED" in line_upper):
            if table_active:
                _finish_table()
            _add_iv_fields(doc)
            continue

        # NOTE inside active table -> merged blue row
        if line_upper.startswith("NOTE:") and table_active and table_ref:
            text = re.sub(r'^NOTE:\s*', '', line_stripped, flags=re.IGNORECASE)
            _add_note_row(table_ref, text, table_col_count)
            continue

        # CAUTION inside active table -> merged yellow row
        if line_upper.startswith("CAUTION:") and table_active and table_ref:
            text = re.sub(r'^CAUTION:\s*', '', line_stripped, flags=re.IGNORECASE)
            _add_caution_row(table_ref, text, table_col_count)
            continue

        # Warning/caution/note outside tables
        if line_upper.startswith("! WARNING:") or line_upper.startswith("WARNING:"):
            if table_active:
                _finish_table()
            text = re.sub(r'^!?\s*WARNING:\s*', '', line_stripped, flags=re.IGNORECASE)
            _add_warning_box(doc, text)
        elif line_upper.startswith("CAUTION:"):
            if table_active:
                _finish_table()
            text = re.sub(r'^CAUTION:\s*', '', line_stripped, flags=re.IGNORECASE)
            _add_caution_box(doc, text)
        elif line_upper.startswith("NOTE:"):
            if table_active:
                _finish_table()
            text = re.sub(r'^NOTE:\s*', '', line_stripped, flags=re.IGNORECASE)
            _add_note_box(doc, text)

        # 7. Attachment headings — page break and centered title
        elif re.match(r'^#{1,2}\s*Attachment\s+\d+', line_stripped, re.IGNORECASE) or \
                (line_stripped.startswith("**") and re.match(r'\*{2}\s*Attachment\s+\d+', line_stripped, re.IGNORECASE)):
            if table_active:
                _finish_table()
            doc.add_page_break()
            att_text = re.sub(r'^#{1,3}\s*', '', line_stripped).strip().strip("*")
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(att_text)
            run.bold = True
            run.font.name = FONT
            run.font.size = Pt(14)

        # Headings (with standard section numbering)
        elif line.startswith("# "):
            if table_active:
                _finish_table()
            doc.add_heading(_normalize_heading(line[2:]), level=1)
        elif line.startswith("## "):
            if table_active:
                _finish_table()
            doc.add_heading(_normalize_heading(line[3:]), level=2)
        elif line.startswith("### "):
            if table_active:
                _finish_table()
            doc.add_heading(line[4:], level=3)

        # Bold section labels
        elif line_stripped.startswith("**") and line_stripped.endswith("**"):
            if table_active:
                _finish_table()
            p = doc.add_paragraph()
            run = p.add_run(line_stripped.strip("*"))
            run.bold = True
            run.font.name = FONT
            run.font.color.rgb = BLACK
            run.font.size = Pt(12)

        # Bullets
        elif line_stripped.startswith("- ") or line_stripped.startswith("• "):
            if table_active:
                _finish_table()
            doc.add_paragraph(line_stripped[2:], style="List Bullet")

        # Numbered steps
        elif re.match(r'^\d+\.\s', line_stripped):
            if table_active:
                _finish_table()
            text = re.sub(r'^\d+\.\s+', '', line_stripped)
            doc.add_paragraph(text, style="List Number")

        # Tables
        elif line_stripped.startswith("|") and line_stripped.endswith("|"):
            cells = [c.strip() for c in line_stripped.strip("|").split("|")]
            if all(set(c) <= set("- :") for c in cells):
                continue
            if not table_active:
                if _detect_action_table(cells):
                    _build_action_table_header(doc, cells)
                else:
                    _build_generic_table_header(doc, cells)
            else:
                if is_action_table:
                    _build_action_table_row(cells)
                else:
                    _build_generic_table_row(cells)

        # Plain text
        else:
            if table_active:
                _finish_table()
            doc.add_paragraph(line)

    if table_active:
        _finish_table()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_title = "".join(c for c in title if c.isalnum() or c in " -_")[:50].strip() or "procedure"
    return Response(
        content=buf.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{safe_title}.docx"'},
    )
