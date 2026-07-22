#!/usr/bin/env python3
"""
PDFHelper — Secure web API for AI-powered PDF search and flagging.

Deployed on Railway. Provides endpoints to upload PDFs, search them
with keywords or AI, and retrieve flagged results.

Security features:
- API key authentication (required in production)
- File encryption at rest (AES via Fernet)
- PDF magic-byte verification
- Rate limiting, CORS, security headers
- Audit logging of all sensitive operations
- Auto-cleanup of old files
"""

import hashlib
import io
import json
import os
import re
import secrets
import time
import uuid
from datetime import datetime, timezone, timedelta
from pathlib import Path
from typing import Literal

try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass  # python-dotenv not installed; env vars must be set directly

from contextlib import asynccontextmanager

from fastapi import (
    BackgroundTasks, FastAPI, File, Form, UploadFile, Depends, HTTPException, Query, Request,
)
from fastapi.middleware.cors import CORSMiddleware
from fastapi.middleware.trustedhost import TrustedHostMiddleware
from fastapi.responses import JSONResponse, HTMLResponse, StreamingResponse, Response
from pydantic import BaseModel, Field
from starlette.middleware.base import BaseHTTPMiddleware

from database import (
    SessionLocal, engine, Base, DBUser, DBDocument, DBSearchResult, DBAnalysisReport,
    DBChatSession, DBChatMessage, DBDrawing, DBIsolationPackage, DBUpdateSession,
    DBAgentCache, DBPoster,
)
from audit import log_upload, log_search, log_delete, log_auth_failure, log_access
from ocr import extract_text_with_ocr_fallback, extract_structured_text
from search import keyword_search, ai_search

# ---------------------------------------------------------------------------
# Config
# ---------------------------------------------------------------------------

ENVIRONMENT = os.getenv("ENVIRONMENT", "development")
IS_PRODUCTION = ENVIRONMENT == "production"

API_KEY = (os.getenv("PDF_HELPER_API_KEY") or "").strip() or None
ENCRYPTION_KEY = os.getenv("ENCRYPTION_KEY")

MAX_FILE_SIZE = int(os.getenv("MAX_FILE_SIZE_MB", "20")) * 1024 * 1024
MAX_FILES_PER_REQUEST = int(os.getenv("MAX_FILES_PER_REQUEST", "20"))
CHAT_MODEL = os.getenv("CHAT_MODEL", "claude-sonnet-5")
CHAT_MAX_TOKENS = int(os.getenv("CHAT_MAX_TOKENS", "32000"))
AGENT_MAX_TOKENS = int(os.getenv("AGENT_MAX_TOKENS", "12000"))
CHAT_WEB_SEARCH = os.getenv("CHAT_WEB_SEARCH", "true").lower() == "true"
UPLOAD_DIR = Path(os.getenv("UPLOAD_DIR", "/tmp/pdfhelper_uploads"))
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

# Auto-cleanup: delete uploads older than this many hours (0 = disabled)
AUTO_CLEANUP_HOURS = int(os.getenv("AUTO_CLEANUP_HOURS", "72"))

# JWT auth config
JWT_SECRET = os.getenv("JWT_SECRET", "").strip() or secrets.token_hex(32)
JWT_ALGORITHM = "HS256"
JWT_EXPIRY_HOURS = int(os.getenv("JWT_EXPIRY_HOURS", "24"))
# Set to "true" to allow new user registration (default: disabled for single-user)
ALLOW_REGISTRATION = os.getenv("ALLOW_REGISTRATION", "false").lower() == "true"
# Auto-create admin account on startup (set both to enable)
ADMIN_USERNAME = os.getenv("ADMIN_USERNAME", "").strip()
ADMIN_PASSWORD = os.getenv("ADMIN_PASSWORD", "").strip()

# ---------------------------------------------------------------------------
# App setup
# ---------------------------------------------------------------------------

_startup_errors: list[str] = []


async def _retry_db_init(logger):
    """Retry database table creation in the background."""
    import asyncio
    for attempt in range(2, 6):  # attempts 2–5
        await asyncio.sleep(2 ** attempt)  # 4, 8, 16, 32 seconds
        try:
            Base.metadata.create_all(bind=engine)
            logger.info("Database tables initialized successfully (attempt %d).", attempt)
            return
        except Exception as exc:
            logger.warning("Database init attempt %d/5 failed: %s", attempt, exc)
    _startup_errors.append("WARNING: Database initialization failed after 5 attempts.")


@asynccontextmanager
async def lifespan(app: FastAPI):
    """Run safety checks and initialize DB on startup — errors are logged, not fatal."""
    import asyncio
    import logging
    logger = logging.getLogger("pdfhelper")

    if IS_PRODUCTION and not API_KEY:
        msg = (
            "WARNING: PDF_HELPER_API_KEY not set in production. "
            "API endpoints will reject requests until it is configured."
        )
        logger.warning(msg)
        _startup_errors.append(msg)

    if IS_PRODUCTION and not os.getenv("JWT_SECRET", "").strip():
        msg = (
            "WARNING: JWT_SECRET not set in production. "
            "A random secret was generated — JWTs will not survive restarts."
        )
        logger.warning(msg)
        _startup_errors.append(msg)

    if IS_PRODUCTION and not ENCRYPTION_KEY:
        msg = (
            "WARNING: ENCRYPTION_KEY not set in production. "
            "File encryption is disabled until it is configured."
        )
        logger.warning(msg)
        _startup_errors.append(msg)

    # Try once synchronously, then retry in background if needed
    try:
        Base.metadata.create_all(bind=engine)
        logger.info("Database tables initialized successfully.")
    except Exception as exc:
        logger.warning("Database init failed on first attempt: %s — retrying in background", exc)
        asyncio.create_task(_retry_db_init(logger))

    # Auto-create admin if env vars are set and no user with that name exists
    if ADMIN_USERNAME and ADMIN_PASSWORD:
        try:
            db = SessionLocal()
            existing = db.query(DBUser).filter(DBUser.username == ADMIN_USERNAME).first()
            if not existing:
                admin = DBUser(
                    id=str(uuid.uuid4()),
                    username=ADMIN_USERNAME,
                    password_hash=_hash_password(ADMIN_PASSWORD),
                    is_admin=True,
                    created_at=datetime.now(timezone.utc),
                )
                db.add(admin)
                db.commit()
                logger.info("Admin account '%s' created from env vars.", ADMIN_USERNAME)
            db.close()
        except Exception as exc:
            logger.warning("Failed to auto-create admin account: %s", exc)

    yield
    # No shutdown logic needed


app = FastAPI(
    title="PDFHelper",
    description="AI-powered PDF search and flagging tool",
    version="1.0.0",
    docs_url="/docs" if not IS_PRODUCTION else None,
    redoc_url=None,
    lifespan=lifespan,
)

# ---------------------------------------------------------------------------
# Security middleware
# ---------------------------------------------------------------------------

# -- CORS --
ALLOWED_ORIGINS = os.getenv("ALLOWED_ORIGINS", "").split(",")
ALLOWED_ORIGINS = [o.strip() for o in ALLOWED_ORIGINS if o.strip()]
if IS_PRODUCTION and not ALLOWED_ORIGINS:
    ALLOWED_ORIGINS = []  # Block all cross-origin requests by default

app.add_middleware(
    CORSMiddleware,
    allow_origins=ALLOWED_ORIGINS if ALLOWED_ORIGINS else (["*"] if not IS_PRODUCTION else []),
    allow_credentials=True,
    allow_methods=["GET", "POST", "DELETE"],
    allow_headers=["Authorization", "X-API-Key", "Content-Type"],
)

# -- Trusted hosts --
ALLOWED_HOSTS = os.getenv("ALLOWED_HOSTS", "").split(",")
ALLOWED_HOSTS = [h.strip() for h in ALLOWED_HOSTS if h.strip()]
if ALLOWED_HOSTS:
    app.add_middleware(TrustedHostMiddleware, allowed_hosts=ALLOWED_HOSTS)


# -- Security headers --
class SecurityHeadersMiddleware(BaseHTTPMiddleware):
    """Add security headers to every response."""
    async def dispatch(self, request: Request, call_next):
        response = await call_next(request)
        response.headers["X-Content-Type-Options"] = "nosniff"
        response.headers["X-Frame-Options"] = "DENY"
        response.headers["X-XSS-Protection"] = "1; mode=block"
        response.headers["Referrer-Policy"] = "strict-origin-when-cross-origin"
        response.headers["Permissions-Policy"] = "camera=(), microphone=(), geolocation=()"
        if IS_PRODUCTION:
            response.headers["Strict-Transport-Security"] = "max-age=63072000; includeSubDomains"
        # Never reveal server tech
        if "server" in response.headers:
            del response.headers["server"]
        return response


app.add_middleware(SecurityHeadersMiddleware)


# -- HTTPS redirect (production) --
class HTTPSRedirectMiddleware(BaseHTTPMiddleware):
    """Redirect HTTP to HTTPS in production using X-Forwarded-Proto."""
    async def dispatch(self, request: Request, call_next):
        if IS_PRODUCTION:
            proto = request.headers.get("x-forwarded-proto", "https")
            if proto != "https":
                url = request.url.replace(scheme="https")
                return JSONResponse(
                    status_code=301,
                    headers={"Location": str(url)},
                    content={"detail": "Use HTTPS"},
                )
        return await call_next(request)


app.add_middleware(HTTPSRedirectMiddleware)


# -- Rate limiting --
class RateLimitMiddleware(BaseHTTPMiddleware):
    """In-memory rate limiter per IP with periodic cleanup."""
    def __init__(self, app, max_requests: int = 60, window_seconds: int = 60):
        super().__init__(app)
        self.max_requests = max_requests
        self.window = window_seconds
        self.requests: dict[str, list[float]] = {}
        self._last_cleanup = time.time()
        self._cleanup_interval = 300  # purge stale IPs every 5 minutes

    async def dispatch(self, request: Request, call_next):
        client_ip = _get_client_ip(request)
        now = time.time()
        window_start = now - self.window

        # Periodically purge IPs with no recent requests to prevent memory leak
        if now - self._last_cleanup > self._cleanup_interval:
            self.requests = {
                ip: [t for t in ts if t > window_start]
                for ip, ts in self.requests.items()
                if any(t > window_start for t in ts)
            }
            self._last_cleanup = now

        hits = self.requests.get(client_ip, [])
        hits = [t for t in hits if t > window_start]

        if len(hits) >= self.max_requests:
            return JSONResponse(
                status_code=429,
                content={"detail": "Too many requests. Try again later."},
            )
        hits.append(now)
        self.requests[client_ip] = hits

        response = await call_next(request)
        log_access(client_ip, request.method, request.url.path, response.status_code)
        return response


app.add_middleware(RateLimitMiddleware, max_requests=60, window_seconds=60)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _get_client_ip(request: Request) -> str:
    """Get real client IP, respecting proxy headers."""
    forwarded = request.headers.get("x-forwarded-for")
    if forwarded:
        return forwarded.split(",")[0].strip()
    return request.client.host if request.client else "unknown"


def _sanitize_filename(filename: str) -> str:
    """Strip path traversal and dangerous characters from filenames."""
    # Take only the basename
    name = Path(filename).name
    # Remove anything that isn't alphanumeric, dash, underscore, dot, or space
    name = re.sub(r"[^\w\-. ]", "_", name)
    # Prevent hidden files
    name = name.lstrip(".")
    return name or "unnamed.pdf"


def _validate_filepath(filepath: Path) -> Path:
    """Validate that a filepath is within UPLOAD_DIR. Prevents path traversal."""
    try:
        resolved = filepath.resolve()
        if not resolved.is_relative_to(UPLOAD_DIR.resolve()):
            raise ValueError(f"Path escapes upload directory: {filepath}")
        if resolved.is_symlink():
            raise ValueError(f"Symlinks not allowed: {filepath}")
        return resolved
    except (OSError, ValueError):
        raise


def _safe_unlink(filepath: Path) -> None:
    """Safely delete a file, handling race conditions (TOCTOU)."""
    try:
        validated = _validate_filepath(filepath)
        validated.unlink()
    except FileNotFoundError:
        pass  # Already deleted — not an error
    except ValueError:
        import logging
        logging.getLogger("pdfhelper").warning("Blocked path traversal attempt: %s", filepath)


PDF_MAGIC_BYTES = b"%PDF-"

ALLOWED_IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".tiff", ".tif", ".bmp", ".webp"}
ALLOWED_SPREADSHEET_EXTENSIONS = {".xlsx", ".xls", ".csv"}


def _verify_pdf_content(data: bytes) -> bool:
    """Check that the file actually starts with the PDF magic bytes."""
    return data[:5] == PDF_MAGIC_BYTES


def _is_image_file(filename: str) -> bool:
    """Check if a filename has an image extension."""
    return any(filename.lower().endswith(ext) for ext in ALLOWED_IMAGE_EXTENSIONS)


def _image_to_pdf(image_bytes: bytes) -> bytes:
    """Convert an image file to a single-page PDF."""
    from PIL import Image
    import io
    img = Image.open(io.BytesIO(image_bytes))
    if img.mode != "RGB":
        img = img.convert("RGB")
    buf = io.BytesIO()
    img.save(buf, format="PDF", resolution=200)
    return buf.getvalue()


def _extract_image_base64(image_bytes: bytes) -> str:
    """Return a base64-encoded version of the image for Claude vision."""
    import base64
    return base64.b64encode(image_bytes).decode("ascii")


def _detect_image_media_type(image_bytes: bytes) -> str:
    """Detect image MIME type from bytes."""
    if image_bytes[:8] == b'\x89PNG\r\n\x1a\n':
        return "image/png"
    if image_bytes[:2] == b'\xff\xd8':
        return "image/jpeg"
    if image_bytes[:4] in (b'II*\x00', b'MM\x00*'):
        return "image/tiff"
    if image_bytes[:4] == b'RIFF' and image_bytes[8:12] == b'WEBP':
        return "image/webp"
    if image_bytes[:2] == b'BM':
        return "image/bmp"
    return "image/png"  # fallback


def _is_spreadsheet_file(filename: str) -> bool:
    return any(filename.lower().endswith(ext) for ext in ALLOWED_SPREADSHEET_EXTENSIONS)


def _extract_spreadsheet_text(content: bytes, filename: str) -> list[dict]:
    """Extract text from Excel/CSV files as markdown tables, one 'page' per sheet."""
    import io
    import pandas as pd

    pages = []
    lower = filename.lower()

    if lower.endswith(".csv"):
        df = pd.read_csv(io.BytesIO(content))
        text = df.to_markdown(index=False) if hasattr(df, 'to_markdown') else df.to_string(index=False)
        pages.append({"page": 1, "text": f"[Sheet: CSV Data]\n\n{text}"})
    else:
        xls = pd.ExcelFile(io.BytesIO(content))
        for i, sheet_name in enumerate(xls.sheet_names, 1):
            df = pd.read_excel(xls, sheet_name=sheet_name)
            if df.empty:
                continue
            text = df.to_markdown(index=False) if hasattr(df, 'to_markdown') else df.to_string(index=False)
            pages.append({"page": i, "text": f"[Sheet: {sheet_name}]\n\n{text}"})

    if not pages:
        pages.append({"page": 1, "text": "(Empty spreadsheet)"})

    return pages


# ---------------------------------------------------------------------------
# Password hashing & JWT helpers
# ---------------------------------------------------------------------------

_HASH_ITERATIONS = 260_000  # OWASP recommended for PBKDF2-SHA256


def _hash_password(password: str) -> str:
    """Hash a password with PBKDF2-SHA256 + random salt."""
    salt = secrets.token_hex(16)
    h = hashlib.pbkdf2_hmac("sha256", password.encode(), salt.encode(), _HASH_ITERATIONS)
    return f"{salt}:{h.hex()}"


def _verify_password(password: str, stored_hash: str) -> bool:
    """Verify a password against a stored PBKDF2 hash."""
    try:
        salt, hash_hex = stored_hash.split(":", 1)
        h = hashlib.pbkdf2_hmac("sha256", password.encode(), salt.encode(), _HASH_ITERATIONS)
        return secrets.compare_digest(h.hex(), hash_hex)
    except (ValueError, AttributeError):
        return False


def _create_jwt(user_id: str, username: str, is_admin: bool = False) -> str:
    """Create a signed JWT token for authenticated users."""
    import jwt
    payload = {
        "sub": user_id,
        "username": username,
        "admin": is_admin,
        "exp": datetime.now(timezone.utc) + timedelta(hours=JWT_EXPIRY_HOURS),
        "iat": datetime.now(timezone.utc),
    }
    return jwt.encode(payload, JWT_SECRET, algorithm=JWT_ALGORITHM)


def _decode_jwt(token: str) -> dict | None:
    """Decode and verify a JWT token. Returns payload or None."""
    import jwt
    try:
        return jwt.decode(token, JWT_SECRET, algorithms=[JWT_ALGORITHM])
    except (jwt.ExpiredSignatureError, jwt.InvalidTokenError):
        return None


# ---------------------------------------------------------------------------
# Auth dependency
# ---------------------------------------------------------------------------

async def verify_auth(request: Request):
    """Authenticate via JWT Bearer token OR legacy API key."""
    # Dev mode: skip auth only if no API_KEY, no JWT_SECRET, AND no users exist
    if not API_KEY and os.getenv("JWT_SECRET", "").strip() == "":
        try:
            db = SessionLocal()
            has_users = db.query(DBUser).first() is not None
            db.close()
            if not has_users:
                return
        except Exception:
            return

    auth = request.headers.get("Authorization", "")

    # Try JWT Bearer token first
    if auth.startswith("Bearer "):
        token = auth[7:]
        payload = _decode_jwt(token)
        if payload:
            request.state.user_id = payload.get("sub")
            request.state.username = payload.get("username")
            request.state.is_admin = payload.get("admin", False)
            return

    # Fall back to legacy API key (X-API-Key header)
    api_key_token = request.headers.get("X-API-Key", "")
    if API_KEY and api_key_token and secrets.compare_digest(api_key_token, API_KEY):
        request.state.user_id = None  # API key users have no user_id
        request.state.username = "api_key_user"
        request.state.is_admin = False
        return

    # Also accept Bearer token as API key for backward compatibility
    if API_KEY and auth.startswith("Bearer "):
        token = auth[7:]
        if secrets.compare_digest(token, API_KEY):
            request.state.user_id = None
            request.state.username = "api_key_user"
            request.state.is_admin = False
            return

    log_auth_failure(_get_client_ip(request), request.url.path)
    raise HTTPException(status_code=401, detail="Invalid or missing credentials")


# Backward-compatible alias
verify_api_key = verify_auth


# ---------------------------------------------------------------------------
# DB dependency
# ---------------------------------------------------------------------------

def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()


# ---------------------------------------------------------------------------
# Encryption helpers
# ---------------------------------------------------------------------------

def _encrypt_and_save(data: bytes, path: Path) -> None:
    """Encrypt data and write to path. Falls back to plaintext if no key."""
    if ENCRYPTION_KEY:
        from encryption import encrypt_bytes
        path.write_bytes(encrypt_bytes(data))
    else:
        path.write_bytes(data)


def _encrypt_text(text: str) -> str:
    """Encrypt a string for database storage. Returns base64-encoded ciphertext."""
    if ENCRYPTION_KEY:
        from encryption import encrypt_bytes
        import base64
        return base64.b64encode(encrypt_bytes(text.encode("utf-8"))).decode("ascii")
    return text


def _decrypt_text(stored: str) -> str:
    """Decrypt a string read from the database."""
    if ENCRYPTION_KEY:
        from encryption import decrypt_bytes
        import base64
        return decrypt_bytes(base64.b64decode(stored)).decode("utf-8")
    return stored


# ---------------------------------------------------------------------------
# PDF processing
# ---------------------------------------------------------------------------

def extract_text_from_bytes(pdf_bytes: bytes) -> list[dict]:
    """Extract text from PDF bytes. Falls back to OCR for scanned PDFs."""
    return extract_text_with_ocr_fallback(pdf_bytes)



# keyword_search and ai_search are imported from search.py


# ---------------------------------------------------------------------------
# Auto-cleanup
# ---------------------------------------------------------------------------

def _run_cleanup(db) -> int:
    """Delete documents older than AUTO_CLEANUP_HOURS. Returns count deleted."""
    if not AUTO_CLEANUP_HOURS:
        return 0
    cutoff = datetime.now(timezone.utc) - timedelta(hours=AUTO_CLEANUP_HOURS)
    old_docs = db.query(DBDocument).filter(DBDocument.uploaded_at < cutoff).all()
    count = 0
    for doc in old_docs:
        _safe_unlink(Path(doc.filepath))
        db.delete(doc)
        count += 1
    if count:
        db.commit()
    return count


def _run_cleanup_background():
    """Run cleanup with its own DB session (safe for background tasks)."""
    db = SessionLocal()
    try:
        _run_cleanup(db)
    finally:
        db.close()


# ---------------------------------------------------------------------------
# Request / response models
# ---------------------------------------------------------------------------

class SearchRequest(BaseModel):
    search_terms: list[str] = Field(default=[], max_length=50, description="Exact keywords to search")
    ai_query: str | None = Field(default=None, max_length=2000, description="AI concept search query")
    case_sensitive: bool = False


class AnalyzeRequest(BaseModel):
    compliance_context: str | None = Field(
        default=None,
        max_length=2000,
        description="Optional compliance standard to check against, e.g. 'OSHA 2024', 'HIPAA', 'FDA 21 CFR Part 11'",
    )
    search_terms: list[str] = Field(default=[], max_length=50, description="Optional keywords to search for")
    ai_query: str | None = Field(default=None, max_length=2000, description="Optional AI concept search query")


class ChatMessage(BaseModel):
    role: Literal["user", "assistant"] = Field(description="'user' or 'assistant'")
    content: str = Field(max_length=200000)


class ChatRequest(BaseModel):
    message: str = Field(max_length=50000, description="The user's message")
    doc_ids: list[str] = Field(default=[], max_length=100, description="Document IDs to use as context (empty = all)")
    conversation_history: list[ChatMessage] = Field(default=[], max_length=200, description="Previous messages for context")
    session_id: str | None = Field(default=None, max_length=100, description="Chat session ID to continue (omit to create new)")
    model: str = Field(default="", pattern=r"^(sonnet|haiku|)$", description="Model to use: 'sonnet' or 'haiku' (empty = default)")


class RegisterRequest(BaseModel):
    username: str = Field(min_length=3, max_length=50, pattern=r"^[a-zA-Z0-9_]+$")
    password: str = Field(min_length=8, max_length=128)


class LoginRequest(BaseModel):
    username: str = Field(max_length=50)
    password: str = Field(max_length=128)


class PosterCreateRequest(BaseModel):
    prompt: str = Field(max_length=30000, description="Describe the poster you want to create")
    size: str = Field(default="letter", pattern=r"^(letter|a4|a3|wide|square|banner)$")
    style: str = Field(default="", max_length=50, description="Optional style preset")
    model: str = Field(default="haiku", pattern=r"^(sonnet|haiku)$")

class PosterUpdateRequest(BaseModel):
    prompt: str = Field(max_length=30000, description="Describe what to change on the poster")
    model: str = Field(default="haiku", pattern=r"^(sonnet|haiku)$")


class HealthResponse(BaseModel):
    status: str
    version: str
    api_key_required: bool = False
    has_users: bool = False
    warnings: list[str] = []


# Valid work types for isolation requests
_VALID_WORK_TYPES = {
    "MAINTENANCE", "HOT WORK", "CONFINED SPACE ENTRY", "PRESSURE TEST",
    "INSPECTION", "EQUIPMENT REMOVAL", "ELECTRICAL ISOLATION", "INSTRUMENT MAINTENANCE",
}


class IsolationRequest(BaseModel):
    equipment_tag: str = Field(max_length=200, description="Equipment tag, e.g. HB-P-1001A")
    work_description: str = Field(max_length=5000, description="Description of work to be performed")
    work_type: str = Field(max_length=100, description="MAINTENANCE, HOT WORK, CONFINED SPACE ENTRY, PRESSURE TEST, INSPECTION, EQUIPMENT REMOVAL, ELECTRICAL ISOLATION, INSTRUMENT MAINTENANCE")
    fluid_service: str = Field(default="Not specified", max_length=200, description="Fluid service, e.g. Crude Oil, HC Gas, Produced Water")
    special_requirements: str = Field(default="None", max_length=5000, description="Any special requirements")
    facility: str = Field(default="Hebron", max_length=200, description="Facility name")
    regime: str = Field(default="C-NLOPB / C-NLOER", max_length=200, description="Regulatory regime")
    drawing_ids: list[str] = Field(default=[], max_length=20, description="Specific drawing IDs to use (empty = auto-select via Pass 1)")


# ---------------------------------------------------------------------------
# Validation
# ---------------------------------------------------------------------------

def validate_upload(file: UploadFile, content: bytes) -> tuple[str, bytes, bool, bool]:
    """Validate an uploaded file. Returns (sanitized_name, file_bytes, is_image, is_spreadsheet)."""
    if not file.filename:
        raise HTTPException(status_code=400, detail="Filename is required")

    clean_name = _sanitize_filename(file.filename)
    is_image = _is_image_file(clean_name)
    is_spreadsheet = _is_spreadsheet_file(clean_name)

    if not clean_name.lower().endswith(".pdf") and not is_image and not is_spreadsheet:
        raise HTTPException(status_code=400,
                            detail=f"Only PDF, image, and Excel/CSV files allowed, got: {clean_name}")

    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=400,
            detail=f"{clean_name} exceeds max size of {MAX_FILE_SIZE // (1024*1024)} MB",
        )

    if is_spreadsheet:
        return clean_name, content, False, True
    elif is_image:
        try:
            pdf_bytes = _image_to_pdf(content)
        except Exception as exc:
            raise HTTPException(status_code=400,
                                detail=f"Could not process image {clean_name}: {exc}")
        return clean_name, pdf_bytes, True, False
    else:
        if not _verify_pdf_content(content):
            raise HTTPException(status_code=400,
                                detail="File does not appear to be a valid PDF")
        return clean_name, content, False, False


# ---------------------------------------------------------------------------
# Root UI (loaded from templates/index.html)
# ---------------------------------------------------------------------------

_TEMPLATE_DIR = Path(__file__).parent / "templates"
_cached_html: str | None = None


def _load_root_html() -> str:
    """Load the frontend HTML template. Cached in production to avoid disk reads."""
    global _cached_html
    if IS_PRODUCTION and _cached_html is not None:
        return _cached_html
    html = (_TEMPLATE_DIR / "index.html").read_text(encoding="utf-8")
    if IS_PRODUCTION:
        _cached_html = html
    return html



# ---------------------------------------------------------------------------
# Global error handler for DB failures
# ---------------------------------------------------------------------------

from sqlalchemy.exc import OperationalError as SAOperationalError

@app.exception_handler(SAOperationalError)
async def db_error_handler(request: Request, exc: SAOperationalError):
    """Return a clear error message when the database is unreachable."""
    return JSONResponse(
        status_code=503,
        content={
            "detail": "Database is unavailable. Check that DATABASE_URL is set correctly "
                      "(use the public URL, not the internal .railway.internal hostname)."
        },
    )

@app.exception_handler(Exception)
async def general_error_handler(request: Request, exc: Exception):
    """Catch-all so errors always return JSON, never HTML."""
    import traceback, logging
    logging.getLogger("pdfhelper").error("Unhandled error on %s: %s\n%s", request.url.path, exc, traceback.format_exc())
    if IS_PRODUCTION:
        detail = "Internal server error"
    else:
        detail = f"Internal server error: {type(exc).__name__}: {str(exc)}"
    return JSONResponse(
        status_code=500,
        content={"detail": detail},
    )


# ---------------------------------------------------------------------------
# Endpoints
# ---------------------------------------------------------------------------

@app.get("/health", response_model=HealthResponse)
async def health_check():
    """Health check — also tests current DB connectivity."""
    db_ok = True
    db_err = None
    try:
        from sqlalchemy import text
        db = SessionLocal()
        db.execute(text("SELECT 1"))
        db.close()
    except Exception as e:
        db_ok = False
        # Log full error for operators, but don't expose connection details to clients
        import logging
        logging.getLogger("pdfhelper").error("Health check DB failure: %s", e)
        db_err = "Database connection failed"

    warnings = list(_startup_errors)
    if db_err:
        warnings.append(db_err)

    has_users = False
    if db_ok:
        try:
            db = SessionLocal()
            has_users = db.query(DBUser).first() is not None
            db.close()
        except Exception:
            pass

    status = "ok" if (not warnings and db_ok) else "degraded"
    return {
        "status": status,
        "version": "1.0.0",
        "api_key_required": True,
        "has_users": has_users,
        "warnings": warnings,
    }


@app.get("/verify-key", dependencies=[Depends(verify_api_key)])
async def verify_key():
    """Lightweight API key check — no database required."""
    import asyncio
    from sqlalchemy import text

    def _check_db():
        db = SessionLocal()
        try:
            db.execute(text("SELECT 1"))
        finally:
            db.close()

    db_ok = True
    try:
        await asyncio.wait_for(asyncio.to_thread(_check_db), timeout=3)
    except Exception:
        db_ok = False
    return {"valid": True, "db_ok": db_ok}


@app.get("/", response_class=HTMLResponse)
async def root():
    """Interactive operating interface for PDFHelper."""
    return _load_root_html()


@app.get("/bot")
async def bot_page():
    """Redirect old bot page to main app."""
    from starlette.responses import RedirectResponse
    return RedirectResponse(url="/", status_code=301)


# ---------------------------------------------------------------------------
# User registration & login
# ---------------------------------------------------------------------------

@app.get("/setup-needed")
async def setup_needed(db=Depends(get_db)):
    """Check if initial setup is required (no users exist yet)."""
    any_user = db.query(DBUser).first()
    return {"setup_needed": any_user is None}


@app.post("/setup")
async def setup(body: RegisterRequest, db=Depends(get_db)):
    """First-run setup: create the initial admin account. Only works when no users exist."""
    any_user = db.query(DBUser).first()
    if any_user:
        raise HTTPException(status_code=403, detail="Setup already completed. Use /login instead.")

    user_id = str(uuid.uuid4())
    now = datetime.now(timezone.utc)
    user = DBUser(
        id=user_id,
        username=body.username,
        password_hash=_hash_password(body.password),
        is_admin=True,
        created_at=now,
    )
    db.add(user)
    db.commit()

    token = _create_jwt(user_id, body.username, is_admin=True)
    return {"user_id": user_id, "username": body.username, "token": token}


@app.post("/register")
async def register(body: RegisterRequest, db=Depends(get_db)):
    """Create a new user account. Returns a JWT token."""
    if not ALLOW_REGISTRATION:
        raise HTTPException(status_code=403, detail="Registration is disabled. Contact an admin.")

    existing = db.query(DBUser).filter(DBUser.username == body.username).first()
    if existing:
        raise HTTPException(status_code=409, detail="Username already taken")

    user_id = str(uuid.uuid4())
    now = datetime.now(timezone.utc)
    user = DBUser(
        id=user_id,
        username=body.username,
        password_hash=_hash_password(body.password),
        is_admin=False,
        created_at=now,
    )
    db.add(user)
    db.commit()

    token = _create_jwt(user_id, body.username)
    return {"user_id": user_id, "username": body.username, "token": token}


@app.post("/login")
async def login(body: LoginRequest, db=Depends(get_db)):
    """Authenticate and receive a JWT token."""
    user = db.query(DBUser).filter(DBUser.username == body.username).first()
    if not user or not _verify_password(body.password, user.password_hash):
        raise HTTPException(status_code=401, detail="Invalid username or password")

    token = _create_jwt(user.id, user.username, user.is_admin)
    return {"user_id": user.id, "username": user.username, "token": token}


@app.get("/me", dependencies=[Depends(verify_auth)])
async def get_current_user(request: Request):
    """Return the currently authenticated user's info."""
    return {
        "user_id": getattr(request.state, "user_id", None),
        "username": getattr(request.state, "username", None),
        "is_admin": getattr(request.state, "is_admin", False),
    }


@app.post("/upload", dependencies=[Depends(verify_api_key)])
async def upload_pdfs(
    request: Request,
    files: list[UploadFile] = File(...),
    background_tasks: BackgroundTasks = None,
    db=Depends(get_db),
):
    """Upload one or more PDFs, images, or Excel/CSV files for later searching.

    Supported formats: PDF, JPG, PNG, TIFF, BMP, WebP, XLSX, XLS, CSV.
    Images are automatically converted to PDF and OCR'd for text extraction.
    Spreadsheets are converted to markdown tables for AI consumption.
    """
    if len(files) > MAX_FILES_PER_REQUEST:
        raise HTTPException(status_code=400,
                            detail=f"Max {MAX_FILES_PER_REQUEST} files per request")

    client_ip = _get_client_ip(request)

    if background_tasks:
        background_tasks.add_task(_run_cleanup_background)

    uploaded = []
    for file in files:
        raw_content = await file.read()
        clean_name, file_bytes, is_image, is_spreadsheet = validate_upload(file, raw_content)

        if is_spreadsheet:
            pages = _extract_spreadsheet_text(raw_content, clean_name)
        else:
            pages = extract_text_from_bytes(file_bytes)

        doc_id = str(uuid.uuid4())
        save_path = UPLOAD_DIR / f"{doc_id}.pdf.enc"

        _encrypt_and_save(raw_content, save_path)

        if is_image:
            img_save_path = UPLOAD_DIR / f"{doc_id}.img.enc"
            _encrypt_and_save(raw_content, img_save_path)

        content_hash = hashlib.sha256(raw_content).hexdigest()
        file_type = "spreadsheet" if is_spreadsheet else ("image" if is_image else "pdf")

        db_doc = DBDocument(
            id=doc_id,
            filename=_encrypt_text(clean_name),
            filepath=str(save_path),
            page_count=len(pages),
            text_content=_encrypt_text(json.dumps(pages)),
            content_hash=content_hash,
            uploaded_at=datetime.now(timezone.utc),
        )
        db.add(db_doc)
        db.commit()

        log_upload(client_ip, clean_name, doc_id, len(pages))

        uploaded.append({
            "id": doc_id,
            "filename": clean_name,
            "pages": len(pages),
            "type": file_type,
        })

    return {"uploaded": uploaded, "count": len(uploaded)}


@app.post("/search", dependencies=[Depends(verify_api_key)])
async def search_documents(
    request: Request,
    body: SearchRequest,
    doc_ids: list[str] = Query(default=[], description="Document IDs to search (empty = all)"),
    db=Depends(get_db),
):
    """Search uploaded PDFs with keywords and/or AI."""
    if not body.search_terms and not body.ai_query:
        raise HTTPException(status_code=400,
                            detail="Provide search_terms and/or ai_query")

    query = db.query(DBDocument)
    if doc_ids:
        query = query.filter(DBDocument.id.in_(doc_ids))
    documents = query.all()

    if not documents:
        raise HTTPException(status_code=404, detail="No documents found")

    client_ip = _get_client_ip(request)
    all_keyword_results = []
    all_ai_results = []

    for doc in documents:
        pages = json.loads(_decrypt_text(doc.text_content))
        decrypted_name = _decrypt_text(doc.filename)

        if body.search_terms:
            matches = keyword_search(pages, body.search_terms, body.case_sensitive)
            for m in matches:
                m["document_id"] = doc.id
                m["filename"] = decrypted_name
            all_keyword_results.extend(matches)

        if body.ai_query:
            findings = ai_search(pages, body.ai_query, decrypted_name)
            for f in findings:
                f["document_id"] = doc.id
                f["filename"] = decrypted_name
            all_ai_results.extend(findings)

    search_id = str(uuid.uuid4())
    flagged_count = len([r for r in all_ai_results if r.get("needs_review")])

    db_result = DBSearchResult(
        id=search_id,
        search_terms=_encrypt_text(json.dumps(body.search_terms)) if body.search_terms else None,
        ai_query=_encrypt_text(body.ai_query) if body.ai_query else None,
        keyword_results=_encrypt_text(json.dumps(all_keyword_results)),
        ai_results=_encrypt_text(json.dumps(all_ai_results)),
        total_keyword_matches=len(all_keyword_results),
        total_ai_findings=len(all_ai_results),
        flagged_for_review=flagged_count,
        searched_at=datetime.now(timezone.utc),
    )
    db.add(db_result)
    db.commit()

    log_search(client_ip, search_id, body.search_terms, body.ai_query,
               len(documents), len(all_keyword_results) + len(all_ai_results),
               flagged_count)

    return {
        "search_id": search_id,
        "summary": {
            "documents_searched": len(documents),
            "total_keyword_matches": len(all_keyword_results),
            "total_ai_findings": len(all_ai_results),
            "flagged_for_review": flagged_count,
        },
        "keyword_results": all_keyword_results,
        "ai_results": all_ai_results,
    }


@app.get("/documents", dependencies=[Depends(verify_api_key)])
async def list_documents(db=Depends(get_db)):
    """List all uploaded documents."""
    docs = db.query(DBDocument).order_by(DBDocument.uploaded_at.desc()).all()
    return {
        "documents": [
            {
                "id": d.id,
                "filename": _decrypt_text(d.filename),
                "pages": d.page_count,
                "uploaded_at": d.uploaded_at.isoformat(),
            }
            for d in docs
        ]
    }


@app.get("/documents/{doc_id}", dependencies=[Depends(verify_api_key)])
async def get_document(doc_id: str, db=Depends(get_db)):
    """Get details for a specific document."""
    doc = db.query(DBDocument).filter(DBDocument.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    return {
        "id": doc.id,
        "filename": _decrypt_text(doc.filename),
        "pages": doc.page_count,
        "uploaded_at": doc.uploaded_at.isoformat(),
    }


@app.delete("/documents/{doc_id}", dependencies=[Depends(verify_api_key)])
async def delete_document(doc_id: str, request: Request, db=Depends(get_db)):
    """Delete an uploaded document and its encrypted file."""
    doc = db.query(DBDocument).filter(DBDocument.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    _safe_unlink(Path(doc.filepath))

    log_delete(_get_client_ip(request), doc_id, _decrypt_text(doc.filename))

    db.delete(doc)
    db.commit()
    return {"deleted": doc_id}


@app.get("/history", dependencies=[Depends(verify_api_key)])
async def search_history(limit: int = Query(default=20, le=100), db=Depends(get_db)):
    """Get past search results."""
    results = (
        db.query(DBSearchResult)
        .order_by(DBSearchResult.searched_at.desc())
        .limit(limit)
        .all()
    )
    return {
        "searches": [
            {
                "id": r.id,
                "search_terms": json.loads(_decrypt_text(r.search_terms)) if r.search_terms else [],
                "ai_query": _decrypt_text(r.ai_query) if r.ai_query else None,
                "total_keyword_matches": r.total_keyword_matches,
                "total_ai_findings": r.total_ai_findings,
                "flagged_for_review": r.flagged_for_review,
                "searched_at": r.searched_at.isoformat(),
            }
            for r in results
        ]
    }


@app.get("/history/{search_id}", dependencies=[Depends(verify_api_key)])
async def get_search_result(search_id: str, db=Depends(get_db)):
    """Get full details of a past search."""
    result = db.query(DBSearchResult).filter(DBSearchResult.id == search_id).first()
    if not result:
        raise HTTPException(status_code=404, detail="Search result not found")
    return {
        "id": result.id,
        "search_terms": json.loads(_decrypt_text(result.search_terms)) if result.search_terms else [],
        "ai_query": _decrypt_text(result.ai_query) if result.ai_query else None,
        "keyword_results": json.loads(_decrypt_text(result.keyword_results)),
        "ai_results": json.loads(_decrypt_text(result.ai_results)),
        "summary": {
            "total_keyword_matches": result.total_keyword_matches,
            "total_ai_findings": result.total_ai_findings,
            "flagged_for_review": result.flagged_for_review,
        },
        "searched_at": result.searched_at.isoformat(),
    }


# ---------------------------------------------------------------------------
# Procedure Chatbot
# ---------------------------------------------------------------------------

@app.post("/chat", dependencies=[Depends(verify_api_key)])
async def chat_with_documents(
    request: Request,
    body: ChatRequest,
    db=Depends(get_db),
):
    """Chat with your uploaded documents using AI.

    Sends the user's message along with selected document content to Claude
    and returns a context-aware response with procedure citations.

    If session_id is provided, continues that session (loading history from DB).
    Otherwise creates a new session. All messages are persisted.
    """
    api_key = os.getenv("ANTHROPIC_API_KEY")
    if not api_key:
        raise HTTPException(status_code=500, detail="ANTHROPIC_API_KEY not configured on server")

    query = db.query(DBDocument)
    if body.doc_ids:
        query = query.filter(DBDocument.id.in_(body.doc_ids))
    documents = query.all()

    if not documents:
        raise HTTPException(status_code=404, detail="No documents found. Upload documents first.")

    # Resolve or create chat session
    now = datetime.now(timezone.utc)
    current_user_id = getattr(request.state, "user_id", None)
    session = None
    if body.session_id:
        session = db.query(DBChatSession).filter(DBChatSession.id == body.session_id).first()
        # Enforce session ownership: user can only access their own sessions
        if session and current_user_id and session.user_id and session.user_id != current_user_id:
            raise HTTPException(status_code=403, detail="You do not own this chat session")

    if session is None:
        session = DBChatSession(
            id=str(uuid.uuid4()),
            user_id=current_user_id,
            title=body.message[:100],
            doc_ids=json.dumps(body.doc_ids),
            created_at=now,
            updated_at=now,
        )
        db.add(session)
        db.flush()

    # Build procedure context from selected documents using structured extraction
    procedure_parts = []
    image_content_blocks = []  # Claude vision blocks for uploaded images
    for doc in documents:
        decrypted_name = _decrypt_text(doc.filename)
        try:
            pdf_bytes = _load_pdf_bytes(doc)
            structured = extract_structured_text(pdf_bytes)
            parts = []
            for page_data in structured:
                parts.append(f"\n--- Page {page_data['page']} ---")
                for block in page_data.get("blocks", []):
                    prefix = f"[{block['type'].upper()}] " if block['type'] != 'paragraph' else ""
                    parts.append(f"{prefix}{block['text']}")
            full_text = "\n".join(parts)
        except Exception:
            pages = json.loads(_decrypt_text(doc.text_content))
            full_text = "\n".join(p["text"] for p in pages if p.get("text"))
        if len(full_text) > 80000:
            full_text = full_text[:80000] + "\n\n[... content truncated for context window ...]"
        procedure_parts.append(
            f'--- PROCEDURE: "{decrypted_name}" ---\n{full_text}\n--- END OF "{decrypted_name}" ---'
        )

        # Check if this document has an associated image file for vision
        img_path = Path(doc.filepath.replace(".pdf.enc", ".img.enc"))
        if img_path.exists() and len(image_content_blocks) < 10:  # max 10 images
            try:
                img_bytes = _decrypt_and_load(img_path)
                media_type = _detect_image_media_type(img_bytes)
                b64 = _extract_image_base64(img_bytes)
                image_content_blocks.append({
                    "type": "text",
                    "text": f"[Image: {decrypted_name}]"
                })
                image_content_blocks.append({
                    "type": "image",
                    "source": {
                        "type": "base64",
                        "media_type": media_type,
                        "data": b64,
                    }
                })
            except Exception:
                pass  # skip if image can't be loaded

    doc_index_lines = ["DOCUMENT INDEX:"]
    for i, doc in enumerate(documents, 1):
        decrypted_name = _decrypt_text(doc.filename)
        doc_index_lines.append(f"  {i}. \"{decrypted_name}\" — {doc.page_count} pages")
    doc_index = "\n".join(doc_index_lines)
    procedure_context = doc_index + "\n\n" + "\n\n".join(procedure_parts)

    # Build conversation from DB history (prefer DB over client-sent history)
    db_messages = (
        db.query(DBChatMessage)
        .filter(DBChatMessage.session_id == session.id)
        .order_by(DBChatMessage.created_at)
        .all()
    )

    if db_messages:
        conversation = [
            {"role": m.role, "content": _decrypt_text(m.content)}
            for m in db_messages[-40:]
        ]
    else:
        conversation = [
            {"role": m.role, "content": m.content}
            for m in body.conversation_history[-40:]
            if m.role in ("user", "assistant")
        ]
    # Build the user message — include images via Claude vision only on the
    # first message of a session so they aren't re-sent on every turn.
    include_images = image_content_blocks and not db_messages
    if include_images:
        user_content = list(image_content_blocks)  # copy
        user_content.append({"type": "text", "text": body.message})
        conversation.append({"role": "user", "content": user_content})
    else:
        conversation.append({"role": "user", "content": body.message})

    # Budget the total context to stay within the model's context window.
    # Reserve chars for the system prompt template, response tokens, and safety margin.
    # Approximate: 1 token ≈ 4 chars.  Model context ≈ 200K tokens ≈ 800K chars.
    # Each image ≈ 1600 tokens, so subtract from budget when included.
    MAX_TOTAL_CHARS = 600000  # leave headroom for response + system template
    if include_images:
        image_char_budget = len(image_content_blocks) // 2 * 6400  # ~1600 tokens * 4 chars per image
        MAX_TOTAL_CHARS -= image_char_budget

    def _msg_text_len(m):
        c = m["content"]
        if isinstance(c, str):
            return len(c)
        if isinstance(c, list):
            return sum(len(b.get("text", "")) for b in c if isinstance(b, dict) and b.get("type") == "text")
        return 0

    conv_chars = sum(_msg_text_len(m) for m in conversation)

    # If conversation history alone is too large, trim older messages (keep latest)
    while conv_chars > 200000 and len(conversation) > 1:
        removed = conversation.pop(0)
        conv_chars -= _msg_text_len(removed)

    budget_for_procedures = MAX_TOTAL_CHARS - conv_chars
    if budget_for_procedures < 10000:
        budget_for_procedures = 10000  # always keep at least some procedure context

    # Truncate procedure context if it exceeds budget
    if len(procedure_context) > budget_for_procedures:
        procedure_context = procedure_context[:budget_for_procedures] + "\n\n[... procedures truncated to fit context window ...]"

    system_prompt = """You are a Procedure Knowledge Assistant. You have two modes:

1. **Q&A MODE** (default): Answer questions about the procedure documents loaded below. Ground every answer in the documents with quotes and citations.
2. **BUILD MODE**: When the user explicitly asks you to build, create, generate, or develop code (HTML, CSS, JavaScript, etc.), generate the requested code using ONLY data and content from the loaded procedure documents.

IMPORTANT RULES:
- NEVER switch to Build Mode on your own. Only generate code when the user explicitly asks for it (e.g. "build an HTML page", "create a form", "generate a dashboard").
- In Build Mode, ALL data, text, steps, and content in the generated code MUST come from the loaded procedure documents — do NOT invent or fabricate content.
- Use web search ONLY to find regulatory references or standards mentioned in the procedures. NEVER use web search to find software templates or unrelated tools.
- When answering questions (Q&A Mode), NEVER offer to build an application unless the user asks.

CODE QUALITY RULES (Build Mode):
- Generate COMPLETE, working, single-file HTML with all CSS and JavaScript embedded inline. Never use external CDN links.
- Test your logic mentally before writing — ensure all variables are defined, all functions are called correctly, and all event listeners are properly attached.
- Use clean, modern HTML5 with semantic elements. Include responsive CSS so it works on mobile and desktop.
- Add clear section headings and professional styling. Use a clean color scheme.
- Make sure all buttons, forms, and interactive elements actually work — wire up every onclick/onsubmit handler.
- Include all the data from the procedures — never use placeholder text like "lorem ipsum" or "TODO".
- Output the COMPLETE code in a single code block. Never say "rest of code here" or truncate the output.

CRITICAL ACCURACY RULES:
1. GROUND EVERY CLAIM: Every factual statement you make MUST be traceable to a specific page in the loaded procedures. If you cannot find it in the documents, say "I could not find this in the loaded procedures" — do NOT guess, infer, or fill in from general knowledge.
2. QUOTE BEFORE PARAPHRASING: When answering, first provide the exact relevant quote from the source (in a blockquote), then explain it. This forces accuracy and lets the user verify.
3. CITE PRECISELY: Always cite the procedure name AND page number (e.g. **"WMS Manual 4.0.1" — Page 12**). Every answer must have at least one citation.
4. DISTINGUISH SOURCES: Clearly separate what comes from the loaded procedures vs. web search vs. your general knowledge:
   - Procedure content: cite normally with document name + page
   - Web search results: label as *(Source: web search)*
   - General knowledge (ONLY if explicitly asked): label as *(General knowledge — verify against your procedures)*
5. NEVER HALLUCINATE PROCEDURE CONTENT: If you're unsure whether something is in the documents, re-read the relevant sections before answering. It is better to say "I'm not sure" than to state something incorrectly.
6. CONFLICT DETECTION: If a question spans multiple procedures, reference all relevant ones and explicitly flag any differences or contradictions between them.

RESPONSE FORMAT:
- Simple factual questions: 1-3 sentences with a direct quote and citation
- Step-by-step procedure questions: numbered list preserving the exact steps from the source
- Comparison questions: markdown table with citations per cell
- Analysis questions: structured answer with headings, but every claim cited
- Use **bold** for procedure names and key terms
- Preserve table structure from the source when a [TABLE] block is referenced
- When uncertain: "**Note:** This requires verification — the procedure is unclear on this point."
- Do NOT add follow-up question suggestions unless the question is genuinely broad"""

    # Use structured system prompt with cache_control for Anthropic prompt caching.
    # The rules block is cached (stable across messages), and the procedures block
    # is cached separately (stable within a session). This means follow-up messages
    # in the same session reuse cached tokens instead of re-processing everything,
    # cutting input costs by up to 90%.
    system_blocks = [
        {
            "type": "text",
            "text": system_prompt,
            "cache_control": {"type": "ephemeral"},
        },
        {
            "type": "text",
            "text": f"LOADED PROCEDURES:\n{procedure_context}",
            "cache_control": {"type": "ephemeral"},
        },
    ]

    from anthropic import Anthropic
    client = Anthropic(api_key=api_key)

    # Save user message now (assistant message saved after stream completes)
    db.add(DBChatMessage(
        id=str(uuid.uuid4()), session_id=session.id,
        role="user", content=_encrypt_text(body.message), created_at=now,
    ))
    db.commit()

    session_id = session.id
    doc_info = [{"id": d.id, "filename": _decrypt_text(d.filename)} for d in documents]

    # Configure tools — optionally include web search
    chat_tools = []
    if CHAT_WEB_SEARCH:
        chat_tools.append({"type": "web_search_20250305", "name": "web_search"})

    async def stream_chat():
        """Stream the AI response as Server-Sent Events."""
        full_reply = ""

        try:
            yield f"data: {json.dumps({'type': 'meta', 'session_id': session_id, 'documents_used': doc_info})}\n\n"

            chat_model = AGENT_MODELS.get(body.model, CHAT_MODEL) if body.model else CHAT_MODEL
            create_kwargs = dict(
                model=chat_model,
                max_tokens=CHAT_MAX_TOKENS,
                system=system_blocks,
                messages=conversation,
            )
            if chat_tools:
                create_kwargs["tools"] = chat_tools

            max_retries = 3
            for attempt in range(max_retries):
                try:
                    with client.messages.stream(**create_kwargs) as stream:
                        for event in stream:
                            if hasattr(event, 'type'):
                                if event.type == 'content_block_start':
                                    if hasattr(event.content_block, 'type'):
                                        if event.content_block.type == 'server_tool_use':
                                            yield f"data: {json.dumps({'type': 'status', 'message': 'Searching the web...'})}\n\n"
                                        elif event.content_block.type == 'thinking':
                                            yield f"data: {json.dumps({'type': 'status', 'message': 'Thinking...'})}\n\n"
                                elif event.type == 'content_block_delta':
                                    if hasattr(event.delta, 'text'):
                                        full_reply += event.delta.text
                                        yield f"data: {json.dumps({'type': 'chunk', 'text': event.delta.text})}\n\n"
                    break
                except Exception as retry_err:
                    err_str = str(retry_err)
                    is_retryable = "overloaded" in err_str.lower() or "529" in err_str or "rate" in err_str.lower()
                    if is_retryable and attempt < max_retries - 1:
                        import asyncio
                        wait_time = 2 ** (attempt + 1)
                        yield f"data: {json.dumps({'type': 'status', 'message': f'Server busy, retrying in {wait_time}s...'})}\n\n"
                        await asyncio.sleep(wait_time)
                        full_reply = ""
                        continue
                    raise

            yield f"data: {json.dumps({'type': 'done'})}\n\n"
        except Exception as e:
            import logging
            logging.getLogger("pdfhelper").error("Chat AI stream failed: %s", e)
            err_msg = "AI request failed — please try again" if IS_PRODUCTION else f"AI request failed: {str(e)}"
            yield f"data: {json.dumps({'type': 'error', 'detail': err_msg})}\n\n"
            full_reply = full_reply or err_msg
        finally:
            # Persist assistant reply after stream completes
            save_db = SessionLocal()
            try:
                save_db.add(DBChatMessage(
                    id=str(uuid.uuid4()), session_id=session_id,
                    role="assistant",
                    content=_encrypt_text(full_reply or "Sorry, I couldn't generate a response."),
                    created_at=datetime.now(timezone.utc),
                ))
                sess = save_db.query(DBChatSession).filter(DBChatSession.id == session_id).first()
                if sess:
                    sess.updated_at = datetime.now(timezone.utc)
                save_db.commit()
            except Exception:
                save_db.rollback()
            finally:
                save_db.close()

    return StreamingResponse(stream_chat(), media_type="text/event-stream")


# ---------------------------------------------------------------------------
# Chat History Endpoints
# ---------------------------------------------------------------------------

@app.get("/chat/sessions", dependencies=[Depends(verify_api_key)])
async def list_chat_sessions(request: Request, limit: int = Query(default=30, le=100), db=Depends(get_db)):
    """List past chat sessions, most recent first. Filtered to current user."""
    current_user_id = getattr(request.state, "user_id", None)
    query = db.query(DBChatSession)
    if current_user_id:
        query = query.filter(
            (DBChatSession.user_id == current_user_id) | (DBChatSession.user_id.is_(None))
        )
    sessions = query.order_by(DBChatSession.updated_at.desc()).limit(limit).all()
    return {
        "sessions": [
            {
                "id": s.id,
                "title": s.title,
                "doc_ids": json.loads(s.doc_ids),
                "message_count": len(s.messages),
                "created_at": s.created_at.isoformat(),
                "updated_at": s.updated_at.isoformat(),
            }
            for s in sessions
        ]
    }


@app.get("/chat/sessions/{session_id}", dependencies=[Depends(verify_api_key)])
async def get_chat_session(session_id: str, request: Request, db=Depends(get_db)):
    """Get full message history for a chat session."""
    session = db.query(DBChatSession).filter(DBChatSession.id == session_id).first()
    if not session:
        raise HTTPException(status_code=404, detail="Chat session not found")
    current_user_id = getattr(request.state, "user_id", None)
    if current_user_id and session.user_id and session.user_id != current_user_id:
        raise HTTPException(status_code=403, detail="You do not own this chat session")
    return {
        "id": session.id,
        "title": session.title,
        "doc_ids": json.loads(session.doc_ids),
        "created_at": session.created_at.isoformat(),
        "updated_at": session.updated_at.isoformat(),
        "messages": [
            {
                "role": m.role,
                "content": _decrypt_text(m.content),
                "created_at": m.created_at.isoformat(),
            }
            for m in session.messages
        ],
    }


@app.delete("/chat/sessions/{session_id}", dependencies=[Depends(verify_api_key)])
async def delete_chat_session(session_id: str, request: Request, db=Depends(get_db)):
    """Delete a chat session and all its messages."""
    session = db.query(DBChatSession).filter(DBChatSession.id == session_id).first()
    if not session:
        raise HTTPException(status_code=404, detail="Chat session not found")
    current_user_id = getattr(request.state, "user_id", None)
    if current_user_id and session.user_id and session.user_id != current_user_id:
        raise HTTPException(status_code=403, detail="You do not own this chat session")
    db.delete(session)
    db.commit()
    return {"deleted": session_id}


# ---------------------------------------------------------------------------
# Chat → Word Document Export
# ---------------------------------------------------------------------------

class ExportChatRequest(BaseModel):
    session_id: str = Field(max_length=100, description="Chat session to export")
    format: str = Field(default="docx", pattern=r"^docx$", description="Export format (docx)")


class GenerateDocRequest(BaseModel):
    session_id: str = Field(max_length=100, description="Chat session for context")
    instructions: str = Field(max_length=10000, description="What the document should contain, e.g. 'Create a safety procedure for valve isolation'")
    doc_ids: list[str] = Field(default=[], max_length=100, description="Document IDs to reference")
    title: str = Field(default="Generated Document", max_length=255)
    include_vba: bool = Field(default=False, description="Include VBA macro code for Word formatting automation")


def _markdown_to_docx(text: str, title: str = "Document") -> bytes:
    """Convert markdown-ish AI text to a formatted Word document."""
    from docx import Document as DocxDocument
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import re
    import io

    doc = DocxDocument()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Add title
    t = doc.add_heading(title, level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add generation date
    from datetime import datetime, timezone
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run(f"Generated: {datetime.now(timezone.utc).strftime('%B %d, %Y at %H:%M UTC')}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)
    doc.add_paragraph()  # spacer

    # Parse markdown-like content into Word elements
    lines = text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Headings
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        # Bullet points
        elif stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(stripped[2:], style="List Bullet")
        # Numbered list
        elif re.match(r"^\d+[\.\)]\s", stripped):
            text_content = re.sub(r"^\d+[\.\)]\s", "", stripped)
            p = doc.add_paragraph(text_content, style="List Number")
        # Bold line (like **Section Title**)
        elif stripped.startswith("**") and stripped.endswith("**") and len(stripped) > 4:
            p = doc.add_paragraph()
            run = p.add_run(stripped[2:-2])
            run.bold = True
            run.font.size = Pt(12)
        # Horizontal rule
        elif stripped in ("---", "***", "___"):
            doc.add_paragraph("_" * 50)
        # Empty line
        elif not stripped:
            pass  # skip blank lines (natural spacing from paragraphs)
        # Normal paragraph — handle inline bold/italic
        else:
            p = doc.add_paragraph()
            # Split on **bold** and *italic* markers
            parts = re.split(r"(\*\*.*?\*\*|\*.*?\*)", stripped)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith("*") and part.endswith("*") and len(part) > 2:
                    run = p.add_run(part[1:-1])
                    run.italic = True
                else:
                    p.add_run(part)
        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@app.post("/chat/export", dependencies=[Depends(verify_api_key)])
async def export_chat_to_docx(body: ExportChatRequest, request: Request, db=Depends(get_db)):
    """Export a chat session's AI responses as a Word document.

    Collects all assistant messages from the session and formats them
    into a downloadable .docx file.
    """
    session = db.query(DBChatSession).filter(DBChatSession.id == body.session_id).first()
    if not session:
        raise HTTPException(status_code=404, detail="Chat session not found")
    current_user_id = getattr(request.state, "user_id", None)
    if current_user_id and session.user_id and session.user_id != current_user_id:
        raise HTTPException(status_code=403, detail="You do not own this chat session")

    messages = (
        db.query(DBChatMessage)
        .filter(DBChatMessage.session_id == session.id)
        .order_by(DBChatMessage.created_at)
        .all()
    )

    # Build document content from the conversation
    parts = []
    for m in messages:
        content = _decrypt_text(m.content)
        if m.role == "user":
            parts.append(f"**Question:** {content}")
        else:
            parts.append(content)
        parts.append("")  # blank line separator

    full_text = "\n".join(parts)
    title = session.title or "Chat Export"
    docx_bytes = _markdown_to_docx(full_text, title)

    safe_title = re.sub(r'[^\w\s-]', '', title)[:50].strip() or "chat-export"
    filename = f"{safe_title}.docx"

    return Response(
        content=docx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


class MarkdownToDocxRequest(BaseModel):
    markdown: str = Field(max_length=500000)
    title: str = Field(default="Document", max_length=255)


@app.post("/chat/markdown-to-docx", dependencies=[Depends(verify_api_key)])
async def markdown_to_docx(body: MarkdownToDocxRequest):
    """Convert markdown text to a downloadable Word document."""
    docx_bytes = _markdown_to_docx(body.markdown, body.title)
    safe_title = re.sub(r'[^\w\s-]', '', body.title)[:50].strip() or "document"
    filename = f"{safe_title}.docx"
    return Response(
        content=docx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.post("/chat/generate-doc", dependencies=[Depends(verify_api_key)])
async def generate_document_from_chat(body: GenerateDocRequest, request: Request, db=Depends(get_db)):
    """Use AI to generate a Word document based on chat context and instructions.

    The AI writes a complete document (procedure, report, summary, etc.)
    using the uploaded procedures and conversation history as context,
    then returns it as a downloadable .docx file.
    """
    api_key = os.getenv("ANTHROPIC_API_KEY")
    if not api_key:
        raise HTTPException(status_code=500, detail="ANTHROPIC_API_KEY not configured on server")

    # Load session history for context
    session = db.query(DBChatSession).filter(DBChatSession.id == body.session_id).first()
    if not session:
        raise HTTPException(status_code=404, detail="Chat session not found")
    current_user_id = getattr(request.state, "user_id", None)
    if current_user_id and session.user_id and session.user_id != current_user_id:
        raise HTTPException(status_code=403, detail="You do not own this chat session")

    # Get conversation context
    db_messages = (
        db.query(DBChatMessage)
        .filter(DBChatMessage.session_id == session.id)
        .order_by(DBChatMessage.created_at)
        .all()
    )
    chat_context = "\n\n".join(
        f"{'User' if m.role == 'user' else 'Assistant'}: {_decrypt_text(m.content)}"
        for m in db_messages[-10:]
    )

    # Get procedure context from selected documents
    query = db.query(DBDocument)
    if body.doc_ids:
        query = query.filter(DBDocument.id.in_(body.doc_ids))
    documents = query.all()

    procedure_parts = []
    for doc in documents:
        decrypted_name = _decrypt_text(doc.filename)
        pages = json.loads(_decrypt_text(doc.text_content))
        full_text = "\n".join(p["text"] for p in pages if p.get("text"))
        if len(full_text) > 40000:
            full_text = full_text[:40000] + "\n[... truncated ...]"
        procedure_parts.append(f'--- "{decrypted_name}" ---\n{full_text}')
    procedure_context = "\n\n".join(procedure_parts) if procedure_parts else "(No procedures loaded)"

    from anthropic import Anthropic
    client = Anthropic(api_key=api_key)

    system = f"""You are a professional document writer. You create well-structured, detailed documents based on the user's instructions and the reference materials provided.

REFERENCE PROCEDURES:
{procedure_context[:200000]}

RECENT CHAT CONTEXT:
{chat_context[:50000]}

INSTRUCTIONS:
- Write the document in clean, professional language
- Use markdown headings (#, ##, ###), bold (**text**), bullet points (- item), and numbered lists (1. item)
- Include all relevant details from the reference procedures
- Structure the document logically with clear sections
- The document should be complete and ready to use — not a draft or outline"""

    create_kwargs = dict(
        model=CHAT_MODEL,
        max_tokens=AGENT_MAX_TOKENS,
        system=system,
        messages=[{"role": "user", "content": body.instructions}],
    )
    if CHAT_WEB_SEARCH:
        create_kwargs["tools"] = [{"type": "web_search_20250305", "name": "web_search"}]

    response = client.messages.create(**create_kwargs)
    full_text = ""
    for block in response.content:
        if hasattr(block, "text"):
            full_text += block.text

    if not full_text.strip():
        raise HTTPException(status_code=500, detail="AI failed to generate document content")

    docx_bytes = _markdown_to_docx(full_text, body.title)

    safe_title = re.sub(r'[^\w\s-]', '', body.title)[:50].strip() or "generated-document"

    if body.include_vba:
        import zipfile
        zip_buf = io.BytesIO()
        vba_code = _build_vba_module(body.title)
        with zipfile.ZipFile(zip_buf, 'w', zipfile.ZIP_DEFLATED) as zf:
            zf.writestr(f"{safe_title}.docx", docx_bytes)
            zf.writestr(f"{safe_title}_macros.bas", vba_code)
            zf.writestr("README.txt",
                f"DOCUMENT PACKAGE\n"
                f"================\n\n"
                f"1. {safe_title}.docx — The generated document\n"
                f"2. {safe_title}_macros.bas — VBA macros for Word formatting\n\n"
                f"To use macros: Open .docx in Word, press Alt+F11,\n"
                f"File > Import File, select the .bas, then Alt+F8 > FormatProcedure > Run\n"
            )
        return Response(
            content=zip_buf.getvalue(),
            media_type="application/zip",
            headers={"Content-Disposition": f'attachment; filename="{safe_title}_package.zip"'},
        )

    return Response(
        content=docx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{safe_title}.docx"'},
    )


class ImproveProcedureRequest(BaseModel):
    session_id: str = Field(max_length=100, description="Chat session for context")
    procedure_doc_id: str = Field(max_length=100, description="Document ID of the procedure to improve")
    reference_doc_ids: list[str] = Field(default=[], max_length=100, description="Document IDs of reference materials (standards, regs, other procedures)")
    focus_areas: str = Field(default="", max_length=5000, description="Specific areas to improve: clarity, safety steps, compliance, formatting, etc.")
    title: str = Field(default="Improved Procedure", max_length=255)
    include_vba: bool = Field(default=False, description="Include VBA macro code for Word formatting automation")


def _build_vba_module(title: str) -> str:
    """Generate a VBA module for Word that automates procedure formatting."""
    safe_title = title.replace('"', '""')
    return (
        'Attribute VB_Name = "ProcedureFormatter"\n'
        "' =============================================================================\n"
        "' Procedure Formatter VBA Module - Generated by PDFHelper AI\n"
        "'\n"
        "' HOW TO USE:\n"
        "' 1. Open the .docx file in Word\n"
        "' 2. Press Alt+F11 to open the VBA editor\n"
        "' 3. Go to Insert > Module, paste this code (or File > Import this .bas)\n"
        "' 4. Close VBA editor, press Alt+F8, run FormatProcedure\n"
        "' =============================================================================\n"
        "\n"
        "Public Sub FormatProcedure()\n"
        "    ' Master formatter - runs all formatting steps in order.\n"
        "    Application.ScreenUpdating = False\n"
        "    FormatHeadings\n"
        "    FormatStepNumbers\n"
        "    AddCheckboxes\n"
        "    FormatWarningsAndCautions\n"
        "    FormatTables\n"
        "    AddHeaderFooter\n"
        "    SetDocumentProperties\n"
        "    Application.ScreenUpdating = True\n"
        '    MsgBox "Procedure formatting complete!", vbInformation, "Procedure Formatter"\n'
        "End Sub\n"
        "\n"
        "Private Sub FormatHeadings()\n"
        "    ' Apply consistent heading styles throughout the document.\n"
        "    Dim para As Paragraph\n"
        "    For Each para In ActiveDocument.Paragraphs\n"
        "        Dim txt As String\n"
        "        txt = Trim(para.Range.Text)\n"
        '        If UCase(txt) = txt And Len(txt) > 3 And Len(txt) < 80 And Not IsNumeric(Left(txt, 1)) Then\n'
        '            para.Style = ActiveDocument.Styles("Heading 1")\n'
        "        End If\n"
        '        If txt Like "#.0*" Or txt Like "##.0*" Or UCase(Left(txt, 7)) = "SECTION" Then\n'
        '            para.Style = ActiveDocument.Styles("Heading 1")\n'
        "        End If\n"
        '        If txt Like "#.#*" And Not txt Like "#.0*" Then\n'
        '            para.Style = ActiveDocument.Styles("Heading 2")\n'
        "        End If\n"
        '        If txt Like "#.#.#*" Then\n'
        '            para.Style = ActiveDocument.Styles("Heading 3")\n'
        "        End If\n"
        "    Next para\n"
        "End Sub\n"
        "\n"
        "Private Sub FormatStepNumbers()\n"
        "    ' Format step patterns (Step 1, Step 2, etc.) consistently.\n"
        "    Dim para As Paragraph\n"
        "    For Each para In ActiveDocument.Paragraphs\n"
        "        Dim txt As String\n"
        "        txt = Trim(para.Range.Text)\n"
        '        If UCase(Left(txt, 4)) = "STEP" Or txt Like "Step #*" Or txt Like "Step ##*" Then\n'
        "            para.Range.Font.Bold = True\n"
        "            para.Range.Font.Size = 11\n"
        "            para.Range.Font.Color = RGB(0, 51, 102)\n"
        "        End If\n"
        "    Next para\n"
        "End Sub\n"
        "\n"
        "Private Sub AddCheckboxes()\n"
        "    ' Add checkbox characters before verification/checklist items.\n"
        "    Dim para As Paragraph\n"
        "    For Each para In ActiveDocument.Paragraphs\n"
        "        Dim txt As String\n"
        "        txt = Trim(para.Range.Text)\n"
        '        If Left(txt, 3) = "[ ]" Then\n'
        '            para.Range.Text = ChrW(&H2610) & " " & Mid(txt, 4)\n'
        '        ElseIf UCase(Left(txt, 6)) = "VERIFY" Or UCase(Left(txt, 7)) = "CONFIRM" _\n'
        '               Or UCase(Left(txt, 5)) = "CHECK" Or UCase(Left(txt, 6)) = "ENSURE" Then\n'
        '            para.Range.InsertBefore ChrW(&H2610) & " "\n'
        "        End If\n"
        "    Next para\n"
        "End Sub\n"
        "\n"
        "Private Sub FormatWarningsAndCautions()\n"
        "    ' Highlight WARNING, CAUTION, DANGER, and NOTE paragraphs.\n"
        "    Dim para As Paragraph\n"
        "    For Each para In ActiveDocument.Paragraphs\n"
        "        Dim txt As String\n"
        "        txt = UCase(Trim(para.Range.Text))\n"
        '        If Left(txt, 7) = "WARNING" Or Left(txt, 6) = "DANGER" Then\n'
        "            para.Range.Font.Bold = True\n"
        "            para.Range.Font.Color = RGB(204, 0, 0)\n"
        "            para.Shading.BackgroundPatternColor = RGB(255, 235, 235)\n"
        '        ElseIf Left(txt, 7) = "CAUTION" Then\n'
        "            para.Range.Font.Bold = True\n"
        "            para.Range.Font.Color = RGB(204, 102, 0)\n"
        "            para.Shading.BackgroundPatternColor = RGB(255, 248, 230)\n"
        '        ElseIf Left(txt, 4) = "NOTE" Then\n'
        "            para.Range.Font.Italic = True\n"
        "            para.Range.Font.Color = RGB(0, 51, 153)\n"
        "            para.Shading.BackgroundPatternColor = RGB(235, 243, 255)\n"
        "        End If\n"
        "    Next para\n"
        "End Sub\n"
        "\n"
        "Private Sub FormatTables()\n"
        "    ' Apply consistent formatting to all tables.\n"
        "    Dim tbl As Table\n"
        "    For Each tbl In ActiveDocument.Tables\n"
        "        tbl.Borders.Enable = True\n"
        "        tbl.Borders.InsideLineStyle = wdLineStyleSingle\n"
        "        tbl.Borders.OutsideLineStyle = wdLineStyleSingle\n"
        "        If tbl.Rows.Count > 0 Then\n"
        "            tbl.Rows(1).Range.Font.Bold = True\n"
        "            tbl.Rows(1).Shading.BackgroundPatternColor = RGB(0, 51, 102)\n"
        "            tbl.Rows(1).Range.Font.Color = RGB(255, 255, 255)\n"
        "        End If\n"
        "        Dim i As Long\n"
        "        For i = 2 To tbl.Rows.Count\n"
        "            If i Mod 2 = 0 Then\n"
        "                tbl.Rows(i).Shading.BackgroundPatternColor = RGB(242, 246, 250)\n"
        "            End If\n"
        "        Next i\n"
        "        tbl.AutoFitBehavior wdAutoFitWindow\n"
        "    Next tbl\n"
        "End Sub\n"
        "\n"
        "Private Sub AddHeaderFooter()\n"
        "    ' Add professional header and footer.\n"
        "    Dim sec As Section\n"
        "    For Each sec In ActiveDocument.Sections\n"
        "        sec.Headers(wdHeaderFooterPrimary).Range.Text = _\n"
        f'            "{safe_title}" & vbTab & vbTab & "CONTROLLED DOCUMENT"\n'
        "        sec.Headers(wdHeaderFooterPrimary).Range.Font.Size = 9\n"
        "        sec.Headers(wdHeaderFooterPrimary).Range.Font.Color = RGB(128, 128, 128)\n"
        "        sec.Footers(wdHeaderFooterPrimary).Range.Text = _\n"
        '            "Page " & vbTab & vbTab & "Revision Date: " & Format(Date, "yyyy-mm-dd")\n'
        "        sec.Footers(wdHeaderFooterPrimary).Range.Font.Size = 9\n"
        "        sec.Footers(wdHeaderFooterPrimary).Range.Font.Color = RGB(128, 128, 128)\n"
        "        Dim rng As Range\n"
        "        Set rng = sec.Footers(wdHeaderFooterPrimary).Range\n"
        "        rng.Collapse Direction:=wdCollapseStart\n"
        "        rng.MoveEnd Unit:=wdCharacter, Count:=5\n"
        "        rng.Collapse Direction:=wdCollapseEnd\n"
        "        ActiveDocument.Fields.Add Range:=rng, Type:=wdFieldPage\n"
        '        rng.InsertAfter " of "\n'
        "        rng.Collapse Direction:=wdCollapseEnd\n"
        "        ActiveDocument.Fields.Add Range:=rng, Type:=wdFieldNumPages\n"
        "    Next sec\n"
        "End Sub\n"
        "\n"
        "Private Sub SetDocumentProperties()\n"
        "    ' Set document metadata.\n"
        "    With ActiveDocument.BuiltInDocumentProperties\n"
        f'        .Item("Title").Value = "{safe_title}"\n'
        '        .Item("Subject").Value = "Operating Procedure"\n'
        '        .Item("Category").Value = "Procedure Document"\n'
        "    End With\n"
        '    ActiveDocument.Styles("Normal").Font.Name = "Calibri"\n'
        '    ActiveDocument.Styles("Normal").Font.Size = 11\n'
        "End Sub\n"
        "\n"
        "Public Sub InsertRevisionTable()\n"
        "    ' Insert a revision history table at the cursor.\n"
        "    Dim tbl As Table\n"
        "    Set tbl = ActiveDocument.Tables.Add( _\n"
        "        Range:=Selection.Range, NumRows:=4, NumColumns:=4)\n"
        '    tbl.Cell(1, 1).Range.Text = "Rev"\n'
        '    tbl.Cell(1, 2).Range.Text = "Date"\n'
        '    tbl.Cell(1, 3).Range.Text = "Description"\n'
        '    tbl.Cell(1, 4).Range.Text = "Author"\n'
        '    tbl.Cell(2, 1).Range.Text = "0"\n'
        '    tbl.Cell(2, 2).Range.Text = Format(Date, "yyyy-mm-dd")\n'
        '    tbl.Cell(2, 3).Range.Text = "Initial release - AI-generated from source documents"\n'
        '    tbl.Cell(2, 4).Range.Text = ""\n'
        "    tbl.Rows(1).Range.Font.Bold = True\n"
        "    tbl.Rows(1).Shading.BackgroundPatternColor = RGB(0, 51, 102)\n"
        "    tbl.Rows(1).Range.Font.Color = RGB(255, 255, 255)\n"
        "    tbl.Borders.Enable = True\n"
        "    tbl.AutoFitBehavior wdAutoFitWindow\n"
        "End Sub\n"
        "\n"
        "Public Sub InsertSignOffBlock()\n"
        "    ' Insert a signature/approval block at the cursor.\n"
        "    Dim tbl As Table\n"
        "    Set tbl = ActiveDocument.Tables.Add( _\n"
        "        Range:=Selection.Range, NumRows:=4, NumColumns:=3)\n"
        '    tbl.Cell(1, 1).Range.Text = "Role"\n'
        '    tbl.Cell(1, 2).Range.Text = "Name / Signature"\n'
        '    tbl.Cell(1, 3).Range.Text = "Date"\n'
        '    tbl.Cell(2, 1).Range.Text = "Prepared By"\n'
        '    tbl.Cell(3, 1).Range.Text = "Reviewed By"\n'
        '    tbl.Cell(4, 1).Range.Text = "Approved By"\n'
        "    tbl.Rows(1).Range.Font.Bold = True\n"
        "    tbl.Rows(1).Shading.BackgroundPatternColor = RGB(0, 51, 102)\n"
        "    tbl.Rows(1).Range.Font.Color = RGB(255, 255, 255)\n"
        "    tbl.Borders.Enable = True\n"
        "    Dim i As Long\n"
        "    For i = 2 To 4\n"
        "        tbl.Rows(i).Height = CentimetersToPoints(1.5)\n"
        "    Next i\n"
        "    tbl.AutoFitBehavior wdAutoFitWindow\n"
        "End Sub\n"
    )


@app.post("/chat/improve-procedure", dependencies=[Depends(verify_api_key)])
async def improve_procedure(body: ImproveProcedureRequest, request: Request, db=Depends(get_db)):
    """Cross-reference uploaded documents to produce an improved procedure.

    Takes one procedure document as the base, compares it against reference
    documents (standards, regulations, other procedures), and generates an
    improved version as a Word document. Optionally includes VBA macro code
    for automated formatting in Word.
    """
    api_key = os.getenv("ANTHROPIC_API_KEY")
    if not api_key:
        raise HTTPException(status_code=500, detail="ANTHROPIC_API_KEY not configured on server")

    # Load the procedure to improve
    proc_doc = db.query(DBDocument).filter(DBDocument.id == body.procedure_doc_id).first()
    if not proc_doc:
        raise HTTPException(status_code=404, detail="Procedure document not found")
    proc_name = _decrypt_text(proc_doc.filename)
    proc_pages = json.loads(_decrypt_text(proc_doc.text_content))
    proc_text = "\n".join(p["text"] for p in proc_pages if p.get("text"))
    if len(proc_text) > 100000:
        proc_text = proc_text[:100000] + "\n[... truncated ...]"

    # Load reference documents
    ref_parts = []
    if body.reference_doc_ids:
        ref_docs = db.query(DBDocument).filter(DBDocument.id.in_(body.reference_doc_ids)).all()
        for doc in ref_docs:
            name = _decrypt_text(doc.filename)
            pages = json.loads(_decrypt_text(doc.text_content))
            text = "\n".join(p["text"] for p in pages if p.get("text"))
            if len(text) > 60000:
                text = text[:60000] + "\n[... truncated ...]"
            ref_parts.append(f'--- REFERENCE: "{name}" ---\n{text}\n--- END ---')
    ref_context = "\n\n".join(ref_parts) if ref_parts else "(No additional reference documents)"

    # Load chat context
    session = db.query(DBChatSession).filter(DBChatSession.id == body.session_id).first()
    chat_context = ""
    if session:
        current_user_id = getattr(request.state, "user_id", None)
        if current_user_id and session.user_id and session.user_id != current_user_id:
            raise HTTPException(status_code=403, detail="You do not own this chat session")
        db_messages = (
            db.query(DBChatMessage)
            .filter(DBChatMessage.session_id == session.id)
            .order_by(DBChatMessage.created_at)
            .all()
        )
        chat_context = "\n\n".join(
            f"{'User' if m.role == 'user' else 'Assistant'}: {_decrypt_text(m.content)}"
            for m in db_messages[-10:]
        )

    focus = body.focus_areas or "clarity, completeness, safety, step-by-step structure, compliance"

    from anthropic import Anthropic
    client = Anthropic(api_key=api_key)

    system = f"""You are an expert technical procedure writer. Your task is to IMPROVE an existing procedure document by cross-referencing it against other reference documents.

=== ORIGINAL PROCEDURE TO IMPROVE ===
Document: "{proc_name}"
{proc_text}

=== REFERENCE DOCUMENTS ===
{ref_context[:200000]}

=== RECENT DISCUSSION ===
{chat_context[:30000]}

=== YOUR TASK ===
Produce an IMPROVED version of the original procedure. Focus on: {focus}

RULES:
1. Cross-reference the original procedure against ALL reference documents
2. Incorporate missing steps, safety requirements, compliance items found in references
3. Keep the improved procedure clear, concise, and actionable
4. Use proper procedure structure:
   - Title and document number
   - Purpose / Scope
   - Definitions / Abbreviations
   - Responsibilities
   - Required Tools / Materials / PPE
   - Precautions (DANGER / WARNING / CAUTION / NOTE)
   - Step-by-step instructions (numbered, with verification checkboxes)
   - Acceptance criteria
   - References
   - Revision history placeholder
5. Highlight what changed from the original with [ADDED], [MODIFIED], or [IMPROVED] tags
6. Use markdown: # headings, **bold**, - bullets, 1. numbered lists
7. Be specific — include actual values, part numbers, limits from the reference documents
8. Every safety-critical step should have a verification checkbox [ ]"""

    create_kwargs = dict(
        model=CHAT_MODEL,
        max_tokens=AGENT_MAX_TOKENS,
        system=system,
        messages=[{"role": "user", "content": f"Improve this procedure. Focus areas: {focus}\n\nProduce the complete improved procedure now."}],
    )
    if CHAT_WEB_SEARCH:
        create_kwargs["tools"] = [{"type": "web_search_20250305"}]

    response = client.messages.create(**create_kwargs)
    full_text = ""
    for block in response.content:
        if hasattr(block, "text"):
            full_text += block.text

    if not full_text.strip():
        raise HTTPException(status_code=500, detail="AI failed to generate the improved procedure")

    docx_bytes = _markdown_to_docx(full_text, body.title)

    # Build response files
    safe_title = re.sub(r'[^\w\s-]', '', body.title)[:50].strip() or "improved-procedure"

    if body.include_vba:
        # Return a ZIP containing the .docx and the .bas VBA module
        import zipfile
        zip_buf = io.BytesIO()
        vba_code = _build_vba_module(body.title)
        with zipfile.ZipFile(zip_buf, 'w', zipfile.ZIP_DEFLATED) as zf:
            zf.writestr(f"{safe_title}.docx", docx_bytes)
            zf.writestr(f"{safe_title}_macros.bas", vba_code)
            zf.writestr("README.txt",
                f"IMPROVED PROCEDURE PACKAGE\n"
                f"=========================\n\n"
                f"This package contains:\n\n"
                f"1. {safe_title}.docx\n"
                f"   The improved procedure document. Open in Word to review and edit.\n\n"
                f"2. {safe_title}_macros.bas\n"
                f"   VBA macro module for Word. To use:\n"
                f"   a) Open the .docx in Word\n"
                f"   b) Press Alt+F11 to open the VBA editor\n"
                f"   c) Go to File > Import File, select the .bas file\n"
                f"   d) Close VBA editor\n"
                f"   e) Press Alt+F8, select 'FormatProcedure', click Run\n\n"
                f"   Available macros:\n"
                f"   - FormatProcedure: Applies all formatting at once\n"
                f"   - InsertRevisionTable: Adds a revision history table\n"
                f"   - InsertSignOffBlock: Adds a signature/approval block\n\n"
                f"3. Save as .docm (macro-enabled) to keep the VBA macros.\n"
            )
        return Response(
            content=zip_buf.getvalue(),
            media_type="application/zip",
            headers={"Content-Disposition": f'attachment; filename="{safe_title}_package.zip"'},
        )

    return Response(
        content=docx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{safe_title}.docx"'},
    )


# ---------------------------------------------------------------------------
# Full Analysis Pipeline (multi-agent)
# ---------------------------------------------------------------------------

@app.post("/analyze", dependencies=[Depends(verify_api_key)])
async def analyze_documents(
    request: Request,
    body: AnalyzeRequest,
    doc_ids: list[str] = Query(default=[], description="Document IDs to analyze (empty = all)"),
    db=Depends(get_db),
):
    """Run the full multi-agent analysis pipeline on uploaded documents.

    This runs 4 specialized AI agents:
    1. Document Analyzer — deep analysis of each document
    2. Cross-Reference Checker — finds conflicts between documents
    3. Compliance Checker — flags regulatory/policy issues
    4. Summary Report Generator — produces an actionable executive report

    Optionally also runs keyword and AI search.
    """
    query = db.query(DBDocument)
    if doc_ids:
        query = query.filter(DBDocument.id.in_(doc_ids))
    documents = query.all()

    if not documents:
        raise HTTPException(status_code=404, detail="No documents found")

    client_ip = _get_client_ip(request)

    # Build a cache key from document content hashes + analysis parameters
    # This lets us skip re-analysis when the same documents are analyzed
    # with the same compliance context (search is fast enough to always re-run)
    doc_hashes = sorted(d.content_hash or d.id for d in documents)
    cache_key_input = json.dumps({
        "hashes": doc_hashes,
        "compliance_context": body.compliance_context,
    }, sort_keys=True)
    cache_key = hashlib.sha256(cache_key_input.encode()).hexdigest()

    # Check for a cached analysis with the same content + parameters
    cached_report = (
        db.query(DBAnalysisReport)
        .filter(DBAnalysisReport.cache_key == cache_key)
        .order_by(DBAnalysisReport.analyzed_at.desc())
        .first()
    )
    if cached_report:
        cached_analysis = json.loads(_decrypt_text(cached_report.report_data))
        # Re-run search if requested (cheap), but reuse the cached analysis
        if body.search_terms or body.ai_query:
            docs_for_agents: dict[str, list[dict]] = {}
            for doc in documents:
                decrypted_name = _decrypt_text(doc.filename)
                pages = json.loads(_decrypt_text(doc.text_content))
                docs_for_agents[decrypted_name] = pages
            from search import keyword_search, ai_search
            search_results = {"keyword_results": [], "ai_results": []}
            for filename, pages in docs_for_agents.items():
                if body.search_terms:
                    kw_matches = keyword_search(pages, body.search_terms)
                    for m in kw_matches:
                        m["filename"] = filename
                    search_results["keyword_results"].extend(kw_matches)
                if body.ai_query:
                    import asyncio
                    ai_matches = await asyncio.to_thread(ai_search, pages, body.ai_query, filename)
                    for m in ai_matches:
                        m["filename"] = filename
                    search_results["ai_results"].extend(ai_matches)
            cached_analysis["search_results"] = search_results

        return {
            "report_id": cached_report.id,
            "cached": True,
            "report": cached_analysis.get("report"),
            "document_analyses": cached_analysis.get("document_analyses"),
            "cross_reference_findings": cached_analysis.get("cross_reference_findings"),
            "compliance_findings": cached_analysis.get("compliance_findings"),
            "search_results": cached_analysis.get("search_results"),
        }

    # Build documents dict for the agent pipeline
    docs_for_agents: dict[str, list[dict]] = {}
    for doc in documents:
        decrypted_name = _decrypt_text(doc.filename)
        pages = json.loads(_decrypt_text(doc.text_content))
        docs_for_agents[decrypted_name] = pages

    # Run the full pipeline in a thread pool to avoid blocking the event loop
    # (run_full_analysis makes multiple synchronous Anthropic API calls)
    import asyncio
    from agents import run_full_analysis
    analysis = await asyncio.to_thread(
        run_full_analysis,
        documents=docs_for_agents,
        compliance_context=body.compliance_context,
        search_terms=body.search_terms if body.search_terms else None,
        ai_query=body.ai_query,
    )

    # Save to DB
    report_id = str(uuid.uuid4())
    db_report = DBAnalysisReport(
        id=report_id,
        doc_ids=json.dumps([d.id for d in documents]),
        compliance_context=_encrypt_text(body.compliance_context) if body.compliance_context else None,
        report_data=_encrypt_text(json.dumps(analysis)),
        documents_analyzed=len(documents),
        total_issues=analysis.get("report", {}).get("total_issues_found", 0),
        critical_issues=analysis.get("report", {}).get("critical_issues", 0),
        risk_level=analysis.get("report", {}).get("overall_risk_level", "unknown"),
        cache_key=cache_key,
        analyzed_at=datetime.now(timezone.utc),
    )
    db.add(db_report)
    db.commit()

    log_search(client_ip, report_id, body.search_terms, body.ai_query,
               len(documents), db_report.total_issues, db_report.critical_issues)

    response = {
        "report_id": report_id,
        "cached": False,
        "report": analysis.get("report"),
        "document_analyses": analysis.get("document_analyses"),
        "cross_reference_findings": analysis.get("cross_reference_findings"),
        "compliance_findings": analysis.get("compliance_findings"),
        "search_results": analysis.get("search_results"),
    }
    if analysis.get("warnings"):
        response["warnings"] = analysis["warnings"]
    return response


@app.get("/reports", dependencies=[Depends(verify_api_key)])
async def list_reports(limit: int = Query(default=20, le=100), db=Depends(get_db)):
    """List past analysis reports."""
    reports = (
        db.query(DBAnalysisReport)
        .order_by(DBAnalysisReport.analyzed_at.desc())
        .limit(limit)
        .all()
    )
    return {
        "reports": [
            {
                "id": r.id,
                "documents_analyzed": r.documents_analyzed,
                "total_issues": r.total_issues,
                "critical_issues": r.critical_issues,
                "risk_level": r.risk_level,
                "analyzed_at": r.analyzed_at.isoformat(),
            }
            for r in reports
        ]
    }


@app.get("/reports/{report_id}", dependencies=[Depends(verify_api_key)])
async def get_report(report_id: str, db=Depends(get_db)):
    """Get full details of a past analysis report."""
    report = db.query(DBAnalysisReport).filter(DBAnalysisReport.id == report_id).first()
    if not report:
        raise HTTPException(status_code=404, detail="Report not found")
    full_data = json.loads(_decrypt_text(report.report_data))
    return {
        "id": report.id,
        "documents_analyzed": report.documents_analyzed,
        "total_issues": report.total_issues,
        "critical_issues": report.critical_issues,
        "risk_level": report.risk_level,
        "analyzed_at": report.analyzed_at.isoformat(),
        **full_data,
    }


# ---------------------------------------------------------------------------
# IsoIntel — P&ID Drawing Management
# ---------------------------------------------------------------------------

@app.post("/drawings/upload", dependencies=[Depends(verify_api_key)])
async def upload_drawing(
    request: Request,
    file: UploadFile = File(...),
    title: str = Query(default="", description="Drawing title"),
    drawing_number: str = Query(default="", description="Drawing number, e.g. HEB-PID-1234"),
    equipment_tags: str = Query(default="", description="Comma-separated equipment tags on this drawing"),
    description: str = Query(default="", description="What system this drawing covers"),
    db=Depends(get_db),
):
    """Upload a P&ID drawing (PDF or image) for use in isolation packages."""
    content = await file.read()

    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(status_code=413, detail=f"File too large. Max {MAX_FILE_SIZE // (1024*1024)}MB.")

    drawing_id = str(uuid.uuid4())
    now = datetime.now(timezone.utc)

    # Determine file type and process
    fname = (file.filename or "drawing.pdf").strip()
    ext = fname.rsplit(".", 1)[-1].lower() if "." in fname else ""

    # Save encrypted file
    safe_name = _sanitize_filename(fname) if fname else "drawing"
    dest = UPLOAD_DIR / f"{drawing_id}_{safe_name}"
    if ENCRYPTION_KEY:
        from encryption import encrypt_bytes
        dest.write_bytes(encrypt_bytes(content))
    else:
        dest.write_bytes(content)

    # Extract text — convert images to PDF first for OCR
    text_content = "[]"
    page_count = 1
    try:
        if ext == "pdf":
            pages = extract_text_from_bytes(content)
        elif ext in ("jpg", "jpeg", "png", "tiff", "tif", "bmp", "webp"):
            pdf_bytes = _image_to_pdf(content)
            pages = extract_text_from_bytes(pdf_bytes)
        else:
            pages = []
        if pages:
            text_content = json.dumps(pages)
            page_count = len(pages)
    except Exception:
        pass

    drawing = DBDrawing(
        id=drawing_id,
        filename=_encrypt_text(fname),
        filepath=_encrypt_text(str(dest)),
        title=_encrypt_text(title) if title else _encrypt_text(fname),
        drawing_number=_encrypt_text(drawing_number) if drawing_number else None,
        equipment_tags=_encrypt_text(equipment_tags) if equipment_tags else None,
        description=_encrypt_text(description) if description else None,
        page_count=page_count,
        text_content=_encrypt_text(text_content),
        uploaded_at=now,
    )
    db.add(drawing)
    db.commit()

    log_upload(_get_client_ip(request), fname, drawing_id, 1)

    return {
        "id": drawing_id,
        "filename": fname,
        "title": title or fname,
        "drawing_number": drawing_number,
        "equipment_tags": equipment_tags,
        "page_count": page_count,
        "uploaded_at": now.isoformat(),
    }


@app.post("/drawings/upload-batch", dependencies=[Depends(verify_api_key)])
async def upload_drawings_batch(
    request: Request,
    files: list[UploadFile] = File(...),
    equipment_tags: str = Query(default="", description="Comma-separated equipment tags shared across all drawings"),
    description: str = Query(default="", description="Shared description for these drawings"),
    db=Depends(get_db),
):
    """Upload multiple P&ID drawings at once. Each file gets its filename as the title and drawing number."""
    if len(files) > MAX_FILES_PER_REQUEST:
        raise HTTPException(status_code=400, detail=f"Max {MAX_FILES_PER_REQUEST} files per request.")

    uploaded = []
    errors = []

    for file in files:
        try:
            content = await file.read()
            if len(content) > MAX_FILE_SIZE:
                errors.append({"filename": file.filename, "error": f"File too large. Max {MAX_FILE_SIZE // (1024*1024)}MB."})
                continue

            drawing_id = str(uuid.uuid4())
            now = datetime.now(timezone.utc)

            fname = (file.filename or "drawing.pdf").strip()
            ext = fname.rsplit(".", 1)[-1].lower() if "." in fname else ""

            # Save encrypted file
            safe_name = _sanitize_filename(fname) if fname else "drawing"
            dest = UPLOAD_DIR / f"{drawing_id}_{safe_name}"
            if ENCRYPTION_KEY:
                from encryption import encrypt_bytes
                dest.write_bytes(encrypt_bytes(content))
            else:
                dest.write_bytes(content)

            # Extract text if PDF
            text_content = "[]"
            page_count = 1
            if ext == "pdf":
                try:
                    pages = extract_text_from_bytes(content)
                    text_content = json.dumps(pages)
                    page_count = len(pages)
                except Exception:
                    pass

            # Use filename (without extension) as default title
            default_title = fname.rsplit(".", 1)[0] if "." in fname else fname

            drawing = DBDrawing(
                id=drawing_id,
                filename=_encrypt_text(fname),
                filepath=_encrypt_text(str(dest)),
                title=_encrypt_text(default_title),
                drawing_number=_encrypt_text(default_title),
                equipment_tags=_encrypt_text(equipment_tags) if equipment_tags else None,
                description=_encrypt_text(description) if description else None,
                page_count=page_count,
                text_content=_encrypt_text(text_content),
                uploaded_at=now,
            )
            db.add(drawing)
            db.commit()

            log_upload(_get_client_ip(request), fname, drawing_id, 1)

            uploaded.append({
                "id": drawing_id,
                "filename": fname,
                "title": default_title,
                "page_count": page_count,
                "uploaded_at": now.isoformat(),
            })
        except Exception as e:
            errors.append({"filename": file.filename or "unknown", "error": str(e)})

    return {
        "uploaded": uploaded,
        "errors": errors,
        "total_uploaded": len(uploaded),
        "total_errors": len(errors),
    }


@app.get("/drawings", dependencies=[Depends(verify_api_key)])
async def list_drawings(db=Depends(get_db)):
    """List all uploaded P&ID drawings."""
    drawings = db.query(DBDrawing).order_by(DBDrawing.uploaded_at.desc()).all()
    return {
        "drawings": [
            {
                "id": d.id,
                "filename": _decrypt_text(d.filename),
                "title": _decrypt_text(d.title) if d.title else None,
                "drawing_number": _decrypt_text(d.drawing_number) if d.drawing_number else None,
                "equipment_tags": _decrypt_text(d.equipment_tags) if d.equipment_tags else None,
                "description": _decrypt_text(d.description) if d.description else None,
                "page_count": d.page_count,
                "uploaded_at": d.uploaded_at.isoformat(),
            }
            for d in drawings
        ]
    }


@app.delete("/drawings/{drawing_id}", dependencies=[Depends(verify_api_key)])
async def delete_drawing(drawing_id: str, db=Depends(get_db)):
    """Delete a P&ID drawing."""
    drawing = db.query(DBDrawing).filter(DBDrawing.id == drawing_id).first()
    if not drawing:
        raise HTTPException(status_code=404, detail="Drawing not found")
    # Delete file
    try:
        _safe_unlink(Path(_decrypt_text(drawing.filepath)))
    except Exception:
        pass
    db.delete(drawing)
    db.commit()
    return {"deleted": drawing_id}


# ---------------------------------------------------------------------------
# IsoIntel — Isolation Package Generation
# ---------------------------------------------------------------------------

@app.post("/isolations/generate", dependencies=[Depends(verify_api_key)])
async def generate_isolation(
    request: Request,
    body: IsolationRequest,
    db=Depends(get_db),
):
    """Generate a full isolation package using the two-pass AI pipeline.

    Pass 1 (if no drawing_ids provided): AI selects relevant drawings from library.
    Pass 2: AI reads P&ID images and generates the complete isolation package.
    Response is streamed as Server-Sent Events.
    """
    api_key = os.getenv("ANTHROPIC_API_KEY")
    if not api_key:
        raise HTTPException(status_code=500, detail="ANTHROPIC_API_KEY not configured on server")

    from anthropic import Anthropic
    client = Anthropic(api_key=api_key)

    # Load all drawings for metadata
    all_drawings = db.query(DBDrawing).all()
    if not all_drawings:
        raise HTTPException(status_code=404, detail="No P&ID drawings uploaded. Upload drawings first.")

    now = datetime.now(timezone.utc)
    iso_id = str(uuid.uuid4())
    cert_number = f"ISO-{now.strftime('%Y%m%d')}-{iso_id[:8].upper()}"

    job = {
        "equipment_tag": body.equipment_tag,
        "work_description": body.work_description,
        "work_type": body.work_type,
        "fluid_service": body.fluid_service,
        "special_requirements": body.special_requirements,
    }

    # Determine which drawings to use
    selected_ids = list(body.drawing_ids)

    async def stream_isolation():
        """Stream the isolation generation as SSE."""
        import base64
        import logging
        logger = logging.getLogger("pdfhelper")
        from isointel import run_pass1, run_pass2_stream

        nonlocal selected_ids
        full_output = ""

        try:
            # Pass 1: drawing discovery (if not manually specified)
            if not selected_ids:
                drawings_meta = []
                for d in all_drawings:
                    drawings_meta.append({
                        "id": d.id,
                        "title": _decrypt_text(d.title) if d.title else None,
                        "drawingNumber": _decrypt_text(d.drawing_number) if d.drawing_number else None,
                        "equipmentTags": _decrypt_text(d.equipment_tags) if d.equipment_tags else None,
                        "description": _decrypt_text(d.description) if d.description else None,
                    })

                yield f"data: {json.dumps({'type': 'status', 'message': 'Pass 1: Searching drawing library...'})}\n\n"
                selected_ids = run_pass1(client, drawings_meta, job)

                if not selected_ids:
                    yield f"data: {json.dumps({'type': 'error', 'detail': 'Pass 1 could not identify relevant drawings. Try specifying drawing IDs manually.'})}\n\n"
                    return

                selected_titles = []
                for d in all_drawings:
                    if d.id in selected_ids:
                        selected_titles.append(_decrypt_text(d.title) if d.title else d.id)
                yield f"data: {json.dumps({'type': 'pass1_result', 'drawing_ids': selected_ids, 'drawing_titles': selected_titles})}\n\n"

            # Load drawing images for Pass 2
            yield f"data: {json.dumps({'type': 'status', 'message': 'Pass 2: Reading P&ID images and generating isolation package...'})}\n\n"

            drawing_images = []
            for d in all_drawings:
                if d.id not in selected_ids:
                    continue
                try:
                    fpath = _validate_filepath(Path(_decrypt_text(d.filepath)))
                    if ENCRYPTION_KEY:
                        from encryption import decrypt_bytes as dec_bytes
                        raw = dec_bytes(fpath.read_bytes())
                    else:
                        raw = fpath.read_bytes()

                    fname = _decrypt_text(d.filename).lower()
                    if fname.endswith(".pdf"):
                        # Convert first page of PDF to image
                        try:
                            import fitz  # PyMuPDF
                            pdf_doc = fitz.open(stream=raw, filetype="pdf")
                            for page_num in range(min(pdf_doc.page_count, 2)):
                                page = pdf_doc[page_num]
                                pix = page.get_pixmap(dpi=200)
                                img_bytes = pix.tobytes("png")
                                drawing_images.append({
                                    "id": d.id,
                                    "title": _decrypt_text(d.title) if d.title else d.id,
                                    "image_b64": base64.b64encode(img_bytes).decode("ascii"),
                                    "media_type": "image/png",
                                })
                            pdf_doc.close()
                        except ImportError:
                            logger.warning("PyMuPDF not available for PDF-to-image conversion")
                    else:
                        # Image file — use directly
                        media = "image/png"
                        if fname.endswith(".jpg") or fname.endswith(".jpeg"):
                            media = "image/jpeg"
                        elif fname.endswith(".webp"):
                            media = "image/webp"
                        drawing_images.append({
                            "id": d.id,
                            "title": _decrypt_text(d.title) if d.title else d.id,
                            "image_b64": base64.b64encode(raw).decode("ascii"),
                            "media_type": media,
                        })
                except Exception as e:
                    logger.warning("Failed to load drawing %s: %s", d.id, e)

            if not drawing_images:
                yield f"data: {json.dumps({'type': 'error', 'detail': 'Could not load any drawing images. Ensure drawings are uploaded as PDF or image files.'})}\n\n"
                return

            # Send meta
            yield f"data: {json.dumps({'type': 'meta', 'isolation_id': iso_id, 'cert_number': cert_number, 'drawings_used': len(drawing_images)})}\n\n"

            # Pass 2: stream the isolation generation
            for chunk in run_pass2_stream(
                client, job, drawing_images,
                facility=body.facility,
                regime=body.regime,
                cert_number=cert_number,
            ):
                full_output += chunk
                yield f"data: {json.dumps({'type': 'chunk', 'text': chunk})}\n\n"

            yield f"data: {json.dumps({'type': 'done'})}\n\n"

        except Exception as e:
            import logging
            logging.getLogger("pdfhelper").error("Isolation generation failed: %s", e)
            err_msg = "Isolation generation failed" if IS_PRODUCTION else f"Isolation generation failed: {str(e)}"
            yield f"data: {json.dumps({'type': 'error', 'detail': err_msg})}\n\n"
        finally:
            # Save the isolation package to DB
            if full_output:
                save_db = SessionLocal()
                try:
                    # Try to parse stats from the JSON output
                    hazard = "HIGH"
                    v_count = b_count = s_count = e_count = 0
                    try:
                        parsed = json.loads(full_output)
                        stats = parsed.get("stats", {})
                        v_count = stats.get("valveCount", 0)
                        b_count = stats.get("blindCount", 0)
                        s_count = stats.get("stepCount", 0)
                        e_count = stats.get("energySourceCount", 0)
                        hazard = parsed.get("hazardClassification", "HIGH")
                    except (json.JSONDecodeError, AttributeError):
                        pass

                    pkg = DBIsolationPackage(
                        id=iso_id,
                        cert_number=_encrypt_text(cert_number),
                        equipment_tag=_encrypt_text(body.equipment_tag),
                        work_description=_encrypt_text(body.work_description),
                        work_type=_encrypt_text(body.work_type),
                        fluid_service=_encrypt_text(body.fluid_service) if body.fluid_service else None,
                        facility=_encrypt_text(body.facility) if body.facility else None,
                        regime=_encrypt_text(body.regime) if body.regime else None,
                        special_requirements=_encrypt_text(body.special_requirements) if body.special_requirements else None,
                        drawing_ids=json.dumps(selected_ids),
                        package_data=_encrypt_text(full_output),
                        hazard_classification=hazard,
                        valve_count=v_count,
                        blind_count=b_count,
                        step_count=s_count,
                        energy_source_count=e_count,
                        status="draft",
                        created_at=now,
                        updated_at=now,
                    )
                    save_db.add(pkg)
                    save_db.commit()
                except Exception:
                    save_db.rollback()
                finally:
                    save_db.close()

    return StreamingResponse(stream_isolation(), media_type="text/event-stream")


@app.get("/isolations", dependencies=[Depends(verify_api_key)])
async def list_isolations(
    limit: int = Query(default=30, le=100),
    db=Depends(get_db),
):
    """List past isolation packages, most recent first."""
    packages = (
        db.query(DBIsolationPackage)
        .order_by(DBIsolationPackage.created_at.desc())
        .limit(limit)
        .all()
    )
    return {
        "isolations": [
            {
                "id": p.id,
                "cert_number": _decrypt_text(p.cert_number),
                "equipment_tag": _decrypt_text(p.equipment_tag),
                "work_description": _decrypt_text(p.work_description),
                "work_type": _decrypt_text(p.work_type),
                "hazard_classification": p.hazard_classification,
                "valve_count": p.valve_count,
                "blind_count": p.blind_count,
                "step_count": p.step_count,
                "status": p.status,
                "created_at": p.created_at.isoformat(),
            }
            for p in packages
        ]
    }


@app.get("/isolations/{isolation_id}", dependencies=[Depends(verify_api_key)])
async def get_isolation(isolation_id: str, db=Depends(get_db)):
    """Get the full isolation package by ID."""
    pkg = db.query(DBIsolationPackage).filter(DBIsolationPackage.id == isolation_id).first()
    if not pkg:
        raise HTTPException(status_code=404, detail="Isolation package not found")

    # Try to parse the package data as JSON
    raw_data = _decrypt_text(pkg.package_data)
    try:
        package_json = json.loads(raw_data)
    except (json.JSONDecodeError, ValueError):
        package_json = {"raw": raw_data}

    return {
        "id": pkg.id,
        "cert_number": _decrypt_text(pkg.cert_number),
        "equipment_tag": _decrypt_text(pkg.equipment_tag),
        "work_description": _decrypt_text(pkg.work_description),
        "work_type": _decrypt_text(pkg.work_type),
        "fluid_service": _decrypt_text(pkg.fluid_service) if pkg.fluid_service else None,
        "facility": _decrypt_text(pkg.facility) if pkg.facility else None,
        "regime": _decrypt_text(pkg.regime) if pkg.regime else None,
        "special_requirements": _decrypt_text(pkg.special_requirements) if pkg.special_requirements else None,
        "drawing_ids": json.loads(pkg.drawing_ids),
        "hazard_classification": pkg.hazard_classification,
        "valve_count": pkg.valve_count,
        "blind_count": pkg.blind_count,
        "step_count": pkg.step_count,
        "energy_source_count": pkg.energy_source_count,
        "status": pkg.status,
        "created_at": pkg.created_at.isoformat(),
        "updated_at": pkg.updated_at.isoformat(),
        "package": package_json,
    }


@app.delete("/isolations/{isolation_id}", dependencies=[Depends(verify_api_key)])
async def delete_isolation(isolation_id: str, db=Depends(get_db)):
    """Delete an isolation package."""
    pkg = db.query(DBIsolationPackage).filter(DBIsolationPackage.id == isolation_id).first()
    if not pkg:
        raise HTTPException(status_code=404, detail="Isolation package not found")
    db.delete(pkg)
    db.commit()
    return {"deleted": isolation_id}


# ---------------------------------------------------------------------------
# PDF Tools — Download, Merge, Split, Annotate
# ---------------------------------------------------------------------------

def _decrypt_and_load(filepath: Path) -> bytes:
    """Load and decrypt any encrypted file, returning raw bytes."""
    if not filepath.exists():
        raise FileNotFoundError(f"File not found: {filepath}")
    if ENCRYPTION_KEY:
        from encryption import decrypt_file
        return decrypt_file(str(filepath))
    return filepath.read_bytes()


def _load_pdf_bytes(doc) -> bytes:
    """Load and decrypt a stored PDF, returning raw bytes."""
    filepath = Path(doc.filepath)
    if not filepath.exists():
        raise HTTPException(status_code=404, detail="PDF file not found on disk")
    return _decrypt_and_load(filepath)


@app.get("/documents/{doc_id}/download", dependencies=[Depends(verify_api_key)])
async def download_document(doc_id: str, db=Depends(get_db)):
    """Download the original PDF file."""
    doc = db.query(DBDocument).filter(DBDocument.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    pdf_bytes = _load_pdf_bytes(doc)
    filename = _decrypt_text(doc.filename)
    return Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.get("/documents/{doc_id}/view")
async def view_document_pdf(
    doc_id: str,
    request: Request,
    token: str = Query(default=""),
    key: str = Query(default=""),
    db=Depends(get_db),
):
    """Serve decrypted PDF inline for in-browser viewing.

    Accepts auth via query params (token= or key=) for iframe embedding,
    in addition to the standard Authorization header.
    """
    # Try standard header auth first
    authed = False
    try:
        await verify_auth(request)
        authed = True
    except HTTPException:
        pass

    # Fall back to query-param auth for iframe usage
    if not authed and token:
        payload = _decode_jwt(token)
        if payload:
            authed = True
    if not authed and key and API_KEY:
        if secrets.compare_digest(key, API_KEY):
            authed = True

    if not authed:
        raise HTTPException(status_code=401, detail="Invalid or missing credentials")

    doc = db.query(DBDocument).filter(DBDocument.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    pdf_bytes = _load_pdf_bytes(doc)
    filename = _decrypt_text(doc.filename)
    return Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={"Content-Disposition": f'inline; filename="{filename}"'},
    )


@app.get("/dashboard/stats", dependencies=[Depends(verify_api_key)])
async def dashboard_stats(db=Depends(get_db)):
    """Return summary stats for the dashboard."""
    doc_count = db.query(DBDocument).count()
    chat_count = db.query(DBChatSession).count()
    cache_count = db.query(DBAgentCache).count()
    search_count = db.query(DBSearchResult).count()
    analysis_count = db.query(DBAnalysisReport).count()

    recent_docs = db.query(DBDocument).order_by(DBDocument.uploaded_at.desc()).limit(5).all()
    recent_chats = db.query(DBChatSession).order_by(DBChatSession.updated_at.desc()).limit(5).all()
    recent_cache = db.query(DBAgentCache).order_by(DBAgentCache.created_at.desc()).limit(5).all()

    return {
        "documents": doc_count,
        "chat_sessions": chat_count,
        "agent_cache": cache_count,
        "searches": search_count,
        "analyses": analysis_count,
        "recent_documents": [
            {"id": d.id, "filename": _decrypt_text(d.filename),
             "page_count": d.page_count, "uploaded_at": d.uploaded_at.isoformat() if d.uploaded_at else None}
            for d in recent_docs
        ],
        "recent_chats": [
            {"id": s.id, "title": s.title, "updated_at": s.updated_at.isoformat() if s.updated_at else None}
            for s in recent_chats
        ],
        "recent_agents": [
            {"id": c.id, "agent_type": c.agent_type, "model_used": c.model_used,
             "params_summary": c.params_summary, "created_at": c.created_at.isoformat() if c.created_at else None}
            for c in recent_cache
        ],
    }


class BulkAuditRequest(BaseModel):
    doc_ids: list[str] = Field(default=[], description="Document IDs to audit (empty = all)")
    focus_areas: str = Field(default="", max_length=2000)
    model: str = Field(default="haiku", pattern=r"^(sonnet|haiku)$")


@app.post("/agents/bulk-audit", dependencies=[Depends(verify_api_key)])
async def agent_bulk_audit(body: BulkAuditRequest, request: Request, db=Depends(get_db)):
    """Run compliance audit on multiple documents, streaming progress."""
    api_key = os.getenv("ANTHROPIC_API_KEY")
    if not api_key:
        raise HTTPException(status_code=500, detail="ANTHROPIC_API_KEY not configured on server")

    current_user_id = getattr(request.state, "user_id", None)

    if body.doc_ids:
        docs = db.query(DBDocument).filter(DBDocument.id.in_(body.doc_ids)).all()
    else:
        docs = db.query(DBDocument).all()

    if not docs:
        raise HTTPException(status_code=404, detail="No documents found")

    model = _resolve_agent_model(body.model)
    focus = body.focus_areas

    from anthropic import Anthropic
    client = Anthropic(api_key=api_key)
    web_tools = [{"type": "web_search_20250305", "name": "web_search"}] if CHAT_WEB_SEARCH else None

    async def run_bulk():
        total = len(docs)
        yield f"data: {json.dumps({'type': 'bulk_start', 'total': total})}\n\n"

        for idx, doc in enumerate(docs):
            doc_name = _decrypt_text(doc.filename)
            yield f"data: {json.dumps({'type': 'bulk_progress', 'current': idx + 1, 'total': total, 'doc_name': doc_name, 'status': 'running'})}\n\n"

            doc_hash = _get_doc_hash(doc)
            cache_key = _agent_cache_key("audit", model, [doc_hash], focus)
            save_db = SessionLocal()
            try:
                cached = _check_agent_cache(save_db, cache_key, current_user_id)
            finally:
                save_db.close()

            if cached:
                yield f"data: {json.dumps({'type': 'bulk_result', 'current': idx + 1, 'doc_name': doc_name, 'doc_id': doc.id, 'cached': True, 'summary': cached[:500]})}\n\n"
                continue

            try:
                pdf_bytes = _load_pdf_bytes(doc)
                structured = extract_structured_text(pdf_bytes)
                doc_content = ""
                for page_data in structured:
                    doc_content += f"\n--- Page {page_data['page']} ---\n"
                    for block in page_data.get("blocks", []):
                        prefix = f"[{block['type'].upper()}] " if block['type'] != 'paragraph' else ""
                        doc_content += f"{prefix}{block['text']}\n"
                if len(doc_content) > 50000:
                    doc_content = doc_content[:50000] + "\n[... truncated ...]"

                report = _call_claude(client,
                    f"""You are a compliance auditor. Analyze this document against current regulations.

DOCUMENT: {doc_name}
{doc_content}

Provide a concise compliance audit with:
1. Overall compliance rating (percentage)
2. Risk level (HIGH/MEDIUM/LOW)
3. Key findings (max 5 bullet points)
4. Critical issues requiring immediate attention

Use markdown formatting.""",
                    "Audit this document for compliance.", tools=web_tools,
                    max_tokens=AGENT_MAX_TOKENS, model=model)

                save_db = SessionLocal()
                try:
                    _save_agent_cache(save_db, cache_key, "audit", model,
                                      report, [doc.id], f"bulk|focus: {focus}" if focus else "bulk",
                                      user_id=current_user_id)
                finally:
                    save_db.close()

                yield f"data: {json.dumps({'type': 'bulk_result', 'current': idx + 1, 'doc_name': doc_name, 'doc_id': doc.id, 'cached': False, 'summary': report[:500]})}\n\n"

            except Exception as e:
                import logging
                logging.getLogger("pdfhelper").error("Bulk audit failed for %s: %s", doc_name, e)
                err = "Audit failed" if IS_PRODUCTION else str(e)
                yield f"data: {json.dumps({'type': 'bulk_result', 'current': idx + 1, 'doc_name': doc_name, 'doc_id': doc.id, 'error': err})}\n\n"

        yield _agent_done()

    return StreamingResponse(run_bulk(), media_type="text/event-stream")


class MergeRequest(BaseModel):
    doc_ids: list[str] = Field(..., min_length=2, description="IDs of documents to merge (in order)")
    output_filename: str = Field(default="merged.pdf", max_length=255)


@app.post("/documents/merge", dependencies=[Depends(verify_api_key)])
async def merge_documents(body: MergeRequest, request: Request, db=Depends(get_db)):
    """Merge multiple PDFs into a single new document."""
    import fitz  # PyMuPDF

    docs_db = []
    for did in body.doc_ids:
        doc = db.query(DBDocument).filter(DBDocument.id == did).first()
        if not doc:
            raise HTTPException(status_code=404, detail=f"Document {did} not found")
        docs_db.append(doc)

    merged = fitz.open()
    try:
        for doc in docs_db:
            pdf_bytes = _load_pdf_bytes(doc)
            src = fitz.open(stream=pdf_bytes, filetype="pdf")
            merged.insert_pdf(src)
            src.close()

        merged_bytes = merged.tobytes()
    finally:
        merged.close()

    # Save the merged PDF as a new document
    clean_name = _sanitize_filename(body.output_filename)
    if not clean_name.lower().endswith(".pdf"):
        clean_name += ".pdf"

    doc_id = str(uuid.uuid4())
    save_path = UPLOAD_DIR / f"{doc_id}.pdf.enc"
    _encrypt_and_save(merged_bytes, save_path)

    pages = extract_text_from_bytes(merged_bytes)
    content_hash = hashlib.sha256(merged_bytes).hexdigest()

    db_doc = DBDocument(
        id=doc_id,
        filename=_encrypt_text(clean_name),
        filepath=str(save_path),
        page_count=len(pages),
        text_content=_encrypt_text(json.dumps(pages)),
        content_hash=content_hash,
        uploaded_at=datetime.now(timezone.utc),
    )
    db.add(db_doc)
    db.commit()

    log_upload(_get_client_ip(request), clean_name, doc_id, len(pages))

    return {
        "id": doc_id,
        "filename": clean_name,
        "pages": len(pages),
        "merged_from": body.doc_ids,
    }


class SplitRequest(BaseModel):
    pages: list[int] = Field(..., min_length=1, description="Page numbers to extract (1-based)")
    output_filename: str = Field(default="split.pdf", max_length=255)


@app.post("/documents/{doc_id}/split", dependencies=[Depends(verify_api_key)])
async def split_document(doc_id: str, body: SplitRequest, request: Request, db=Depends(get_db)):
    """Extract specific pages from a PDF into a new document."""
    import fitz

    doc = db.query(DBDocument).filter(DBDocument.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    pdf_bytes = _load_pdf_bytes(doc)
    src = fitz.open(stream=pdf_bytes, filetype="pdf")

    # Validate page numbers (convert 1-based to 0-based)
    total_pages = src.page_count
    zero_based = []
    for p in body.pages:
        if p < 1 or p > total_pages:
            src.close()
            raise HTTPException(
                status_code=400,
                detail=f"Page {p} out of range (document has {total_pages} pages)",
            )
        zero_based.append(p - 1)

    new_pdf = fitz.open()
    try:
        for pg in zero_based:
            new_pdf.insert_pdf(src, from_page=pg, to_page=pg)
        split_bytes = new_pdf.tobytes()
    finally:
        new_pdf.close()
        src.close()

    clean_name = _sanitize_filename(body.output_filename)
    if not clean_name.lower().endswith(".pdf"):
        clean_name += ".pdf"

    new_id = str(uuid.uuid4())
    save_path = UPLOAD_DIR / f"{new_id}.pdf.enc"
    _encrypt_and_save(split_bytes, save_path)

    pages = extract_text_from_bytes(split_bytes)
    content_hash = hashlib.sha256(split_bytes).hexdigest()

    db_doc = DBDocument(
        id=new_id,
        filename=_encrypt_text(clean_name),
        filepath=str(save_path),
        page_count=len(pages),
        text_content=_encrypt_text(json.dumps(pages)),
        content_hash=content_hash,
        uploaded_at=datetime.now(timezone.utc),
    )
    db.add(db_doc)
    db.commit()

    log_upload(_get_client_ip(request), clean_name, new_id, len(pages))

    return {
        "id": new_id,
        "filename": clean_name,
        "pages": len(pages),
        "extracted_from": doc_id,
        "page_numbers": body.pages,
    }


@app.post("/documents/{doc_id}/annotate", dependencies=[Depends(verify_api_key)])
async def annotate_document(
    doc_id: str,
    request: Request,
    db=Depends(get_db),
    text: str = Form(..., description="Text to add"),
    page: int = Form(default=1, description="Page number (1-based)"),
    x: float = Form(default=72, description="X position in points from left"),
    y: float = Form(default=72, description="Y position in points from top"),
    font_size: float = Form(default=12, description="Font size"),
    color: str = Form(default="0,0,0", description="RGB color as 'r,g,b' (0-1 range)"),
    save_as_new: bool = Form(default=True, description="Save as new document instead of overwriting"),
    output_filename: str = Form(default="", description="Output filename (only if save_as_new)"),
):
    """Add text annotation/watermark to a PDF page."""
    import fitz

    doc = db.query(DBDocument).filter(DBDocument.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    pdf_bytes = _load_pdf_bytes(doc)
    pdf = fitz.open(stream=pdf_bytes, filetype="pdf")

    if page < 1 or page > pdf.page_count:
        pdf.close()
        raise HTTPException(
            status_code=400,
            detail=f"Page {page} out of range (document has {pdf.page_count} pages)",
        )

    # Parse and validate color
    try:
        rgb = tuple(float(c.strip()) for c in color.split(","))
        if len(rgb) != 3 or not all(0.0 <= v <= 1.0 for v in rgb):
            raise ValueError
    except (ValueError, TypeError):
        pdf.close()
        raise HTTPException(status_code=400, detail="Color must be three comma-separated floats between 0.0 and 1.0, e.g. '0,0,0' for black")

    pg = pdf[page - 1]
    pg.insert_text(
        fitz.Point(x, y),
        text,
        fontsize=font_size,
        color=rgb,
    )

    annotated_bytes = pdf.tobytes()
    pdf.close()

    client_ip = _get_client_ip(request)
    original_name = _decrypt_text(doc.filename)

    if save_as_new:
        clean_name = _sanitize_filename(output_filename or f"annotated_{original_name}")
        if not clean_name.lower().endswith(".pdf"):
            clean_name += ".pdf"
        new_id = str(uuid.uuid4())
        save_path = UPLOAD_DIR / f"{new_id}.pdf.enc"
        _encrypt_and_save(annotated_bytes, save_path)

        pages_data = extract_text_from_bytes(annotated_bytes)
        content_hash = hashlib.sha256(annotated_bytes).hexdigest()

        db_doc = DBDocument(
            id=new_id,
            filename=_encrypt_text(clean_name),
            filepath=str(save_path),
            page_count=len(pages_data),
            text_content=_encrypt_text(json.dumps(pages_data)),
            content_hash=content_hash,
            uploaded_at=datetime.now(timezone.utc),
        )
        db.add(db_doc)
        db.commit()
        log_upload(client_ip, clean_name, new_id, len(pages_data))
        return {"id": new_id, "filename": clean_name, "pages": len(pages_data)}
    else:
        # Overwrite existing document
        save_path = Path(doc.filepath)
        _encrypt_and_save(annotated_bytes, save_path)
        pages_data = extract_text_from_bytes(annotated_bytes)
        doc.page_count = len(pages_data)
        doc.text_content = _encrypt_text(json.dumps(pages_data))
        doc.content_hash = hashlib.sha256(annotated_bytes).hexdigest()
        db.commit()
        return {"id": doc_id, "filename": original_name, "pages": len(pages_data), "updated": True}


# ---------------------------------------------------------------------------
# Doc Updater — Structure, Regulation Search, Updates, Review, Sessions
# ---------------------------------------------------------------------------

@app.get("/documents/{doc_id}/structure", dependencies=[Depends(verify_api_key)])
async def get_document_structure(doc_id: str, db=Depends(get_db)):
    """Return structured content extraction for a document (headings, paragraphs, lists, tables)."""
    doc = db.query(DBDocument).filter(DBDocument.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    pdf_bytes = _load_pdf_bytes(doc)
    structured = extract_structured_text(pdf_bytes)
    return {"doc_id": doc_id, "filename": _decrypt_text(doc.filename), "pages": structured}


@app.get("/documents/{doc_id}/html", dependencies=[Depends(verify_api_key)])
async def get_document_html(doc_id: str, db=Depends(get_db)):
    """Return an HTML rendering of a document's structured content for the in-browser viewer."""
    doc = db.query(DBDocument).filter(DBDocument.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    pdf_bytes = _load_pdf_bytes(doc)
    structured = extract_structured_text(pdf_bytes)

    html_parts = []
    block_idx = 0
    for page_data in structured:
        page_num = page_data["page"]
        html_parts.append(f'<div class="doc-page" data-page="{page_num}">')
        html_parts.append(f'<div class="page-header">Page {page_num}</div>')
        for block in page_data.get("blocks", []):
            btype = block["type"]
            text = block["text"].replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            bid = f"blk-{block_idx}"
            block_idx += 1
            if btype == "heading":
                html_parts.append(f'<h3 class="doc-heading doc-block" id="{bid}" data-type="heading" data-idx="{block_idx}">{text}</h3>')
            elif btype == "list_item":
                html_parts.append(f'<p class="doc-list-item doc-block" id="{bid}" data-type="list_item" data-idx="{block_idx}">{text}</p>')
            elif btype == "table":
                rows = block.get("rows", [])
                if rows:
                    html_parts.append(f'<table class="doc-table doc-block" id="{bid}" data-type="table" data-idx="{block_idx}"><tbody>')
                    for ri, row in enumerate(rows):
                        tag = "th" if ri == 0 else "td"
                        cells = "".join(
                            f"<{tag}>{str(c).replace('&','&amp;').replace('<','&lt;').replace('>','&gt;')}</{tag}>"
                            for c in row
                        )
                        html_parts.append(f"<tr>{cells}</tr>")
                    html_parts.append("</tbody></table>")
                else:
                    html_parts.append(f'<pre class="doc-table-text doc-block" id="{bid}" data-type="table" data-idx="{block_idx}">{text}</pre>')
            else:
                html_parts.append(f'<p class="doc-para doc-block" id="{bid}" data-type="paragraph" data-idx="{block_idx}">{text}</p>')
        html_parts.append("</div>")

    return HTMLResponse("\n".join(html_parts))


@app.get("/documents/{doc_id}/detect-regulations", dependencies=[Depends(verify_api_key)])
async def detect_regulations(doc_id: str, db=Depends(get_db)):
    """Scan a document and return detected regulation/standard references."""
    doc = db.query(DBDocument).filter(DBDocument.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    pages = json.loads(_decrypt_text(doc.text_content))
    full_text = "\n".join(p["text"] for p in pages if p.get("text"))

    patterns = [
        r'(?:OSHA|29\s*CFR)\s*[\d.]+(?:\([a-z]\))?',
        r'(?:API|ASME|ANSI|NFPA|ISO|IEC|IEEE|ASTM|CSA|CGA|DOT|EPA|MSHA)\s*[\d][\w.\-]*',
        r'(?:AS|BS|EN|DIN|JIS|NF|GB)\s*\d[\w.\-]*',
        r'(?:NEC|NESC|CFR|USC|FR)\s*[\d.]+',
        r'(?:Part|Section|Subpart)\s+\d[\w.\-]*',
        r'(?:29|30|33|40|46|49)\s*CFR\s*[\d.]+',
    ]
    found = set()
    for pattern in patterns:
        for match in re.finditer(pattern, full_text, re.IGNORECASE):
            ref = match.group(0).strip()
            if len(ref) > 3:
                found.add(ref)

    refs = sorted(found)
    suggested_query = ""
    if refs:
        top_refs = refs[:10]
        suggested_query = "Current requirements for: " + ", ".join(top_refs)

    return {
        "doc_id": doc_id,
        "regulations_found": refs,
        "count": len(refs),
        "suggested_query": suggested_query,
    }


class RegulationSearchRequest(BaseModel):
    query: str = Field(max_length=2000)
    doc_id: str | None = Field(default=None)
    context: str = Field(default="", max_length=5000)


@app.post("/regulations/search", dependencies=[Depends(verify_api_key)])
async def search_regulations(body: RegulationSearchRequest, db=Depends(get_db)):
    """Search the web for current regulations relevant to a query or document."""
    api_key = os.getenv("ANTHROPIC_API_KEY")
    if not api_key:
        raise HTTPException(status_code=500, detail="ANTHROPIC_API_KEY not configured on server")

    doc_context = ""
    if body.doc_id:
        doc = db.query(DBDocument).filter(DBDocument.id == body.doc_id).first()
        if doc:
            pages = json.loads(_decrypt_text(doc.text_content))
            full_text = "\n".join(p["text"] for p in pages if p.get("text"))
            if len(full_text) > 30000:
                full_text = full_text[:30000] + "\n[... truncated ...]"
            doc_context = f"\n\nDOCUMENT CONTENT TO CHECK AGAINST:\n{full_text}"

    from anthropic import Anthropic
    client = Anthropic(api_key=api_key)

    system = f"""You are a regulatory compliance researcher. Search the web for current regulations, standards, and requirements related to the user's query.{doc_context}

Return your findings as a structured analysis with these sections:
1. **Regulations Found** — list each regulation/standard with its current version and source
2. **Key Requirements** — summarize the main requirements from each regulation
3. **Relevance to Document** — if a document was provided, explain how each regulation applies
4. **Recommended Updates** — specific changes the document should make for compliance

Use markdown formatting. Cite sources with URLs when available."""

    create_kwargs = dict(
        model=CHAT_MODEL,
        max_tokens=CHAT_MAX_TOKENS,
        system=system,
        messages=[{"role": "user", "content": body.query + (f"\n\nAdditional context: {body.context}" if body.context else "")}],
        tools=[{"type": "web_search_20250305", "name": "web_search"}],
    )

    async def stream_search():
        full_reply = ""
        try:
            with client.messages.stream(**create_kwargs) as stream:
                for event in stream:
                    if hasattr(event, 'type'):
                        if event.type == 'content_block_start':
                            if hasattr(event.content_block, 'type') and event.content_block.type == 'server_tool_use':
                                yield f"data: {json.dumps({'type': 'status', 'message': 'Searching the web for regulations...'})}\n\n"
                        elif event.type == 'content_block_delta':
                            if hasattr(event.delta, 'text'):
                                full_reply += event.delta.text
                                yield f"data: {json.dumps({'type': 'chunk', 'text': event.delta.text})}\n\n"
            yield f"data: {json.dumps({'type': 'done', 'full_text': full_reply})}\n\n"
        except Exception as e:
            import logging
            logging.getLogger("pdfhelper").error("Regulation search failed: %s", e)
            err_msg = "Search failed" if IS_PRODUCTION else f"Search failed: {str(e)}"
            yield f"data: {json.dumps({'type': 'error', 'detail': err_msg})}\n\n"

    return StreamingResponse(stream_search(), media_type="text/event-stream")


class GenerateUpdatesRequest(BaseModel):
    doc_id: str = Field(max_length=100)
    regulation_text: str = Field(max_length=100000)
    additional_instructions: str = Field(default="", max_length=5000)


@app.post("/documents/{doc_id}/generate-updates", dependencies=[Depends(verify_api_key)])
async def generate_updates(doc_id: str, body: GenerateUpdatesRequest, db=Depends(get_db)):
    """Generate proposed updates for a document based on regulation findings."""
    api_key = os.getenv("ANTHROPIC_API_KEY")
    if not api_key:
        raise HTTPException(status_code=500, detail="ANTHROPIC_API_KEY not configured on server")

    doc = db.query(DBDocument).filter(DBDocument.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    pdf_bytes = _load_pdf_bytes(doc)
    structured = extract_structured_text(pdf_bytes)
    doc_content = ""
    for page_data in structured:
        doc_content += f"\n--- Page {page_data['page']} ---\n"
        for block in page_data.get("blocks", []):
            prefix = f"[{block['type'].upper()}] " if block['type'] != 'paragraph' else ""
            doc_content += f"{prefix}{block['text']}\n"
    if len(doc_content) > 80000:
        doc_content = doc_content[:80000] + "\n[... truncated ...]"

    from anthropic import Anthropic
    client = Anthropic(api_key=api_key)

    system = f"""You are a document update specialist. You have the original document structure and regulation findings. Generate specific, localized updates.

ORIGINAL DOCUMENT:
{doc_content}

REGULATION FINDINGS:
{body.regulation_text[:50000]}

{('ADDITIONAL INSTRUCTIONS: ' + body.additional_instructions) if body.additional_instructions else ''}

You MUST respond with a JSON array of update objects. Each object has these fields:
- "id": a unique short identifier like "upd-1", "upd-2", etc.
- "section": the section name or heading this update applies to
- "change_type": one of "replace", "insert", or "delete"
- "original_text": the exact original text being changed (quote it precisely)
- "proposed_text": the new text to replace it with (empty string for deletions)
- "rationale": why this change is needed, citing the specific regulation

Respond ONLY with a valid JSON array. No markdown, no explanation outside the JSON. Example:
[
  {{"id": "upd-1", "section": "PPE Requirements", "change_type": "replace", "original_text": "Hard hats required", "proposed_text": "Hard hats and safety glasses required per OSHA 1926.100", "rationale": "OSHA 1926.100 requires eye protection in addition to head protection"}}
]"""

    create_kwargs = dict(
        model=CHAT_MODEL,
        max_tokens=AGENT_MAX_TOKENS,
        system=system,
        messages=[{"role": "user", "content": "Analyze the document and generate all necessary updates as a JSON array."}],
    )

    response = client.messages.create(**create_kwargs)
    full_text = ""
    for block in response.content:
        if hasattr(block, "text"):
            full_text += block.text

    # Try to parse as JSON array
    updates = []
    try:
        # Strip markdown code fences if present
        cleaned = full_text.strip()
        if cleaned.startswith("```"):
            cleaned = re.sub(r'^```\w*\n?', '', cleaned)
            cleaned = re.sub(r'\n?```$', '', cleaned)
        updates = json.loads(cleaned)
        if not isinstance(updates, list):
            updates = [updates]
    except json.JSONDecodeError:
        # Fallback: return as single raw text block
        updates = [{"id": "upd-1", "section": "Full Document", "change_type": "replace",
                     "original_text": "", "proposed_text": full_text, "rationale": "AI-generated update (could not parse structured blocks)"}]

    return {"updates": updates, "raw_text": full_text}


class ReviewSectionRequest(BaseModel):
    highlighted_text: str = Field(max_length=10000)
    context: str = Field(default="", max_length=5000)
    focus: str = Field(default="compliance,clarity,completeness", max_length=500)


@app.post("/documents/{doc_id}/review-section", dependencies=[Depends(verify_api_key)])
async def review_section(doc_id: str, body: ReviewSectionRequest, db=Depends(get_db)):
    """AI-review a highlighted section of a document for compliance and improvements."""
    api_key = os.getenv("ANTHROPIC_API_KEY")
    if not api_key:
        raise HTTPException(status_code=500, detail="ANTHROPIC_API_KEY not configured on server")

    doc = db.query(DBDocument).filter(DBDocument.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    pages = json.loads(_decrypt_text(doc.text_content))
    full_text = "\n".join(p["text"] for p in pages if p.get("text"))
    if len(full_text) > 40000:
        full_text = full_text[:40000] + "\n[... truncated ...]"

    from anthropic import Anthropic
    client = Anthropic(api_key=api_key)

    focus_areas = [f.strip() for f in body.focus.split(",") if f.strip()]
    focus_str = ", ".join(focus_areas) if focus_areas else "compliance, clarity, completeness"

    system = f"""You are a document review specialist. The user has highlighted a section of a procedure document for review.

FULL DOCUMENT CONTEXT:
{full_text}

Review the highlighted section focusing on: {focus_str}

You MUST respond with a JSON object containing:
- "issues": array of strings describing problems found
- "suggested_replacement": the rewritten section text
- "rationale": explanation of why changes were made
- "regulation_refs": array of relevant regulation references

Respond ONLY with valid JSON. No markdown outside the JSON."""

    create_kwargs = dict(
        model=CHAT_MODEL,
        max_tokens=CHAT_MAX_TOKENS,
        system=system,
        messages=[{"role": "user", "content": f"Review this highlighted section:\n\n{body.highlighted_text}" + (f"\n\nAdditional context: {body.context}" if body.context else "")}],
    )
    if CHAT_WEB_SEARCH:
        create_kwargs["tools"] = [{"type": "web_search_20250305", "name": "web_search"}]

    async def stream_review():
        full_reply = ""
        try:
            with client.messages.stream(**create_kwargs) as stream:
                for event in stream:
                    if hasattr(event, 'type'):
                        if event.type == 'content_block_start':
                            if hasattr(event.content_block, 'type') and event.content_block.type == 'server_tool_use':
                                yield f"data: {json.dumps({'type': 'status', 'message': 'Searching regulations...'})}\n\n"
                        elif event.type == 'content_block_delta':
                            if hasattr(event.delta, 'text'):
                                full_reply += event.delta.text
                                yield f"data: {json.dumps({'type': 'chunk', 'text': event.delta.text})}\n\n"
            yield f"data: {json.dumps({'type': 'done', 'full_text': full_reply})}\n\n"
        except Exception as e:
            import logging
            logging.getLogger("pdfhelper").error("Section review failed: %s", e)
            err_msg = "Review failed" if IS_PRODUCTION else f"Review failed: {str(e)}"
            yield f"data: {json.dumps({'type': 'error', 'detail': err_msg})}\n\n"

    return StreamingResponse(stream_review(), media_type="text/event-stream")


class ApplyUpdatesRequest(BaseModel):
    updates_markdown: str = Field(max_length=200000, description="The full AI-generated updates text to apply")
    title: str = Field(default="Updated Document", max_length=255)
    include_vba: bool = Field(default=False)


@app.post("/documents/{doc_id}/apply-updates", dependencies=[Depends(verify_api_key)])
async def apply_updates_to_document(doc_id: str, body: ApplyUpdatesRequest, db=Depends(get_db)):
    """Generate a Word document with the proposed updates applied."""
    doc = db.query(DBDocument).filter(DBDocument.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    api_key = os.getenv("ANTHROPIC_API_KEY")
    if not api_key:
        raise HTTPException(status_code=500, detail="ANTHROPIC_API_KEY not configured on server")

    pages = json.loads(_decrypt_text(doc.text_content))
    full_text = "\n".join(p["text"] for p in pages if p.get("text"))
    if len(full_text) > 80000:
        full_text = full_text[:80000] + "\n[... truncated ...]"

    from anthropic import Anthropic
    client = Anthropic(api_key=api_key)

    system = f"""You are a professional document writer. You have the original document and a set of proposed updates. Write the COMPLETE updated document incorporating all the accepted changes.

ORIGINAL DOCUMENT:
{full_text}

PROPOSED UPDATES:
{body.updates_markdown[:80000]}

Write the complete updated document using markdown formatting:
- Use # for main title, ## for sections, ### for subsections
- Use **bold** for emphasis
- Use numbered lists (1. 2. 3.) and bullet lists (- item)
- Preserve the original document structure and tone
- Incorporate all the proposed changes
- The output should be the FULL document, not just the changed sections"""

    response = client.messages.create(
        model=CHAT_MODEL,
        max_tokens=AGENT_MAX_TOKENS,
        system=system,
        messages=[{"role": "user", "content": f"Please write the complete updated document titled '{body.title}'."}],
    )
    full_doc_text = ""
    for block in response.content:
        if hasattr(block, "text"):
            full_doc_text += block.text

    if not full_doc_text.strip():
        raise HTTPException(status_code=500, detail="AI failed to generate document")

    docx_bytes = _markdown_to_docx(full_doc_text, body.title)
    safe_title = re.sub(r'[^\w\s-]', '', body.title)[:50].strip() or "updated-document"

    if body.include_vba:
        import zipfile
        zip_buf = io.BytesIO()
        vba_code = _build_vba_module(body.title)
        with zipfile.ZipFile(zip_buf, 'w', zipfile.ZIP_DEFLATED) as zf:
            zf.writestr(f"{safe_title}.docx", docx_bytes)
            zf.writestr(f"{safe_title}_macros.bas", vba_code)
            zf.writestr("README.txt",
                f"UPDATED DOCUMENT PACKAGE\n"
                f"========================\n\n"
                f"1. {safe_title}.docx - The updated document\n"
                f"2. {safe_title}_macros.bas - VBA macros for formatting\n\n"
                f"To use macros: Open .docx in Word, Alt+F11, File > Import File\n"
            )
        return Response(
            content=zip_buf.getvalue(),
            media_type="application/zip",
            headers={"Content-Disposition": f'attachment; filename="{safe_title}_package.zip"'},
        )

    return Response(
        content=docx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{safe_title}.docx"'},
    )


# -- Update Sessions: save / load / list --

class SaveSessionRequest(BaseModel):
    doc_id: str = Field(max_length=100)
    title: str = Field(default="", max_length=255)
    regulation_query: str = Field(default="", max_length=2000)
    regulation_results: str = Field(default="", max_length=200000)
    updates_json: str = Field(default="[]", max_length=200000)
    accepted_ids: list[str] = Field(default=[])


@app.post("/updater/sessions", dependencies=[Depends(verify_api_key)])
async def save_update_session(body: SaveSessionRequest, request: Request, db=Depends(get_db)):
    """Save a Doc Updater session so the user can resume later."""
    now = datetime.now(timezone.utc)
    current_user_id = getattr(request.state, "user_id", None)
    session = DBUpdateSession(
        id=str(uuid.uuid4()),
        doc_id=body.doc_id,
        user_id=current_user_id,
        title=body.title or f"Session {now.strftime('%b %d %H:%M')}",
        regulation_query=body.regulation_query,
        regulation_results=_encrypt_text(body.regulation_results) if body.regulation_results else None,
        updates_json=_encrypt_text(body.updates_json) if body.updates_json else None,
        accepted_ids=json.dumps(body.accepted_ids),
        status="draft",
        created_at=now,
        updated_at=now,
    )
    db.add(session)
    db.commit()
    return {"id": session.id, "title": session.title, "created_at": now.isoformat()}


@app.get("/updater/sessions", dependencies=[Depends(verify_api_key)])
async def list_update_sessions(request: Request, db=Depends(get_db)):
    """List saved Doc Updater sessions."""
    current_user_id = getattr(request.state, "user_id", None)
    query = db.query(DBUpdateSession).order_by(DBUpdateSession.updated_at.desc())
    if current_user_id:
        query = query.filter(DBUpdateSession.user_id == current_user_id)
    sessions = query.limit(50).all()
    return {"sessions": [{
        "id": s.id, "doc_id": s.doc_id, "title": s.title,
        "status": s.status, "regulation_query": s.regulation_query,
        "created_at": s.created_at.isoformat(), "updated_at": s.updated_at.isoformat(),
    } for s in sessions]}


@app.get("/updater/sessions/{session_id}", dependencies=[Depends(verify_api_key)])
async def get_update_session(session_id: str, request: Request, db=Depends(get_db)):
    """Load a saved Doc Updater session."""
    session = db.query(DBUpdateSession).filter(DBUpdateSession.id == session_id).first()
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    current_user_id = getattr(request.state, "user_id", None)
    if current_user_id and session.user_id and session.user_id != current_user_id:
        raise HTTPException(status_code=403, detail="You do not own this session")
    return {
        "id": session.id, "doc_id": session.doc_id, "title": session.title,
        "regulation_query": session.regulation_query,
        "regulation_results": _decrypt_text(session.regulation_results) if session.regulation_results else "",
        "updates_json": _decrypt_text(session.updates_json) if session.updates_json else "[]",
        "accepted_ids": json.loads(session.accepted_ids) if session.accepted_ids else [],
        "status": session.status,
    }


@app.delete("/updater/sessions/{session_id}", dependencies=[Depends(verify_api_key)])
async def delete_update_session(session_id: str, request: Request, db=Depends(get_db)):
    """Delete a saved Doc Updater session."""
    session = db.query(DBUpdateSession).filter(DBUpdateSession.id == session_id).first()
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    current_user_id = getattr(request.state, "user_id", None)
    if current_user_id and session.user_id and session.user_id != current_user_id:
        raise HTTPException(status_code=403, detail="You do not own this session")
    db.delete(session)
    db.commit()
    return {"deleted": True}


# ---------------------------------------------------------------------------
# AI Agents — Multi-step autonomous tasks streamed into chat
# ---------------------------------------------------------------------------

def _agent_step(step: int, total: int, name: str, status: str = "running"):
    return f"data: {json.dumps({'type': 'agent_step', 'step': step, 'total': total, 'name': name, 'status': status})}\n\n"


def _agent_chunk(text: str):
    return f"data: {json.dumps({'type': 'chunk', 'text': text})}\n\n"


def _agent_done():
    return f"data: {json.dumps({'type': 'done'})}\n\n"


def _agent_error(msg: str):
    return f"data: {json.dumps({'type': 'error', 'detail': msg})}\n\n"


AGENT_MODELS = {
    "sonnet": "claude-sonnet-5",
    "haiku": "claude-haiku-4-5-20251001",
}


def _resolve_agent_model(model_choice: str) -> str:
    """Resolve user's model choice to a valid model ID."""
    if model_choice in AGENT_MODELS:
        return AGENT_MODELS[model_choice]
    return CHAT_MODEL


def _call_claude(client, system: str, user_msg: str, tools=None, max_tokens=None, model=None):
    """Synchronous Claude call using streaming internally to avoid SDK timeout on long requests."""
    kwargs = dict(
        model=model or CHAT_MODEL,
        max_tokens=max_tokens or AGENT_MAX_TOKENS,
        system=system,
        messages=[{"role": "user", "content": user_msg}],
    )
    if tools:
        kwargs["tools"] = tools
    result = []
    with client.messages.stream(**kwargs) as stream:
        for event in stream:
            if hasattr(event, 'type') and event.type == 'content_block_delta':
                if hasattr(event.delta, 'text'):
                    result.append(event.delta.text)
    return "".join(result)


async def _call_claude_bg(client, system, user_msg, tools=None, max_tokens=None, model=None):
    """Run _call_claude in a background thread (non-blocking for async code)."""
    import asyncio
    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(
        None, lambda: _call_claude(client, system, user_msg, tools, max_tokens, model)
    )


def _stream_claude(client, system: str, user_msg: str, tools=None, max_tokens=None, model=None):
    """Stream Claude response, yields text chunks."""
    kwargs = dict(
        model=model or CHAT_MODEL,
        max_tokens=max_tokens or AGENT_MAX_TOKENS,
        system=system,
        messages=[{"role": "user", "content": user_msg}],
    )
    if tools:
        kwargs["tools"] = tools
    with client.messages.stream(**kwargs) as stream:
        for event in stream:
            if hasattr(event, 'type') and event.type == 'content_block_delta':
                if hasattr(event.delta, 'text'):
                    yield event.delta.text


def _agent_cache_key(agent_type: str, model: str, doc_hashes: list[str], params: str) -> str:
    """Generate a SHA-256 cache key from agent inputs."""
    import hashlib
    raw = f"{agent_type}|{model}|{'|'.join(sorted(doc_hashes))}|{params}"
    return hashlib.sha256(raw.encode()).hexdigest()


def _get_doc_hash(doc) -> str:
    """Get or compute a content hash for a document."""
    if doc.content_hash:
        return doc.content_hash
    import hashlib
    pdf_bytes = _load_pdf_bytes(doc)
    return hashlib.sha256(pdf_bytes).hexdigest()


def _check_agent_cache(db, cache_key: str, user_id: str | None = None):
    """Look up a cached agent result scoped to user. Returns decrypted text or None."""
    q = db.query(DBAgentCache).filter(DBAgentCache.cache_key == cache_key)
    if user_id:
        q = q.filter(DBAgentCache.user_id == user_id)
    else:
        q = q.filter(DBAgentCache.user_id.is_(None))
    cached = q.first()
    if not cached:
        return None
    if cached.expires_at and cached.expires_at < datetime.now(timezone.utc):
        db.delete(cached)
        db.commit()
        return None
    return _decrypt_text(cached.result_data)


def _save_agent_cache(db, cache_key: str, agent_type: str, model: str,
                      result: str, doc_ids: list[str], params_summary: str,
                      user_id: str | None = None):
    """Save an agent result to encrypted cache, scoped to user."""
    from datetime import timedelta
    db.add(DBAgentCache(
        id=str(uuid.uuid4()),
        user_id=user_id,
        cache_key=cache_key,
        agent_type=agent_type,
        model_used=model,
        result_data=_encrypt_text(result),
        doc_ids=json.dumps(doc_ids),
        params_summary=params_summary[:200] if params_summary else "",
        created_at=datetime.now(timezone.utc),
        expires_at=datetime.now(timezone.utc) + timedelta(days=7),
    ))
    db.commit()


class ComplianceAuditRequest(BaseModel):
    doc_id: str = Field(max_length=100)
    focus_areas: str = Field(default="", max_length=2000)
    model: str = Field(default="sonnet", pattern=r"^(sonnet|haiku)$")


@app.post("/agents/compliance-audit", dependencies=[Depends(verify_api_key)])
async def agent_compliance_audit(body: ComplianceAuditRequest, request: Request, db=Depends(get_db)):
    """Multi-step compliance audit agent."""
    api_key = os.getenv("ANTHROPIC_API_KEY")
    if not api_key:
        raise HTTPException(status_code=500, detail="ANTHROPIC_API_KEY not configured on server")

    doc = db.query(DBDocument).filter(DBDocument.id == body.doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    current_user_id = getattr(request.state, "user_id", None)
    model = _resolve_agent_model(body.model)
    doc_hash = _get_doc_hash(doc)
    cache_key = _agent_cache_key("audit", model, [doc_hash], body.focus_areas)

    cached = _check_agent_cache(db, cache_key, current_user_id)
    if cached:
        async def return_cached():
            yield f"data: {json.dumps({'type': 'cached', 'message': 'Returning cached result'})}\n\n"
            yield _agent_chunk(cached)
            yield _agent_done()
        return StreamingResponse(return_cached(), media_type="text/event-stream")

    pdf_bytes = _load_pdf_bytes(doc)
    structured = extract_structured_text(pdf_bytes)
    doc_name = _decrypt_text(doc.filename)
    focus = body.focus_areas

    doc_content = ""
    for page_data in structured:
        doc_content += f"\n--- Page {page_data['page']} ---\n"
        for block in page_data.get("blocks", []):
            prefix = f"[{block['type'].upper()}] " if block['type'] != 'paragraph' else ""
            doc_content += f"{prefix}{block['text']}\n"
    if len(doc_content) > 80000:
        doc_content = doc_content[:80000] + "\n[... truncated ...]"

    from anthropic import Anthropic
    client = Anthropic(api_key=api_key)
    web_tools = [{"type": "web_search_20250305", "name": "web_search"}] if CHAT_WEB_SEARCH else None

    async def run_audit():
        import asyncio
        full_report = ""
        try:
            yield _agent_step(1, 4, "Analyzing document structure")
            task = asyncio.create_task(_call_claude_bg(client,
                "You are a document analyst. Identify the key sections, scope, and purpose of this procedure document. List each section with a one-line summary.",
                f"Analyze this document:\n\nFILENAME: {doc_name}\n\n{doc_content}", model=model))
            while not task.done():
                yield ":\n\n"
                await asyncio.sleep(3)
            analysis = task.result()
            yield _agent_step(1, 4, "Analyzing document structure", "done")

            yield _agent_step(2, 4, "Searching current regulations")
            reg_query = f"Current regulations and standards applicable to: {doc_name}."
            if focus:
                reg_query += f" Focus areas: {focus}."
            reg_query += f"\n\nDocument sections found:\n{analysis[:3000]}"
            task = asyncio.create_task(_call_claude_bg(client,
                "You are a regulatory researcher. Search the web for current, applicable regulations, standards, and industry requirements. List each regulation with its full title, version/year, and key requirements.",
                reg_query, tools=web_tools, model=model))
            while not task.done():
                yield ":\n\n"
                await asyncio.sleep(3)
            regulations = task.result()
            yield _agent_step(2, 4, "Searching current regulations", "done")

            yield _agent_step(3, 4, "Cross-referencing sections against regulations")
            task = asyncio.create_task(_call_claude_bg(client,
                f"""You are a compliance auditor. You have:

DOCUMENT: {doc_name}
{doc_content[:40000]}

APPLICABLE REGULATIONS:
{regulations[:20000]}

For EACH section of the document, determine:
- PASS: Section meets regulatory requirements
- FAIL: Section is missing required content or contradicts regulations
- WARNING: Section is partially compliant or could be improved

Be specific about what's missing or wrong. Cite the exact regulation.""",
                "Perform the cross-reference audit for every section.", model=model))
            while not task.done():
                yield ":\n\n"
                await asyncio.sleep(3)
            cross_ref = task.result()
            yield _agent_step(3, 4, "Cross-referencing sections against regulations", "done")

            yield _agent_step(4, 4, "Generating audit report")
            for chunk in _stream_claude(client,
                f"""You are a compliance audit report writer. Using the analysis below, write a complete, professional audit report.

DOCUMENT ANALYZED: {doc_name}
{('FOCUS AREAS: ' + focus) if focus else ''}

DOCUMENT STRUCTURE ANALYSIS:
{analysis[:5000]}

APPLICABLE REGULATIONS:
{regulations[:10000]}

CROSS-REFERENCE FINDINGS:
{cross_ref[:20000]}

Format the report with these sections:
# Compliance Audit Report: [Document Name]

## Executive Summary
Brief overview with overall compliance rating (percentage) and risk level.

## Regulations Reviewed
Table of all regulations checked.

## Section-by-Section Findings
For each section: status (PASS/FAIL/WARNING), finding detail, regulation reference, recommended action.

## Critical Issues
List any FAIL items that need immediate attention.

## Recommendations
Prioritized list of changes to achieve full compliance.

Use markdown formatting with tables where appropriate.""",
                "Write the complete audit report.", max_tokens=CHAT_MAX_TOKENS * 3, model=model):
                full_report += chunk
                yield _agent_chunk(chunk)

            yield _agent_step(4, 4, "Generating audit report", "done")

            save_db = SessionLocal()
            try:
                _save_agent_cache(save_db, cache_key, "audit", model,
                                  full_report, [body.doc_id], f"focus: {focus}" if focus else "",
                                  user_id=current_user_id)
            finally:
                save_db.close()

            yield _agent_done()

        except Exception as e:
            import logging
            logging.getLogger("pdfhelper").error("Compliance audit agent failed: %s", e)
            err_msg = "Audit failed" if IS_PRODUCTION else f"Audit failed: {str(e)}"
            yield _agent_error(err_msg)

    return StreamingResponse(run_audit(), media_type="text/event-stream")


class CompareDocsRequest(BaseModel):
    doc_id_1: str = Field(max_length=100)
    doc_id_2: str = Field(max_length=100)
    focus_areas: str = Field(default="", max_length=2000)
    model: str = Field(default="sonnet", pattern=r"^(sonnet|haiku)$")


@app.post("/agents/compare-docs", dependencies=[Depends(verify_api_key)])
async def agent_compare_docs(body: CompareDocsRequest, request: Request, db=Depends(get_db)):
    """Multi-step document comparison agent."""
    api_key = os.getenv("ANTHROPIC_API_KEY")
    if not api_key:
        raise HTTPException(status_code=500, detail="ANTHROPIC_API_KEY not configured on server")

    doc1 = db.query(DBDocument).filter(DBDocument.id == body.doc_id_1).first()
    doc2 = db.query(DBDocument).filter(DBDocument.id == body.doc_id_2).first()
    if not doc1 or not doc2:
        raise HTTPException(status_code=404, detail="One or both documents not found")

    current_user_id = getattr(request.state, "user_id", None)
    model = _resolve_agent_model(body.model)
    h1, h2 = _get_doc_hash(doc1), _get_doc_hash(doc2)
    cache_key = _agent_cache_key("compare", model, [h1, h2], body.focus_areas)

    cached = _check_agent_cache(db, cache_key, current_user_id)
    if cached:
        async def return_cached():
            yield f"data: {json.dumps({'type': 'cached', 'message': 'Returning cached result'})}\n\n"
            yield _agent_chunk(cached)
            yield _agent_done()
        return StreamingResponse(return_cached(), media_type="text/event-stream")

    def _get_content(doc):
        pdf_bytes = _load_pdf_bytes(doc)
        structured = extract_structured_text(pdf_bytes)
        text = ""
        for page_data in structured:
            text += f"\n--- Page {page_data['page']} ---\n"
            for block in page_data.get("blocks", []):
                prefix = f"[{block['type'].upper()}] " if block['type'] != 'paragraph' else ""
                text += f"{prefix}{block['text']}\n"
        return text[:60000] if len(text) > 60000 else text

    name1 = _decrypt_text(doc1.filename)
    name2 = _decrypt_text(doc2.filename)
    content1 = _get_content(doc1)
    content2 = _get_content(doc2)
    focus = body.focus_areas

    from anthropic import Anthropic
    client = Anthropic(api_key=api_key)

    async def run_compare():
        import asyncio
        full_report = ""
        try:
            yield _agent_step(1, 3, "Analyzing document structures")
            task = asyncio.create_task(_call_claude_bg(client,
                "You are a document analyst. Compare the structure (sections, headings, organization) of these two documents. List the sections in each and note which sections exist in one but not the other.",
                f"DOCUMENT A: {name1}\n{content1[:30000]}\n\nDOCUMENT B: {name2}\n{content2[:30000]}", model=model))
            while not task.done():
                yield ":\n\n"
                await asyncio.sleep(3)
            structures = task.result()
            yield _agent_step(1, 3, "Analyzing document structures", "done")

            yield _agent_step(2, 3, "Comparing content section by section")
            task = asyncio.create_task(_call_claude_bg(client,
                f"""You are a document comparison specialist. Compare these two documents in detail.

DOCUMENT A: {name1}
{content1[:40000]}

DOCUMENT B: {name2}
{content2[:40000]}

STRUCTURE ANALYSIS:
{structures[:5000]}

For each shared section, identify:
- Content that is IDENTICAL or equivalent
- Content that DIFFERS (quote both versions)
- Content that is MISSING from one document
- Content that CONFLICTS between documents

{('FOCUS AREAS: ' + focus) if focus else ''}""",
                "Perform the detailed comparison.", model=model))
            while not task.done():
                yield ":\n\n"
                await asyncio.sleep(3)
            differences = task.result()
            yield _agent_step(2, 3, "Comparing content section by section", "done")

            yield _agent_step(3, 3, "Generating comparison report")
            for chunk in _stream_claude(client,
                f"""You are a document comparison report writer. Write a professional comparison report using the analysis below.

DOCUMENT A: {name1}
DOCUMENT B: {name2}
{('FOCUS AREAS: ' + focus) if focus else ''}

STRUCTURE ANALYSIS:
{structures[:5000]}

DETAILED DIFFERENCES:
{differences[:20000]}

Format the report as:
# Document Comparison: {name1} vs {name2}

## Summary
Overall similarity rating, key differences count, which document is more comprehensive.

## Structure Comparison
Table showing sections side-by-side.

## Key Differences
Each significant difference with quotes from both documents.

## Conflicts Found
Any contradictions between the documents (these are critical).

## Gaps
Content present in one but missing from the other.

## Recommendation
Which document is more complete and what each needs to match the other.

Use markdown with tables.""",
                "Write the complete comparison report.", max_tokens=CHAT_MAX_TOKENS * 3, model=model):
                full_report += chunk
                yield _agent_chunk(chunk)

            yield _agent_step(3, 3, "Generating comparison report", "done")

            save_db = SessionLocal()
            try:
                _save_agent_cache(save_db, cache_key, "compare", model,
                                  full_report, [body.doc_id_1, body.doc_id_2],
                                  f"focus: {focus}" if focus else "",
                                  user_id=current_user_id)
            finally:
                save_db.close()

            yield _agent_done()

        except Exception as e:
            import logging
            logging.getLogger("pdfhelper").error("Compare docs agent failed: %s", e)
            err_msg = "Comparison failed" if IS_PRODUCTION else f"Comparison failed: {str(e)}"
            yield _agent_error(err_msg)

    return StreamingResponse(run_compare(), media_type="text/event-stream")


class ProcedureWriterRequest(BaseModel):
    description: str = Field(max_length=5000)
    source_doc_id: str = Field(default="", description="Primary document to base the procedure on")
    reference_doc_ids: list[str] = Field(default=[])
    include_regulations: bool = Field(default=True)
    model: str = Field(default="sonnet", pattern=r"^(sonnet|haiku)$")


@app.post("/agents/procedure-writer", dependencies=[Depends(verify_api_key)])
async def agent_procedure_writer(body: ProcedureWriterRequest, request: Request, db=Depends(get_db)):
    """Multi-step procedure writing agent."""
    api_key = os.getenv("ANTHROPIC_API_KEY")
    if not api_key:
        raise HTTPException(status_code=500, detail="ANTHROPIC_API_KEY not configured on server")

    current_user_id = getattr(request.state, "user_id", None)
    model = _resolve_agent_model(body.model)

    # Extract source document content (the document to base the procedure on)
    source_content = ""
    source_name = ""
    doc_hashes = []
    if body.source_doc_id:
        source_doc = db.query(DBDocument).filter(DBDocument.id == body.source_doc_id).first()
        if source_doc:
            source_name = _decrypt_text(source_doc.filename)
            doc_hashes.append(_get_doc_hash(source_doc))
            pdf_bytes = _load_pdf_bytes(source_doc)
            structured = extract_structured_text(pdf_bytes)
            for page_data in structured:
                source_content += f"\n--- Page {page_data['page']} ---\n"
                for block in page_data.get("blocks", []):
                    prefix = f"[{block['type'].upper()}] " if block['type'] != 'paragraph' else ""
                    source_content += f"{prefix}{block['text']}\n"
            if len(source_content) > 60000:
                source_content = source_content[:60000] + "\n[... truncated ...]"

    # Extract reference documents (for style matching)
    ref_content = ""
    ref_names = []
    if body.reference_doc_ids:
        ref_ids = [rid for rid in body.reference_doc_ids if rid != body.source_doc_id]
        if ref_ids:
            docs = db.query(DBDocument).filter(DBDocument.id.in_(ref_ids)).all()
            for doc in docs[:3]:
                name = _decrypt_text(doc.filename)
                ref_names.append(name)
                doc_hashes.append(_get_doc_hash(doc))
                pdf_bytes = _load_pdf_bytes(doc)
                structured = extract_structured_text(pdf_bytes)
                text = ""
                for page_data in structured:
                    for block in page_data.get("blocks", []):
                        prefix = f"[{block['type'].upper()}] " if block['type'] != 'paragraph' else ""
                        text += f"{prefix}{block['text']}\n"
                if len(text) > 40000:
                    text = text[:40000] + "\n[... truncated ...]"
                ref_content += f'\n--- REFERENCE: "{name}" ---\n{text}\n'

    cache_params = f"{body.description}|src={body.source_doc_id}|regs={body.include_regulations}"
    cache_key = _agent_cache_key("writer", model, doc_hashes or ["no-refs"], cache_params)

    cached = _check_agent_cache(db, cache_key, current_user_id)
    if cached:
        async def return_cached():
            yield f"data: {json.dumps({'type': 'cached', 'message': 'Returning cached result'})}\n\n"
            yield _agent_chunk(cached)
            yield _agent_done()
        return StreamingResponse(return_cached(), media_type="text/event-stream")

    from anthropic import Anthropic
    client = Anthropic(api_key=api_key)
    web_tools = [{"type": "web_search_20250305", "name": "web_search"}] if CHAT_WEB_SEARCH else None
    steps = 4 if body.include_regulations else 3

    async def run_writer():
        import asyncio
        full_doc = ""
        try:
            # Step 1: Research and outline
            yield _agent_step(1, steps, "Analyzing document and creating outline")
            outline_prompt = f"Create a detailed outline for this procedure:\n\nDESCRIPTION: {body.description}"
            if source_content:
                outline_prompt += f'\n\nSOURCE DOCUMENT ("{source_name}") — base the procedure on this content:\n{source_content[:30000]}'
            if ref_content:
                outline_prompt += f"\n\nREFERENCE PROCEDURES (match their style and structure):\n{ref_content[:15000]}"
            outline_system = "You are a technical procedure writer. Create a detailed section-by-section outline for the requested procedure."
            if source_content:
                outline_system += " The user has provided a SOURCE DOCUMENT — your outline MUST be based on the actual content, topics, processes, and specifics from that document. Extract the real procedures, steps, equipment, roles, and safety information from it. Do NOT invent generic content — use what the document actually says."
            outline_system += " Include all standard sections (Purpose, Scope, Definitions, Responsibilities, Procedure Steps, Safety, Emergency, References). Note what content goes in each section."
            task = asyncio.create_task(_call_claude_bg(client, outline_system, outline_prompt, model=model))
            while not task.done():
                yield ":\n\n"
                await asyncio.sleep(3)
            outline = task.result()
            yield _agent_step(1, steps, "Analyzing document and creating outline", "done")

            regulations = ""
            step_num = 2
            if body.include_regulations:
                yield _agent_step(2, steps, "Searching applicable regulations")
                task = asyncio.create_task(_call_claude_bg(client,
                    "You are a regulatory researcher. Search the web for all regulations, standards, and industry best practices that apply to this procedure. List each with its key requirements.",
                    f"Find applicable regulations for: {body.description}", tools=web_tools, model=model))
                while not task.done():
                    yield ":\n\n"
                    await asyncio.sleep(3)
                regulations = task.result()
                yield _agent_step(2, steps, "Searching applicable regulations", "done")
                step_num = 3

            yield _agent_step(step_num, steps, "Writing procedure content")
            write_system = f"""You are an expert technical procedure writer. Write a complete, professional procedure document.

DESCRIPTION: {body.description}

OUTLINE:
{outline[:10000]}

{('SOURCE DOCUMENT ("' + source_name + '") — base the procedure on this content:' + chr(10) + source_content[:40000]) if source_content else ''}

{('APPLICABLE REGULATIONS:' + chr(10) + regulations[:10000]) if regulations else ''}

{('REFERENCE PROCEDURES (match their style):' + chr(10) + ref_content[:15000]) if ref_content else ''}

Write a complete procedure document with:
- Professional formatting using markdown headings, numbered lists, and tables
- Clear, actionable steps with responsible parties
- Safety warnings and precautions in bold
- Regulatory references where applicable
- Standard sections: Purpose, Scope, Definitions, Responsibilities, Procedure, Safety Requirements, Emergency Procedures, References
- Specific details (not generic placeholders)"""
            if source_content:
                write_system += f"""

CRITICAL: The SOURCE DOCUMENT contains the actual content you must base this procedure on. Use the real processes, equipment names, roles, locations, safety requirements, and specific details from that document. Do NOT generate generic or hypothetical content — extract and organize what the source document actually describes into a well-structured procedure format."""

            for chunk in _stream_claude(client, write_system,
                "Write the complete procedure document now.", max_tokens=CHAT_MAX_TOKENS * 3, model=model):
                full_doc += chunk
                yield _agent_chunk(chunk)

            yield _agent_step(step_num, steps, "Writing procedure content", "done")

            yield _agent_step(steps, steps, "Running quality review")
            task = asyncio.create_task(_call_claude_bg(client,
                f"""You are a procedure quality reviewer. Review this draft procedure for:
1. Completeness — are any standard sections missing?
2. Clarity — are steps clear and unambiguous?
3. Safety — are all hazards addressed?
4. Compliance — does it meet the regulations found?
5. Consistency — any contradictions?

DRAFT:
{full_doc[:40000]}

{('REGULATIONS:' + chr(10) + regulations[:5000]) if regulations else ''}

If issues are found, list them briefly. If the document is good, say so.""",
                "Review the draft and list any issues.", model=model))
            while not task.done():
                yield ":\n\n"
                await asyncio.sleep(3)
            review = task.result()
            yield _agent_step(steps, steps, "Running quality review", "done")

            if "issue" in review.lower() or "missing" in review.lower() or "should" in review.lower():
                full_doc += "\n\n---\n\n## Quality Review Notes\n\n" + review
                yield _agent_chunk("\n\n---\n\n## Quality Review Notes\n\n" + review)

            save_db = SessionLocal()
            try:
                _save_agent_cache(save_db, cache_key, "writer", model,
                                  full_doc, body.reference_doc_ids,
                                  body.description[:200],
                                  user_id=current_user_id)
            finally:
                save_db.close()

            yield _agent_done()

        except Exception as e:
            import logging
            logging.getLogger("pdfhelper").error("Procedure writer agent failed: %s", e)
            err_msg = "Writing failed" if IS_PRODUCTION else f"Writing failed: {str(e)}"
            yield _agent_error(err_msg)

    return StreamingResponse(run_writer(), media_type="text/event-stream")


# ---------------------------------------------------------------------------
# Code Builder Agent
# ---------------------------------------------------------------------------

class CodeBuilderRequest(BaseModel):
    description: str = Field(max_length=50000)
    doc_ids: list[str] = Field(default=[])
    app_type: str = Field(default="dashboard", pattern=r"^(dashboard|form|tracker|checklist|report|custom)$")
    model: str = Field(default="sonnet", pattern=r"^(sonnet|haiku)$")


@app.post("/agents/code-builder", dependencies=[Depends(verify_api_key)])
async def agent_code_builder(body: CodeBuilderRequest, request: Request, db=Depends(get_db)):
    """Multi-step code builder agent that generates complete HTML applications from document data."""
    api_key = os.getenv("ANTHROPIC_API_KEY")
    if not api_key:
        raise HTTPException(status_code=500, detail="ANTHROPIC_API_KEY not configured on server")

    current_user_id = getattr(request.state, "user_id", None)
    model = _resolve_agent_model(body.model)

    doc_content = ""
    doc_names = []
    doc_hashes = []
    selected_docs = db.query(DBDocument).filter(DBDocument.id.in_(body.doc_ids)).all() if body.doc_ids else []
    if not selected_docs:
        selected_docs = db.query(DBDocument).filter(
            DBDocument.user_id == current_user_id
        ).all() if current_user_id else []

    for doc in selected_docs[:5]:
        name = _decrypt_text(doc.filename)
        doc_names.append(name)
        doc_hashes.append(_get_doc_hash(doc))
        pdf_bytes = _load_pdf_bytes(doc)
        structured = extract_structured_text(pdf_bytes)
        text = ""
        for page_data in structured:
            text += f"\n--- Page {page_data['page']} ---\n"
            for block in page_data.get("blocks", []):
                prefix = f"[{block['type'].upper()}] " if block['type'] != 'paragraph' else ""
                text += f"{prefix}{block['text']}\n"
        if len(text) > 80000:
            text = text[:80000] + "\n[... truncated ...]"
        doc_content += f'\n===== DOCUMENT: "{name}" =====\n{text}\n'

    if len(doc_content) > 300000:
        doc_content = doc_content[:300000] + "\n[... truncated ...]"

    cache_params = f"{body.description}|type={body.app_type}"
    cache_key = _agent_cache_key("code-builder", model, doc_hashes or ["no-docs"], cache_params)

    cached = _check_agent_cache(db, cache_key, current_user_id)
    if cached:
        async def return_cached():
            yield f"data: {json.dumps({'type': 'cached', 'message': 'Returning cached result'})}\n\n"
            yield _agent_chunk(cached)
            yield _agent_done()
        return StreamingResponse(return_cached(), media_type="text/event-stream")

    from anthropic import Anthropic
    client = Anthropic(api_key=api_key)

    async def run_builder():
        import asyncio
        full_code = ""
        try:
            # Step 1: Extract and structure the data
            yield _agent_step(1, 3, "Extracting data from documents")
            extract_system = """You are a data extraction specialist. Extract ALL relevant data from the provided documents and organize it into structured JSON-like format.

Extract:
- All tables (preserve rows and columns)
- All lists and checklists
- All named items, categories, and their properties
- All numerical data, dates, statuses
- All personnel roles, responsibilities
- All procedures, steps, requirements
- All section headings and their content hierarchy

Output the extracted data in a clear, organized format that a code generator can use. Include EVERY piece of data — do not summarize or skip anything."""

            task = asyncio.create_task(_call_claude_bg(client, extract_system,
                f"Extract all data from these documents:\n\n{doc_content[:150000]}",
                max_tokens=CHAT_MAX_TOKENS, model=model))
            while not task.done():
                yield ":\n\n"
                await asyncio.sleep(3)
            extracted_data = task.result()
            yield _agent_step(1, 3, "Extracting data from documents", "done")

            # Step 2: Plan the application
            yield _agent_step(2, 3, "Planning application structure")
            plan_system = f"""You are a senior web application architect. Plan a complete single-file HTML5 application.

APPLICATION TYPE: {body.app_type}
USER REQUEST: {body.description}
DOCUMENTS USED: {', '.join(doc_names)}

Plan the application with:
1. Component layout (what sections, panels, and UI elements)
2. Data model (how the extracted data maps to the UI)
3. Interactivity (filters, search, sorting, tabs, modals)
4. Color scheme and visual design approach
5. Responsive layout strategy

Keep the plan focused and specific to the actual data provided."""

            task = asyncio.create_task(_call_claude_bg(client, plan_system,
                f"Plan the application using this extracted data:\n\n{extracted_data[:50000]}",
                model=model))
            while not task.done():
                yield ":\n\n"
                await asyncio.sleep(3)
            plan = task.result()
            yield _agent_step(2, 3, "Planning application structure", "done")

            # Step 3: Generate the complete code
            yield _agent_step(3, 3, "Generating complete HTML application")
            code_system = f"""You are an expert front-end developer. Generate a COMPLETE, PRODUCTION-READY, single-file HTML5 application.

APPLICATION TYPE: {body.app_type}
USER REQUEST: {body.description}
DOCUMENTS USED: {', '.join(doc_names)}

APPLICATION PLAN:
{plan[:15000]}

ABSOLUTE REQUIREMENTS — YOUR CODE WILL BE REJECTED IF ANY ARE VIOLATED:
1. Output ONLY the HTML code — start with <!DOCTYPE html> and end with </html>. No explanations, no markdown fences, no commentary before or after the code.
2. SINGLE FILE — all CSS in <style> tags, all JavaScript in <script> tags. ZERO external dependencies (no CDN links, no imports, no external fonts).
3. ALL DATA EMBEDDED — every piece of data from the documents must be hardcoded as JavaScript arrays/objects inside the file. Never use placeholder data like "Item 1", "Lorem ipsum", or "TODO".
4. FULLY FUNCTIONAL — every button, filter, search box, tab, and interactive element must work. Test mentally: click each button, type in each input — does it do something?
5. RESPONSIVE — must work on both desktop (1200px+) and mobile (375px). Use CSS Grid/Flexbox, relative units, and media queries.
6. PROFESSIONAL DESIGN — clean modern UI with a cohesive color scheme, proper spacing, shadows, rounded corners, hover states on interactive elements.
7. COMPLETE — do not truncate, abbreviate, or skip any section. The output must be the entire working application.
8. INCLUDE: search/filter functionality, sorting where applicable, print-friendly styles (@media print), and a professional header with the application title.

Remember: Output ONLY the raw HTML code. No markdown, no explanations."""

            for chunk in _stream_claude(client, code_system,
                f"Generate the complete HTML application using this data:\n\n{extracted_data[:80000]}",
                max_tokens=CHAT_MAX_TOKENS * 3, model=model):
                full_code += chunk
                yield _agent_chunk(chunk)

            yield _agent_step(3, 3, "Generating complete HTML application", "done")

            save_db = SessionLocal()
            try:
                _save_agent_cache(save_db, cache_key, "code-builder", model,
                                  full_code, body.doc_ids,
                                  body.description[:200],
                                  user_id=current_user_id)
            finally:
                save_db.close()

            yield _agent_done()

        except Exception as e:
            import logging
            logging.getLogger("pdfhelper").error("Code builder agent failed: %s", e)
            err_msg = "Code generation failed" if IS_PRODUCTION else f"Code generation failed: {str(e)}"
            yield _agent_error(err_msg)

    return StreamingResponse(run_builder(), media_type="text/event-stream")


# ---------------------------------------------------------------------------
# Agent Cache Management
# ---------------------------------------------------------------------------

@app.get("/agents/cache", dependencies=[Depends(verify_api_key)])
async def list_agent_cache(request: Request, db=Depends(get_db)):
    """List cached agent results for the current user (metadata only)."""
    current_user_id = getattr(request.state, "user_id", None)
    q = db.query(DBAgentCache)
    if current_user_id:
        q = q.filter(DBAgentCache.user_id == current_user_id)
    else:
        q = q.filter(DBAgentCache.user_id.is_(None))
    entries = q.order_by(DBAgentCache.created_at.desc()).limit(50).all()
    return {
        "cache_entries": [
            {
                "id": e.id,
                "agent_type": e.agent_type,
                "model_used": e.model_used,
                "doc_ids": json.loads(e.doc_ids),
                "params_summary": e.params_summary,
                "created_at": e.created_at.isoformat() if e.created_at else None,
                "expires_at": e.expires_at.isoformat() if e.expires_at else None,
            }
            for e in entries
        ],
        "total": len(entries),
    }


@app.delete("/agents/cache", dependencies=[Depends(verify_api_key)])
async def clear_all_agent_cache(request: Request, db=Depends(get_db)):
    """Clear all cached agent results for the current user only."""
    current_user_id = getattr(request.state, "user_id", None)
    q = db.query(DBAgentCache)
    if current_user_id:
        q = q.filter(DBAgentCache.user_id == current_user_id)
    else:
        q = q.filter(DBAgentCache.user_id.is_(None))
    count = q.delete()
    db.commit()
    return {"deleted": count}


@app.delete("/agents/cache/{cache_id}", dependencies=[Depends(verify_api_key)])
async def delete_agent_cache_entry(cache_id: str, request: Request, db=Depends(get_db)):
    """Delete a specific cached agent result, enforcing ownership."""
    current_user_id = getattr(request.state, "user_id", None)
    entry = db.query(DBAgentCache).filter(DBAgentCache.id == cache_id).first()
    if not entry:
        raise HTTPException(status_code=404, detail="Cache entry not found")
    if current_user_id and entry.user_id and entry.user_id != current_user_id:
        raise HTTPException(status_code=403, detail="You do not own this cache entry")
    db.delete(entry)
    db.commit()
    return {"deleted": True}


# ---------------------------------------------------------------------------
# Poster Generator
# ---------------------------------------------------------------------------

POSTER_SIZES = {
    "letter": {"width": "8.5in", "height": "11in", "px_w": 816, "px_h": 1056},
    "a4": {"width": "210mm", "height": "297mm", "px_w": 794, "px_h": 1123},
    "a3": {"width": "297mm", "height": "420mm", "px_w": 1123, "px_h": 1587},
    "wide": {"width": "11in", "height": "8.5in", "px_w": 1056, "px_h": 816},
    "square": {"width": "10in", "height": "10in", "px_w": 960, "px_h": 960},
    "banner": {"width": "24in", "height": "8in", "px_w": 2304, "px_h": 768},
}

POSTER_STYLES = {
    "bold": "Use strong contrasting colors (black/yellow, red/white), thick borders, large impactful typography, industrial feel.",
    "clean": "Use a clean modern aesthetic with plenty of white space, subtle shadows, thin lines, and a muted professional color palette.",
    "safety": "Use standard safety colors: red for danger/prohibition, yellow for caution, blue for mandatory, green for safe. Include ISO-style safety symbols using Unicode. Add hazard borders with diagonal stripes.",
    "corporate": "Use a polished corporate style with navy/gray tones, structured grid layout, subtle gradients, and professional serif fonts for headings.",
    "vibrant": "Use bright, eye-catching colors with bold gradients, rounded shapes, playful typography, and high visual energy.",
    "retro": "Use a vintage/retro aesthetic with muted earth tones, textured backgrounds via CSS patterns, bold serif fonts, and decorative borders.",
}

POSTER_SYSTEM = """You are a poster design engine. Your ONLY job is to output a complete HTML/CSS poster. You NEVER ask questions, request clarification, provide commentary, or give advice. You ALWAYS respond with raw HTML code and nothing else.

CRITICAL BEHAVIOR:
- NO MATTER WHAT the user writes — whether it's a simple description, a detailed specification, a review request, a consulting prompt, or anything else — you MUST create a poster from it.
- Extract the key content, messages, and themes from the user's prompt and design a poster around them.
- If the user describes an existing document or infographic, create an IMPROVED version as a poster.
- If the user asks for a "review" or "analysis", create a poster that presents those findings visually.
- NEVER respond with text, questions, bullet points, or explanations. ONLY HTML.

DESIGN PRINCIPLES:
- Visual hierarchy: title largest, key message prominent, details smaller
- High contrast for readability — never place light text on light backgrounds
- Balanced whitespace — don't cram everything together
- Consistent alignment and spacing throughout
- Use CSS shapes, gradients, borders, box-shadows for visual interest
- Use Unicode symbols and emoji for icons (e.g. ⚠️ 🔥 ✅ 🏗️ 📋 ☎️ 🚨 ⛑️ 👷 🔒)
- ALL content MUST fit within the poster dimensions — never overflow
- For dense content: use smaller fonts, tighter spacing, multi-column layouts

{style_instruction}

TECHNICAL RULES:
1. Output ONLY the raw HTML — no markdown fences, no explanation, no commentary. NOTHING except HTML.
2. Complete HTML document: <!DOCTYPE html>, <html>, <head> with <title>, <body>.
3. ALL styles in a <style> block. No external stylesheets, images, or JavaScript.
4. Body dimensions: exactly {width} x {height}, margin:0, overflow:hidden.
5. Use web-safe fonts: Arial, Georgia, Impact, Courier New, Trebuchet MS, Verdana.
6. For print quality: use pt/in/cm units for text sizing where appropriate.
7. Add a <title> tag that summarizes the poster content in 3-6 words.
8. CRITICAL — include these print rules in your <style> block so the poster prints cleanly with no browser headers/footers/URLs:
   @page {{ margin: 0; size: {width} {height}; }}
   @media print {{ html, body {{ margin: 0; padding: 0; -webkit-print-color-adjust: exact; print-color-adjust: exact; }} }}

SIZE: {width} x {height}"""

POSTER_UPDATE_SYSTEM = """You are a poster editing engine. Your ONLY job is to output updated HTML/CSS. You NEVER ask questions, provide commentary, or give advice. You ALWAYS respond with raw HTML code and nothing else.

The user will provide the current poster HTML and a description of changes.

RULES:
1. Output ONLY the updated HTML — no markdown fences, no explanation, no commentary. NOTHING except HTML.
2. Keep the same document structure and poster dimensions.
3. Preserve ALL elements the user did NOT ask to change.
4. Apply requested changes precisely — if they say "make title red", only change the title color.
5. Maintain or improve design quality — don't break alignment, spacing, or contrast.
6. All styles stay in <style> block. No external resources, no JavaScript.
7. If the user asks to add content, integrate it naturally into the existing layout.
8. Keep the <title> tag updated if the poster topic changes.
9. ALWAYS keep or add these print rules in the <style> block:
   @page { margin: 0; }
   @media print { html, body { margin: 0; padding: 0; -webkit-print-color-adjust: exact; print-color-adjust: exact; } }"""


@app.post("/posters", dependencies=[Depends(verify_api_key)])
async def create_poster(body: PosterCreateRequest, request: Request, db=Depends(get_db)):
    """Generate a new poster from a text prompt using AI."""
    api_key = os.getenv("ANTHROPIC_API_KEY")
    if not api_key:
        raise HTTPException(status_code=500, detail="ANTHROPIC_API_KEY not configured")

    import anthropic
    client = anthropic.Anthropic(api_key=api_key)
    model = _resolve_agent_model(body.model)
    size = POSTER_SIZES.get(body.size, POSTER_SIZES["letter"])

    style_instruction = ""
    if body.style and body.style in POSTER_STYLES:
        style_instruction = f"STYLE: {body.style.upper()}\n{POSTER_STYLES[body.style]}"

    system = POSTER_SYSTEM.format(width=size["width"], height=size["height"], style_instruction=style_instruction)

    async def generate():
        import asyncio
        yield f"data: {json.dumps({'type': 'status', 'message': 'Designing your poster...'})}\n\n"

        task = asyncio.create_task(_call_claude_bg(
            client, system, body.prompt, max_tokens=8000, model=model
        ))
        while not task.done():
            yield ":\n\n"
            await asyncio.sleep(3)
        html_content = task.result()

        html_content = html_content.strip()
        if html_content.startswith("```"):
            html_content = re.sub(r"^```(?:html)?\s*\n?", "", html_content)
            html_content = re.sub(r"\n?```\s*$", "", html_content)

        current_user_id = getattr(request.state, "user_id", None)
        poster_id = str(uuid.uuid4())
        now = datetime.now(timezone.utc)

        title_match = re.search(r"<title>(.*?)</title>", html_content, re.IGNORECASE)
        title = title_match.group(1) if title_match else body.prompt[:80]

        poster = DBPoster(
            id=poster_id,
            user_id=current_user_id,
            title=_encrypt_text(title),
            prompt_history=_encrypt_text(json.dumps([body.prompt])),
            html_content=_encrypt_text(html_content),
            created_at=now,
            updated_at=now,
        )
        db.add(poster)
        db.commit()

        yield f"data: {json.dumps({'type': 'result', 'poster_id': poster_id, 'title': title, 'html': html_content, 'size': body.size})}\n\n"
        yield f"data: {json.dumps({'type': 'done'})}\n\n"

    return StreamingResponse(generate(), media_type="text/event-stream")


@app.post("/posters/{poster_id}/update", dependencies=[Depends(verify_api_key)])
async def update_poster(poster_id: str, body: PosterUpdateRequest, request: Request, db=Depends(get_db)):
    """Update an existing poster with a new prompt."""
    current_user_id = getattr(request.state, "user_id", None)
    poster = db.query(DBPoster).filter(DBPoster.id == poster_id).first()
    if not poster:
        raise HTTPException(status_code=404, detail="Poster not found")
    if current_user_id and poster.user_id and poster.user_id != current_user_id:
        raise HTTPException(status_code=403, detail="Access denied")

    api_key = os.getenv("ANTHROPIC_API_KEY")
    if not api_key:
        raise HTTPException(status_code=500, detail="ANTHROPIC_API_KEY not configured")

    import anthropic
    client = anthropic.Anthropic(api_key=api_key)
    model = _resolve_agent_model(body.model)

    current_html = _decrypt_text(poster.html_content)
    prompt_history = json.loads(_decrypt_text(poster.prompt_history))

    user_msg = f"CURRENT POSTER HTML:\n{current_html}\n\nREQUESTED CHANGES:\n{body.prompt}"

    async def generate():
        import asyncio
        yield f"data: {json.dumps({'type': 'status', 'message': 'Updating your poster...'})}\n\n"

        task = asyncio.create_task(_call_claude_bg(
            client, POSTER_UPDATE_SYSTEM, user_msg, max_tokens=8000, model=model
        ))
        while not task.done():
            yield ":\n\n"
            await asyncio.sleep(3)
        new_html = task.result()

        new_html = new_html.strip()
        if new_html.startswith("```"):
            new_html = re.sub(r"^```(?:html)?\s*\n?", "", new_html)
            new_html = re.sub(r"\n?```\s*$", "", new_html)

        prompt_history.append(body.prompt)
        title_match = re.search(r"<title>(.*?)</title>", new_html, re.IGNORECASE)
        title = title_match.group(1) if title_match else _decrypt_text(poster.title)

        poster.title = _encrypt_text(title)
        poster.prompt_history = _encrypt_text(json.dumps(prompt_history))
        poster.html_content = _encrypt_text(new_html)
        poster.updated_at = datetime.now(timezone.utc)
        db.commit()

        yield f"data: {json.dumps({'type': 'result', 'poster_id': poster_id, 'title': title, 'html': new_html})}\n\n"
        yield f"data: {json.dumps({'type': 'done'})}\n\n"

    return StreamingResponse(generate(), media_type="text/event-stream")


@app.get("/posters", dependencies=[Depends(verify_api_key)])
async def list_posters(request: Request, db=Depends(get_db)):
    """List all posters for the current user."""
    current_user_id = getattr(request.state, "user_id", None)
    q = db.query(DBPoster)
    if current_user_id:
        q = q.filter(DBPoster.user_id == current_user_id)
    else:
        q = q.filter(DBPoster.user_id.is_(None))
    posters = q.order_by(DBPoster.updated_at.desc()).all()
    return {
        "posters": [
            {
                "id": p.id,
                "title": _decrypt_text(p.title),
                "prompt_count": len(json.loads(_decrypt_text(p.prompt_history))),
                "created_at": p.created_at.isoformat() if p.created_at else None,
                "updated_at": p.updated_at.isoformat() if p.updated_at else None,
            }
            for p in posters
        ]
    }


@app.get("/posters/{poster_id}", dependencies=[Depends(verify_api_key)])
async def get_poster(poster_id: str, request: Request, db=Depends(get_db)):
    """Get a specific poster with its full HTML content."""
    current_user_id = getattr(request.state, "user_id", None)
    poster = db.query(DBPoster).filter(DBPoster.id == poster_id).first()
    if not poster:
        raise HTTPException(status_code=404, detail="Poster not found")
    if current_user_id and poster.user_id and poster.user_id != current_user_id:
        raise HTTPException(status_code=403, detail="Access denied")
    return {
        "id": poster.id,
        "title": _decrypt_text(poster.title),
        "html": _decrypt_text(poster.html_content),
        "prompts": json.loads(_decrypt_text(poster.prompt_history)),
        "created_at": poster.created_at.isoformat() if poster.created_at else None,
        "updated_at": poster.updated_at.isoformat() if poster.updated_at else None,
    }




@app.delete("/posters/{poster_id}", dependencies=[Depends(verify_api_key)])
async def delete_poster(poster_id: str, request: Request, db=Depends(get_db)):
    """Delete a poster."""
    current_user_id = getattr(request.state, "user_id", None)
    poster = db.query(DBPoster).filter(DBPoster.id == poster_id).first()
    if not poster:
        raise HTTPException(status_code=404, detail="Poster not found")
    if current_user_id and poster.user_id and poster.user_id != current_user_id:
        raise HTTPException(status_code=403, detail="Access denied")
    db.delete(poster)
    db.commit()
    return {"deleted": True}


class PosterSaveHTMLRequest(BaseModel):
    html: str = Field(max_length=500000, description="Updated HTML content")


@app.patch("/posters/{poster_id}", dependencies=[Depends(verify_api_key)])
async def save_poster_html(poster_id: str, body: PosterSaveHTMLRequest, request: Request, db=Depends(get_db)):
    """Save manually edited HTML content for a poster."""
    current_user_id = getattr(request.state, "user_id", None)
    poster = db.query(DBPoster).filter(DBPoster.id == poster_id).first()
    if not poster:
        raise HTTPException(status_code=404, detail="Poster not found")
    if current_user_id and poster.user_id and poster.user_id != current_user_id:
        raise HTTPException(status_code=403, detail="Access denied")

    title_match = re.search(r"<title>(.*?)</title>", body.html, re.IGNORECASE)
    if title_match:
        poster.title = _encrypt_text(title_match.group(1))

    prompt_history = json.loads(_decrypt_text(poster.prompt_history))
    prompt_history.append("[Manual HTML edit]")
    poster.prompt_history = _encrypt_text(json.dumps(prompt_history))
    poster.html_content = _encrypt_text(body.html)
    poster.updated_at = datetime.now(timezone.utc)
    db.commit()
    return {"saved": True, "title": _decrypt_text(poster.title)}


@app.post("/posters/{poster_id}/duplicate", dependencies=[Depends(verify_api_key)])
async def duplicate_poster(poster_id: str, request: Request, db=Depends(get_db)):
    """Duplicate an existing poster."""
    current_user_id = getattr(request.state, "user_id", None)
    poster = db.query(DBPoster).filter(DBPoster.id == poster_id).first()
    if not poster:
        raise HTTPException(status_code=404, detail="Poster not found")
    if current_user_id and poster.user_id and poster.user_id != current_user_id:
        raise HTTPException(status_code=403, detail="Access denied")

    now = datetime.now(timezone.utc)
    orig_title = _decrypt_text(poster.title)
    new_poster = DBPoster(
        id=str(uuid.uuid4()),
        user_id=current_user_id,
        title=_encrypt_text(f"{orig_title} (Copy)"),
        prompt_history=poster.prompt_history,
        html_content=poster.html_content,
        created_at=now,
        updated_at=now,
    )
    db.add(new_poster)
    db.commit()
    return {
        "id": new_poster.id,
        "title": f"{orig_title} (Copy)",
    }
